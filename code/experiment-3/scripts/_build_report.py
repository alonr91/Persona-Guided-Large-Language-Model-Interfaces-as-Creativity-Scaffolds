# -*- coding: utf-8 -*-
"""Build the Experiment 3 report (.docx) — field hackathon, process-forward.

Long-conversation sample (>=10 user turns), OFF-TASK conversations 485 & 693 excluded
=> 18 conversations (13 treatment, 5 control). Length-normalised. Reports manipulation
check + persona fidelity (with Experiment-2 expert-validation context), the divergence-
to-convergence choreography (centrepiece), co-regulation by complementarity, and the
(null, length-normalised) product layer. Explicitly mapped to the thesis RQs.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 3'
FIG = f'{BASE}/report/figures'
OUT = f'{BASE}/report/Experiment3_in_exp2_style.docx'

def _load_rubric():
    """Load the regulated LLM-rubric contrast + audit (ported from Experiment 1)."""
    import csv as _csv
    cp = f'{BASE}/outputs/regulated_rubric_contrast.csv'
    ap = f'{BASE}/outputs/regulated_rubric_audit.csv'
    if not (os.path.exists(cp) and os.path.exists(ap)): return None
    con = {}
    for r in _csv.DictReader(open(cp, encoding='utf-8')):
        if r['hedges_g'] == '': continue
        con[r['criterion']] = dict(t=float(r['treatment_mean']), c=float(r['control_mean']),
                                   g=float(r['hedges_g']), p=float(r['p_value']))
    aud = {}
    for r in _csv.DictReader(open(ap, encoding='utf-8')):
        aud[r['criterion']] = dict(rab=r['A_B_pearson_r'], r2=r['length_bias_R2'])
    if len(con) < 12: return None
    rabs = [float(aud[k]['rab']) for k in aud if aud[k]['rab'] not in ('', None)]
    flags = [k for k in aud if aud[k]['r2'] not in ('', None) and float(aud[k]['r2']) > 0.30]
    return dict(con=con, aud=aud, rab_lo=min(rabs), rab_hi=max(rabs), len_flags=flags)
RUBRIC = _load_rubric()

def _load_user_rubric():
    """Load the USER-only behaviour-rubric contrast + audit (Experiment 1 'Option B')."""
    import csv as _csv
    cp = f'{BASE}/outputs/user_rubric_contrast.csv'
    ap = f'{BASE}/outputs/user_rubric_audit.csv'
    if not (os.path.exists(cp) and os.path.exists(ap)): return None
    con = {}
    for r in _csv.DictReader(open(cp, encoding='utf-8')):
        if r['hedges_g'] == '': continue
        con[r['criterion']] = dict(t=float(r['treatment_mean']), c=float(r['control_mean']),
                                   g=float(r['hedges_g']), p=float(r['p_value']))
    aud = {}
    for r in _csv.DictReader(open(ap, encoding='utf-8')):
        aud[r['criterion']] = dict(rab=r['A_B_pearson_r'], r2=r['length_bias_R2'])
    if len(con) < 6: return None
    rabs = [float(aud[k]['rab']) for k in aud if aud[k]['rab'] not in ('', None)]
    return dict(con=con, aud=aud, rab_lo=min(rabs), rab_hi=max(rabs))
URUBRIC = _load_user_rubric()

doc = Document()
normal = doc.styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11)

def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def P(t):
    p = doc.add_paragraph(t); p.paragraph_format.space_after = Pt(6); return p
def CAP(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = True; r.font.size = Pt(9.5)
    p.paragraph_format.space_after = Pt(8); return p
def FIGURE(fname, cap, width=6.1):
    path = f'{FIG}/{fname}'
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    CAP(cap)
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''; r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, c in enumerate(row):
            cells[i].text = ''; r = cells[i].paragraphs[0].add_run(str(c)); r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ================= TITLE =================
title = doc.add_heading('', level=0)
title.add_run('Experiment 3: Divergent and Convergent LLM Personas in a Field Design '
              'Hackathon: Process Replication and the Choreography of Creative Inquiry')

# ================= ABSTRACT =================
H1('Abstract')
P("Experiment 3 moves the dual-persona interface from Experiment 2 out of the lab and into the field. Two "
  "LLM personas, one divergent (Taylor) and one convergent (Alex), sit behind two send buttons. Design "
  "teams could use them or not, across a four-day hackathon on real problems. No questionnaire was "
  "collected, so we set aside perceived creativity (RQ1) and personality (RQ2); the logs let us ask how "
  "people ran the creative process and whether their ideas were more original (RQ3). We analyse the long, "
  "on-task conversations where the interface was genuinely used (at least ten user turns; 13 treatment, 5 "
  "control), and normalise every measure by length. The personas are identical to Experiment 2’s validated "
  "pair, and stayed sharply distinct even in long, multi-day conversations (g=1.19). Users ran a clear arc: "
  "they leaned on the divergent persona early (71% vs. 31% of first-half messages) and turned to the "
  "convergent one later. In effect they handed the work of converging to Alex while keeping idea-generation "
  "to themselves. A condition-blind LLM rubric scored the treatment conversations much higher on reframing, "
  "exploration, and evaluation, but a version that scored only the user’s own turns found no difference. So "
  "the scaffold works through the personas the user steers, not by changing the user. Originality did not "
  "carry over: no output measure separated the groups (idea rate g=−0.34; topic-controlled originality "
  "g=−0.07; both n.s.). The manipulation changes how people work, not the content of their ideas. The "
  "dual-persona scaffold reliably shapes the creative process even in messy field conditions; whether it "
  "makes the final ideas more original depends on having a task clean enough to measure.")

# ================= 1. INTRODUCTION =================
H1('1. Introduction and rationale')
P('Experiment 2 showed, under controlled laboratory conditions, that simultaneous, on-demand access to a '
  'divergent and a convergent LLM persona reshaped creative work: users credited the divergent persona, an '
  'explore-then-evaluate arc appeared in their questioning, and treatment portfolios were semantically more '
  'original than a single-model control. That study fixed one problem, ran in a single 20-minute session, and '
  'depended on a post-task questionnaire. Two questions follow. Do the behavioural signatures of the '
  'manipulation survive outside the laboratory, when people choose their own problems, work over days, and are '
  'not performing for an experimenter? And does the originality advantage, measured by comparing portfolios '
  'that all addressed the same prompt, generalise when every participant’s problem is different?')
P("Experiment 3 runs the same interface in a four-day creative hackathon. The trade is deliberate: we "
  "give up the tight control and the self-report measures of Experiment 2 to gain ecological validity and a "
  "kind of data the laboratory cannot produce, namely long, self-paced conversations in which switching "
  "between modes is a free user choice. Set against the thesis\u2019s research questions, the experiment is "
  "partial by design. RQ1 (perceived creativity support and ownership) and RQ2 (personality moderation) "
  "need the questionnaire and the Big-Five inventory; neither was collected in the field, so both fall out "
  "of scope here. What the logs can speak to is the manipulation check and the process layer of the "
  "evaluation, which the framework treats as first-class evidence about how persona-guided interaction "
  "reorganises creative work rather than as secondary metadata, together with RQ3, the output question of "
  "originality and diversity. The focus therefore moves from product to process, and lands squarely on the "
  "part of the thesis gap that asks whether differentiated personas reshape how people actually "
  "co-create.")
P('The contributions are three. First, a field replication showing that the persona manipulation and the '
  'explore-then-consolidate process behaviour transfer to naturalistic, self-paced creative work where users '
  'freely chose when to engage each role. Second, a within-conversation account of how users orchestrate two '
  'cognitive modes over long conversations, the divergence-to-convergence choreography and its complementarity '
  'signature, made visible only by field-scale data. Third, a length-normalised, topic-controlled originality '
  'methodology (idea-level extraction, challenge labelling, and residualised embeddings) for evaluating '
  'creative output when participants each solve a different problem, complemented by a condition-blind '
  'regulated LLM-rubric (ported from Experiment 1) that independently triangulates the process layer.')

# ================= 2. METHOD =================
H1('2. Method')

H2('2.1 Design and conditions')
P('Experiment 3 used the same two-condition, between-subjects design as Experiment 2, deployed as a field '
  'study at the level of the conversation. In treatment, two differentiated personas, Taylor (divergent) and '
  'Alex (convergent), were reachable through two send buttons beneath one shared chat thread (persona_id 1 and '
  '2). In control, both buttons routed to the same standard model with a minimal prompt and no role '
  'differentiation (persona_id 3 and 4). Unlike Experiment 2, participants were aware that the two treatment '
  'agents embodied different (divergent versus convergent) thinking styles; this realism comes at the cost of '
  'possible demand characteristics, which we treat as a limitation (Orne, 1962; Section 4.5) and which the '
  'control condition partly addresses: control participants faced two buttons routed to one identical model, so '
  'any expectation of "two styles" there could not be confirmed by behaviour. Because the analysable groups '
  'are small and unequal (the long, on-task sample is 13 treatment, 5 '
  'control; Section 2.3), all between-condition tests use Welch’s t and Hedges’ g, and the small control group '
  'is treated as a power limitation throughout.')
P('Statistical reporting. Effects are reported per comparison with simple tests (Welch’s t-tests with Hedges’ '
  'g for group contrasts and two-sided p-values) without multiplicity correction. Given the exploratory aims, '
  'the modest and unequal samples, and the fact that the key contrasts are pre-specified replications of '
  'Experiment 2 rather than a blind search, family-wise or false-discovery-rate corrections would be '
  'inappropriately conservative and are deliberately not applied; effect sizes accompany every p-value so '
  'magnitude can be weighed independently of significance.')

H2('2.2 Personas and task')
P('The persona manipulation was identical to Experiment 2 (system prompts in Appendix A): Taylor was prompted '
  'for multiple, analogical, non-evaluative exploration at temperature 0.8; Alex for criteria, comparison, and '
  'structured convergence at temperature 0.3; both shared task framing and response rules (including Hebrew '
  'language matching). Unlike Experiment 2’s single fixed prompt, participants brought and framed their own '
  'problems, drawn from ten real-world challenges spanning two themes: rehabilitation, trauma, and disability '
  '(e.g., the Ichilov rehabilitation hospital of the future; Polyron therapeutic sleep; NATAL’s language of '
  'distress; support for amputees and combat-injured users) and social cohesion (e.g., the Joint “Rikma” '
  'project on Jewish–Arab workplace inclusion; Tel Aviv community districts; Upper Galilee integration). Many '
  'participants worked from named user personas and Design-Thinking “how might we…” framings.')

H2('2.3 Data, sample, and preprocessing')
P('From the platform export we isolated the hackathon window (10–13 March 2025): 1,042 messages across 81 '
  'conversations. The group mapping above was applied to every message. Cleaning proceeded in three steps. '
  'First, conversations whose persona identifiers spanned both conditions (mixed) and non-substantive openings '
  '(greetings, single turns, test strings) were removed, leaving 40 single-condition conversations with at '
  'least four substantive user turns. Second, central to this report, the analysis is restricted to the long '
  'conversations, those with at least ten user turns; this threshold sits at a natural gap in the length '
  'distribution (treatment conversations cluster at ≤7 turns or ≥10) and reflects a deliberate scope decision, '
  'since the value of this field dataset is in sustained creative work and short exchanges contribute mostly '
  'noise. Third, a content audit of the long conversations removed two that were long but off-task, a user '
  'probing the system’s capabilities and a meta/identity exchange (both already flagged as fitting no '
  'challenge), because, however lengthy, they are not creative work; the remaining conversations carry a '
  'healthy idea density (median 0.5 extracted ideas per user turn). The resulting analysis sample is 18 '
  'conversations: 13 treatment and 5 control, all on-task. Conversations that span multiple days (three of the '
  'eighteen) are retained: long gaps are a normal feature of a multi-day event. Because the retained '
  'conversations still vary substantially in length (10–34 user turns), every measure is length-normalised '
  '(Section 2.4). All text was normalised to English with the Google Translate API (uniformly across '
  'conditions), and, with no questionnaires or participant identifiers, the unit of analysis is the '
  'conversation.')

H2('2.4 Measures')
P("Stance measurement. The personas are identical to Experiment 2’s, whose blind expert panel validated the divergent–convergent manipulation; we inherit that rather than re-validating (Section 3.1). To measure stance in the field logs, for the process analyses and to check the contract persisted, every turn is scored by a zero-shot stance classifier (multi-label NLI) on five divergent and five convergent markers, giving a divergent score D, a convergent score C, and a balance D−C; the classifier was calibrated against the first author’s hand ratings on a 60-turn sample (Pearson r=0.715); Appendix B and Figure B1 give the marker set, the entailment scoring, and the calibration.")
P('Process. From the logs we quantify how many messages users addressed to each persona, the persona they '
  'ended with, and (exploiting the long conversations) the within-conversation trajectory of stance (assistant '
  'D−C by conversation quartile), the share of messages to the divergent persona in the first versus second '
  'half, and question-asking (% of user messages containing “?”) by quarter and persona. Cross-persona '
  'brokerage moves are reported qualitatively. To trace process in semantic space we additionally embed every '
  'turn (bge-large-en-v1.5) and compute each user idea’s novelty relative to the participant’s own history and '
  'the user’s semantic accommodation to the persona’s preceding turn.')
P('Process-regulation rubric (ported from Experiment 1). As an independent, condition-blind triangulation of '
  'the process layer, we re-ran Experiment 1’s regulated LLM-rubric on the same 18 conversations. Each '
  'conversation was masked (persona names → Assistant_A / Assistant_B, so the scorer cannot infer condition) '
  'and scored 0–4 on twelve anchored process-regulation criteria (exploration opening, reframing quality, '
  'evaluative discipline, agency preservation, anchor management, co-regulation uptake, timing fit, '
  'implementation grounding, cognitive-load clarity, stance integrity, and two reverse-scored risks). Two '
  'paraphrased scorers (A strict, B paraphrased; gemini-3.1-flash-lite, temperature 0.15) return '
  'evidence-grounded JSON, reconciled by Experiment 1’s conservative adjudicator (on a disagreement of two '
  'points or more, the lower score is taken). Per the thesis’s Section 2.10 boundary, these are reported as '
  'proxy indicators of how the dialogue was regulated, not as externally judged creativity (Appendix D). To '
  'separate what the user does from what the dialogue contains, we additionally apply Experiment 1’s '
  'user-behaviour rubric (its “Option B”): the same masked transcript is read for context but only the '
  'participant’s own turns are rated, on six 0–4 creative-behaviour criteria (initiative, question richness, '
  'proposal specificity, yes-and uptake, reframing, engagement depth), with every score quoting a user turn.')
P('Product (idea portfolios). We extracted idea portfolios with Experiment 1’s multi-agent pipeline on '
  'user-only English text (Agent 1 extracts user-originated, evidence-grounded ideas; Agent 2 consolidates '
  'near-duplicates → fluency; Agent 4 categorises → flexibility; Agent 5 computes embedding-centroid '
  'originality). Because the retained conversations still vary widely in length (10–34 user turns), all product '
  'measures are length-normalised: the quantitative measure is the idea-generation rate (ideas per user turn) '
  'and peer-relative originality is recomputed within the sample. A within-treatment dose-response relates how '
  'users split attention between the personas to these measures. To guard against the topic confound from '
  'heterogeneous challenges, each conversation was labelled to its challenge (Gemini, hand-validated) and a '
  'topic-residualised embedding originality computed. As validated semantic-creativity checks we also compute '
  'forward flow (Gray et al., 2019) and divergent semantic integration (Johnson et al., 2022) on each user '
  'idea, tagged by the persona engaged.')
CAP('Table 1. Long, on-task analysis sample (≥10 user turns; two off-task conversations removed) and challenge '
    'distribution. Control is concentrated in the rehabilitation challenge; elsewhere control n ≤ 1.')
table(['Challenge (label)', 'Treatment', 'Control'],
 [['Ichilov rehabilitation (future)', '5', '3'],
  ['Joint Rikma (Jewish–Arab inclusion)', '3', '0'],
  ['NATAL (trauma / distress)', '2', '1'],
  ['Polyron (therapeutic sleep)', '1', '1'],
  ['TA-South / TA-East community', '2', '0'],
  ['Total', '13', '5']])

# ================= 3. RESULTS =================
H1('3. Results')

H2('3.1 The manipulation, inherited from Experiment 2, held in the field')
P("The personas are identical to Experiment 2’s, where a blind expert panel established the divergent–convergent contrast (Taylor M=5.07 vs. Alex M=2.67 on a 1–7 scale, p=.032). We inherit that validation rather than re-running it, and use the field logs only to confirm the contract survived deployment conditions Experiment 2 could not test. It plainly did: across assistant turns of the long conversations the divergent persona scored far higher on divergent stance than the convergent one (Taylor M=0.643 vs. Alex M=0.318; g=1.25, p<.001), along the targeted axis (per-marker breakdown in Appendix B), while in control the two routes did not differ (g=0.28, p=.30, n.s.), ruling out button position or expectation as the source (Figure 1). The stance classifier behind these scores, calibrated against the first author’s hand ratings (r=0.715), is also the instrument for the process analyses that follow.")
P("What the field adds beyond Experiment 2 is a test of persistence. A persona is a stance contract that must hold across turns (Tseng et al., 2024), and drift toward generic behaviour as context accumulates is a recognised internal-validity threat that 20-minute lab sessions cannot probe. The contract held: the divergent–convergent separation stayed large in the late half of long, sometimes multi-day conversations (g=1.19, p<.001), only modestly below the early half (g=1.35). The one slippage is a partial softening of the divergent persona toward the centre late in long sessions, with the convergent persona stable, so the stance-reinforcement mechanism largely survives extended, real-world use.")
FIGURE('fig1_manipulation_check.png',
       'Figure 1. Manipulation check. In treatment the divergent persona (Taylor, green) is strongly divergent '
       'and the convergent persona (Alex, purple) is balanced; in control the two routes do not differ. Means '
       '± 95% CI over assistant turns.')

H2('3.2 Persona preference: the divergent persona is the draw')
P('Users gravitated to the divergent persona, and in these sustained sessions the preference was pronounced. '
  'Treatment users addressed Taylor more than twice as often as Alex (11.7 vs. 4.5 messages per conversation; '
  'g=1.33, p=.002), whereas in control, where the two buttons were behaviourally identical, message counts were '
  'balanced and if anything tilted the other way (4.8 vs. 6.4; g=−0.38, n.s.). The balanced control, where the '
  'two buttons were the same model, shows the skew tracks the agents’ actual behaviour rather than button '
  'position; participants were aware of the two roles, so awareness may also contribute, but expectation alone '
  'did not move control users, who could not confirm a difference that was not there. The same asymmetry '
  'appeared at session end: 85% of treatment sessions (11/13) ended on the divergent persona versus 40% in '
  'control (2/5; Figure 2).')
FIGURE('fig2_persona_engagement.png',
       'Figure 2. Persona engagement. Treatment users address the divergent persona (green) far more than the '
       'convergent persona (purple); control users, facing two identical models, are balanced. Means ± 95% CI.')

H2('3.3 The choreography of inquiry: divergence early, convergence late')
P('The long conversations reveal what a single short session cannot: a within-conversation movement from '
  'divergent exploration to convergent consolidation, enacted through which persona the user chooses to '
  'engage. This is the macro divergence-to-convergence trajectory that process accounts of creativity describe '
  '(the Double Diamond; Sawyer, 2021), here observed directly and built from meso-level stance shifts rather '
  'than imposed as a stage script. Two views establish it. First, treatment users front-loaded the divergent '
  'persona: 71% of first-half messages went to Taylor versus 31% in control (g=1.13, p=.10). Second, the '
  'stance balance of the assistant turns users engaged fell across conversation quartiles in treatment (D−C: '
  '0.45→0.38→0.19→0.22), settling into a markedly more convergent register late, while control stayed '
  'comparatively flat and erratic with no ordered trend (Figure 3). This staged cadence emerged without any '
  'instruction about what the personas were.')
P('How users reach that convergence is what matters: they reallocate attention to the convergent '
  'persona without becoming convergent themselves: co-regulation by complementarity rather than alignment '
  '(Fusaroli & Tylén, 2016; Reitter & Moore, 2014). Users were no more convergent when addressing Alex than '
  'Taylor (user balance 0.29 vs. 0.27, g=0.08, n.s.) and did not mirror the assistants’ turn-level stance; a '
  'turn-level semantic analysis agrees, with users’ ideas growing progressively more independent of the '
  'personas’ content over the session (accommodation r=−0.15, p=.03, treatment only) while converging on '
  'their own developing thread (novelty-to-history r=−0.18, p=.01). Instead of being absorbed by either '
  'agent, users delegate consolidation to the convergent persona while keeping their own generative footing, a '
  'division of cognitive labour the dual-persona interface uniquely affords, and one that contrasts with the '
  'stance alignment seen under assigned personas in Experiment 1.')
FIGURE('fig3_trajectory.png',
       'Figure 3. Divergence-to-convergence trajectory. The stance balance (D−C) of the assistant turns '
       'treatment users engage declines steadily across conversation quartiles; control is flat. Means ± 95% '
       'CI; conversations with ≥4 assistant turns.')
P('Qualitatively, treatment users treated the system as a small creative team, brokering between the personas '
  'on roughly half of conversations, versus never in control. Representative moves include cross-persona '
  'referencing (“what do you think of Taylor’s answer?”; “Alex, can you tell Taylor what a nurse does?”), '
  'mode-targeted delegation (“Taylor, please give me ten more ideas to make Dan’s rehab environment more '
  'comfortable”), and perspective-taking (“in your opinion, what would Alex answer to the same question?”). '
  'Directing specialists and asking them to react to one another is rare with a single generic assistant and '
  'signals active user ownership of the process.')
P('Question-asking did not follow Experiment 2’s divergent-early / convergent-late cadence (Figure 4): over '
  'quarters Q2–Q4 the treatment personas did not out-question control; if anything they asked fewer (divergent '
  'g=−0.41; convergent g=−0.33, n.s.), with control question-rates rising over the session. The one '
  'interpretable signal is consistent with complementarity: users query the convergent persona more than the '
  'divergent one overall (~45% vs. ~35% of messages), interrogating the evaluator while drawing ideas from the '
  'generator. The divergence-to-convergence signature is carried by which persona users engage '
  '(Figure 3), not by their question rate, a measure that separated the conditions in the lab but not here.')
FIGURE('fig4_question_rate.png',
       'Figure 4. Question-asking (% of user messages containing “?”) by conversation quarter, by the persona '
       'addressed in treatment (divergent green, convergent purple) and overall in control (blue dashed). '
       'Unlike Experiment 2, no divergent-early / convergent-late cadence appears and treatment does not '
       'exceed control over Q2–Q4.')

H2('3.4 Idea portfolios: originality, fluency, and the (null) role of interaction style')
P('Addressing RQ3, idea portfolios were extracted with Experiment 1’s pipeline and recomputed within the '
  'sample. Once length-normalised, the dual-persona interface did not change product-level creativity on any '
  'measure: the idea-generation rate did not differ (0.50 vs. 0.56 ideas per user turn; g=−0.34, n.s.), nor '
  'did topic-controlled (challenge-residualised) originality (g=−0.07, n.s.), idea-centroid originality '
  '(g=−0.27, n.s.), within-portfolio diversity (g=−0.14, n.s.), or even the raw idea count (8.2 vs. 6.2; '
  'g=0.65, p=.12, n.s.; Appendix C, Table C1). Within treatment, how users split attention between the '
  'personas predicted no creativity measure: the convergent share was uncorrelated with idea rate (r=−0.21), '
  'originality (idea-centroid r=+0.17; topic-residualised r=−0.28), and diversity (r=−0.44; all n.s.). The '
  'lever, if any, is the presence of the divergent generator, not how users ration the two modes.')

H2('3.5 Thematic analysis: interacting with the divergent versus convergent persona')
P('To characterise how users actually worked with the two agents, we thematically analysed the 13 long '
  'treatment conversations (Braun & Clarke, 2006). Across 211 user turns, 152 addressed to the divergent '
  'persona (Taylor) and 59 to the convergent persona (Alex), each turn was coded by the persona engaged and the '
  'kind of request it made; three themes distinguish the two interactions (Table 2). Because participants knew '
  'which agent embodied which style, these are patterns of informed, deliberate use. The 152:59 split is '
  'informative on its own: the divergent persona was the default creative workspace and the convergent one as '
  'a selectively recruited specialist.')
CAP('Table 2. Request types by persona addressed (share of that persona’s user turns; multi-label keyword '
    'coding, treatment long sample). Novelty-pressing concentrates on the divergent persona; evaluation, '
    'specification, and reality-testing on the convergent one.')
table(['Request type (user turn)', 'Divergent (Taylor)', 'Convergent (Alex)'],
 [['Novelty-pressing ("more creative", "not good enough")', '14%', '5%'],
  ['Idea generation ("give me / more ideas")', '9%', '10%'],
  ['Elaborate / detail a direction', '6%', '14%'],
  ['Evaluate / select ("rate", "which is best")', '3%', '14%'],
  ['Specify / technical ("what device", "where", "cost")', '2%', '14%'],
  ['Reality-check ("does this exist?")', '1%', '5%']])
P("Theme 1: the divergent persona as idea engine and novelty foil. Users came to Taylor to generate and broaden, and distinctively to press for novelty, which was about three times more frequent to Taylor than to Alex (14% vs. 5% of turns). They rejected the ordinary (“it’s not creative enough”; “prosthetics are an option already, maybe think in another direction”; convs 731, 667) and pushed Taylor toward speculative leaps such as brain-computer helmets and tissue regeneration. This is the exploratory stance enacted: generating alternatives and resisting premature closure (cf. White, 2003, on expansion resources).")
P("Theme 2: the convergent persona as evaluator, specifier, and reality-tester. Users brought Alex a categorically different set of demands. Evaluation and selection were about five times more frequent to Alex (14% vs. 3%; e.g. “rate the solutions 1–5 for feasibility, effectiveness, and relevance”, conv 595), and technical specification roughly seven times (14% vs. 2%; “where would you place the infrared sensor, and what is the minimum distance?”, convs 733, 737). Reality-testing also concentrated on Alex (“does this exist?”), as did requests to make a chosen idea concrete. This is convergent thinking: criteria, comparison, feasibility, and closure, with the convergent agent serving as an evaluative anchor and authority for grounding (§2.4.5).")
P("Theme 3: orchestration, routing, and the explore-to-converge handoff. The two profiles were not merely different but managed as a division of labour: users routed generation to Taylor and judgment or grounding to Alex, brokered between them (“alex, what you think about that?” about a Taylor proposal, then “taylor, how can we do it?”; conv 523), and ran a recurring generate-with-Taylor then evaluate-and-specify-with-Alex handoff (convs 595, 733). The user, not either agent, assembles the arc.")
P("These themes give the qualitative grain behind the quantitative results and converge on one point: the user, not either agent, integrates generation and judgment. One caveat: the evaluation/selection code leaned on conv 595, the only session with explicit 1–5 rating, whereas specification and reality-testing spread across conversations 483, 733, and 737, so “the convergent persona as specifier” is the better-supported reading.")

if RUBRIC:
    c = RUBRIC['con']
    H2('3.6 Two LLM-rubric proxies: the dialogue transforms, the user does not')
    P('A final, independent line of evidence triangulates the process findings. Experiment 1’s regulated '
      'LLM-rubric (§2.4; Appendix D), run condition-blind on the masked transcripts, scores each conversation '
      '0–4 on twelve dialogue-regulation criteria. Because it never sees the condition and quotes verbatim '
      'evidence for every score, it is an instrument independent of the stance classifier and the behavioural '
      'counts above. Treatment conversations scored markedly higher on the core regulation criteria (Table 3): '
      f'reframing quality (M={c["reframing_quality"]["t"]:.2f} vs. {c["reframing_quality"]["c"]:.2f}; '
      f'g={c["reframing_quality"]["g"]:+.2f}, p={c["reframing_quality"]["p"]:.3f}), and evaluative discipline, '
      'exploration opening, and anchor management (all g≥+1.0, p<.01). Agency preservation, co-regulation '
      'uptake, stance integrity, and implementation grounding point the same way without reaching significance '
      'against the five controls, and premature-convergence risk is lower in treatment. All ten criteria with '
      'a directional design prediction move the predicted way; the other two (cognitive-load clarity, '
      'runaway-divergence risk) are flat or slightly elevated. The contrast is not a length artefact '
      '(conversation length explains under a third of the variance on every criterion, only evaluative '
      'discipline approaching that bound) and the two paraphrased scorers agree closely '
      f'(A–B r={RUBRIC["rab_lo"]:.2f}–{RUBRIC["rab_hi"]:.2f}; full audit in Appendix D).')
    CAP('Table 3. Regulated LLM-rubric (ported from Experiment 1): condition-blind, evidence-grounded 0–4 '
        'process-regulation scores; treatment vs control (Welch, Hedges’ g; no multiplicity correction). The '
        'two risk criteria are reverse-scored (higher = worse), so treatment’s lower premature-convergence is '
        'favourable.')
    _PRETTY = [('exploration_opening','Exploration opening'), ('reframing_quality','Reframing quality'),
               ('evaluative_discipline','Evaluative discipline'), ('agency_preservation','Agency preservation'),
               ('anchor_management','Anchor management'), ('coregulation_uptake','Co-regulation uptake'),
               ('timing_fit','Timing fit'), ('implementation_grounding','Implementation grounding'),
               ('cognitive_load_clarity','Cognitive-load clarity'), ('stance_integrity','Stance integrity'),
               ('premature_convergence_risk','Premature-convergence risk (rev.)'),
               ('runaway_divergence_risk','Runaway-divergence risk (rev.)')]
    table(['Regulation criterion (0–4)', 'Treatment', 'Control', 'g', 'p'],
          [[lab, f"{c[k]['t']:.2f}", f"{c[k]['c']:.2f}", f"{c[k]['g']:+.2f}", f"{c[k]['p']:.3f}"]
           for k, lab in _PRETTY])
    if URUBRIC:
        u = URUBRIC['con']
        P('Holding the same instrument to the user alone tells a different story. Experiment 1’s user-behaviour '
          'rubric (Appendix D) reads the same masked transcript but scores only the participant’s own '
          'turns, quoting user text only, on six 0–4 creative-behaviour criteria. Here no criterion separates '
          f'the conditions (Table 4): every effect is small and non-significant (the largest a g={max(u[k]["g"] for k in u):+.2f} '
          f'on yes-and uptake, the smallest p={min(u[k]["p"] for k in u):.2f}), while the two paraphrased scorers '
          f'agree almost perfectly (A–B r={URUBRIC["rab_lo"]:.2f}–{URUBRIC["rab_hi"]:.2f}) with no length '
          'confound. This is a well-measured null: the participants’ own creative behaviour is statistically '
          'indistinguishable between the dual-persona and single-model conditions.')
        CAP('Table 4. User-behaviour rubric (Experiment 1 “Option B”): the same masked transcripts scored on the '
            'participant’s turns only (0–4; treatment vs control, Welch, Hedges’ g). All six criteria n.s.')
        _UPRETTY = [('user_initiative','Initiative'), ('user_question_richness','Question richness'),
                    ('user_proposal_specificity','Proposal specificity'), ('user_acceptance_yes_and','Yes-and uptake'),
                    ('user_reframing','Reframing'), ('user_engagement_depth','Engagement depth')]
        table(['User criterion (0–4)', 'Treatment', 'Control', 'g', 'p'],
              [[lab, f"{u[k]['t']:.2f}", f"{u[k]['c']:.2f}", f"{u[k]['g']:+.2f}", f"{u[k]['p']:.3f}"]
               for k, lab in _UPRETTY])
    P('Read side by side, the two rubrics localise the effect: the dialogue is transformed while the user’s own behaviour '
      'is not, so the gains live in the conversation the personas supply rather than in the user (developed in '
      '§4.2). Both are proxy indicators of observable behaviour, not creativity scores (§2.10), and they '
      'reinforce the behavioural choreography (§3.3) even as the product-originality layer (§3.4) stays null. '
      'The one part of Experiment 1’s validation stack not reproduced here is cross-model robustness: both '
      'rubrics run on a single model family (Gemini), so model variance is uncontrolled.')

# ================= 4. DISCUSSION =================
H1('4. Discussion')

H2('4.1 What transfers from laboratory to field')
P("Taken together over the long, on-task sessions, Experiments 2 and 3 separate two kinds of claim. The "
  "process claims hold up well, and if anything look sharper in deep field sessions: the personas stay "
  "reliably distinct (g=1.25) and keep their contract over long context, users clearly prefer the divergent "
  "mode (g=1.33), and engagement follows a divergence-to-convergence arc. A "
  "condition-blind LLM rubric ported from Experiment 1 (§3.6) independently rates treatment higher on "
  "reframing, exploration, and evaluative discipline, a third method agreeing with the stance classifier and "
  "the behavioural counts. These "
  "show up under a tightly controlled single task and a messy four-day hackathon alike, and here they come "
  "from people working on their own real design problems. The product claim is shakier. Once we normalise "
  "for length, no measure of idea rate or originality differs (RQ3), and the shared-task originality "
  "advantage from Experiment 2 does not carry over to heterogeneous field problems with a small control "
  "group. What the interface reliably contributes is to the creative process; whether it improves the "
  "creative product depends on the task it is used for.")

H2('4.2 How users orchestrate the agents')
P("The clearest payoff of the field design is that the macro arc becomes directly observable. A 20-minute "
  "lab session shows only a snapshot of how people use the two modes; a multi-day hackathon lets us watch "
  "the same users steer from exploration to consolidation inside one piece of real work, moving messages "
  "from the divergent to the convergent persona as it matures (Figure 3). Observing the trajectory directly, "
  "rather than inferring it from a stage model, matters because multi-timescale accounts treat macro phases "
  "as something that emerges from many small meso-level stance shifts (Sawyer, 2021; Dorst & Cross, 2001); "
  "here the interface supplies the stances and the user assembles the arc. No one instructed users to explore "
  "first and converge later, so its appearance in self-paced field use is stronger evidence than any lab "
  "analogue that the interface scaffolds genuine creative-process dynamics rather than task compliance.")
P("What the arc reveals is orchestration, and the two rubrics locate where the gains live (Figure 5). The "
  "improvement is supplied by the personas and assembled by the user, not internalised: this is co-regulation "
  "by complementarity rather than alignment (§2.3.3; Fusaroli & Tylén, 2016; Reitter & Moore, 2014), with "
  "exploratory and evaluative labour distributed across the two agents instead of mirrored by the user, which "
  "is why the user-side rubric stays flat while the dialogue-side rubric moves.")
FIGURE('fig5_orchestration.png',
       'Figure 5. Augmentation by orchestration: the user works at the meta level (allocate, switch, broker, '
       'integrate) while the personas perform the object-level creative work, so the dialogue-level rubric '
       'rises sharply but the user-only rubric does not.')
P("Read as a distributed cognitive system (Hollan et al., 2000), the user and the two agents divide the work. "
  "The agents carry the object-level operations (generating alternatives, applying criteria) while the user "
  "keeps the meta-level regulation: when to widen, when to narrow, and which agent to engage. In the "
  "vocabulary of regulated learning this is a shift from self-regulation toward socially shared regulation, "
  "with the human as the regulator (Järvelä & Hadwin, 2013). The behaviour fits: users front-load the "
  "divergent agent and turn to the convergent one later, broker between them and hand work off (§3.5), "
  "and stay semantically independent of both (accommodation r=−0.15). The divergence-to-convergence arc "
  "that process models describe (Dorst & Cross, 2001; Sawyer, 2021) is thus assembled as an allocation "
  "pattern over a control surface, not run as an internal cycle.")
P("Mechanistically, the interface externalises metacontrol. Creative work requires continually tuning "
  "cognitive flexibility against persistence (Sowden et al., 2015); a single blended assistant leaves that "
  "tuning inside the user’s head, whereas two differentiated, user-invokable personas let the user tune the "
  "system’s stance by choosing an agent. This is the socio-technical scaffold the conceptual model "
  "anticipates: stances made legible, differentiable, and user-invokable reorganise the conditions of "
  "co-creation without changing the user’s internal cognition (§2.11). The interface thus augments by "
  "orchestration: it gives users orchestratable cognitive resources and the control surface to deploy them, "
  "and they remain the integrator and author. Two cautions follow: if the gains live in the orchestration "
  "they may not persist once the scaffold is removed, and five control conversations cannot rule out a modest "
  "real user-side change. Both invite a longitudinal, better-powered test.")

H2('4.3 Why output-level originality did not transfer')
P("The product null (RQ3) comes from the field design, not from any failure of the scaffold, which the "
  "strong manipulation and process results rule out. Two things combine. First, the conversations are very "
  "uneven: portfolios run from a handful of ideas to dozens, so raw counts and centroid originality mostly "
  "track length and disappear once we normalise to a per-turn rate. Second, the problems are heterogeneous "
  "and there are only five control conversations, which leaves the between-condition comparisons badly "
  "underpowered. A turn-level semantic analysis then shows where the effect lives and where it does not. "
  "The personas differ sharply in stance (g=1.25) but not in the semantic novelty of the content they "
  "introduce (g=−0.11, n.s.). The validated creativity measures agree: ideas produced with the divergent "
  "versus the convergent persona do not differ in forward flow (Gray et al., 2019) or divergent semantic "
  "integration (Johnson et al., 2022), and users show none of the anchoring to early ideas that "
  "design-fixation accounts warn about (Jansson & Smith, 1991). So the manipulation acts on stance and "
  "process, not on the geometry of ideas: it changes how users explore, switch, and consolidate without "
  "moving their ideas around in semantic space, which is exactly why an embedding-based originality measure "
  "stays put no matter the statistical power. Read plainly, the originality advantage depends on conditions "
  "that allow clean measurement, namely a shared task, idea-level extraction, and an adequate control "
  "group. Here it is bounded, not refuted.")

H2('4.4 Design implications')
P('The process results yield concrete guidance for creativity-support systems. Expose cognitive roles as '
  'separable, user-selectable agents rather than blending them into one assistant; users discover and exploit '
  'the distinction on their own. Keep both modes simultaneously visible with low-friction switching, so the '
  'natural explore-then-commit cadence can play out within a single piece of work. And because users delegate '
  'convergence to the convergent persona rather than adopting it themselves, pair that persona with light, '
  'user-owned commitment prompts at late-session checkpoints (“name the decision,” “justify the choice”) to '
  'convert delegated consolidation into authored conclusions. These behaviours appear both when the '
  'roles are withheld (Experiment 2) and when users are aware of them (Experiment 3, here), suggesting the '
  'role differentiation is useful with or without explicit labelling: a first read, across the two studies, on '
  'the labelling question Experiment 2 left open.')

H2('4.5 Limitations')
P('Restricted to long, on-task conversations. By design the analysis sample is the 18 long conversations '
  '(≥10 user turns) that carry genuine creative content; short and off-task exchanges were set aside. This '
  'sharpens the process picture and matches the dataset’s purpose, but the effects describe sustained creative '
  'sessions, not casual or one-shot use.')
P('Small, unequal control group. Five control conversations limit power for every between-condition test and '
  'are the chief reason the product layer cannot be resolved; the product results are best read as '
  'underpowered nulls rather than evidence of no effect. Field allocation did not permit balance.')
P("Manipulation inherited, not re-validated. The divergent–convergent manipulation is inherited from Experiment 2’s expert validation rather than re-established here; the field stance classifier (calibrated against the first author’s hand ratings, r=0.715) is a proxy instrument, not a fresh expert panel. A field expert panel would strengthen it but was not feasible.")
P('Participant awareness of the roles. Unlike Experiment 2, participants knew the two agents embodied '
  'divergent and convergent styles, so the persona preference and engagement patterns may partly reflect '
  'expectations (demand characteristics; Orne, 1962) rather than purely discovered usefulness. The balanced '
  'control condition mitigates this: knowing there were "two agents" did not make control users favour either '
  'identical button, but a blinded field arm would be needed to separate discovery from expectation fully.')
P('No subjective or personality data. Without a questionnaire, Experiment 3 cannot address RQ1 (perceived '
  'creativity, authorship) or RQ2 (personality moderation), and cannot triangulate behaviour against '
  'self-report.')
P('Proxy instruments and translation. The originality analysis depends on automatic challenge labels (hand-'
  'validated) and idea extraction; with no participant identifiers, sessions cannot be linked to individuals; '
  'and 41% of messages were machine-translated (uniformly across conditions), which may attenuate subtle '
  'stance cues.')

H2('4.6 Conclusion')
P("Experiment 3 is a field replication of the dual-persona LLM interface in a creative design hackathon, "
  "read from the long, on-task conversations where the interface was genuinely used. The personas, validated in Experiment 2, held their contract in the field and over long, accumulating context. The "
  "process layer replicated and went further: users strongly preferred and ended with the divergent "
  "persona, and they traced a divergence-to-convergence arc, shifting attention to the convergent agent "
  "while handing it the work of consolidation instead of taking that on themselves. No one had instructed "
  "them to work this way. On RQ3, the interface produced no detectable advantage in the originality or "
  "diversity of idea portfolios in the field (idea rate g=−0.34; topic-controlled originality g=−0.07; "
  "both n.s.). A turn-level semantic analysis puts this down to the manipulation acting on stance and "
  "process rather than on the geometry of ideas, made harder to detect by problem heterogeneity and limited "
  "power, so the result is bounded rather than a refutation. RQ1 and RQ2 stay open for a future field "
  "deployment that brings back a subjective layer and a larger, balanced control group. The practical "
  "lesson sharpens the one from Experiment 2: structurally differentiated, user-switchable personas are a "
  "robust way to scaffold the creative process, reliably moving users between exploration and "
  "consolidation, while their effect on creative products depends on having a task that lets those "
  "differences be measured.")

# ================= APPENDICES =================
doc.add_page_break()
H1('Appendix A. Persona system prompts')
P('The persona manipulation was identical to Experiment 2; the full divergent (Taylor, T=0.8), convergent '
  '(Alex, T=0.3), and control system prompts are reproduced verbatim in Appendix A of the Experiment 2 report '
  'and are not duplicated here. The only deployment differences were operational: participant-chosen problems '
  'in place of the fixed prompt, and self-paced multi-day sessions in place of a single 20-minute session.')

H1('Appendix B. Stance-classification pipeline and calibration')
P("Stance is measured turn by turn with a zero-shot natural-language-inference (NLI) classifier "
  "(MoritzLaurer/deberta-v3-base-zeroshot-v2.0). The turn is treated as the premise and ten short stance "
  "statements as hypotheses: five describing divergent moves and five convergent. The divergent markers "
  "are broadening the search with alternatives or new angles; asking an open exploratory “what if” "
  "question; reframing the problem or inviting multiple interpretations; using analogy or metaphor to "
  "expand thinking; and keeping options open rather than ranking or selecting. The convergent markers are "
  "articulating explicit criteria or constraints; comparing options and recommending one; prioritising or "
  "ranking ideas; offering structured stepwise planning; and critiquing or identifying weaknesses.")
P("Each hypothesis is scored independently (multi-label), so a turn may register as divergent, convergent, "
  "both, or neither. For a turn t and hypothesis h the NLI head yields an entailment logit and a "
  "contradiction logit, and the marker score is their two-way softmax, e(t, h) = exp(z_ent) / (exp(z_ent) "
  "+ exp(z_con)), read as the probability that the turn supports h and bounded in [0, 1]. The turn’s "
  "divergent and convergent scores are the means over the two marker sets, D(t) = (1/5) Σ e(t, h) over the "
  "five divergent hypotheses and C(t) = (1/5) Σ e(t, h) over the five convergent ones, and its stance "
  "balance is b(t) = D(t) − C(t), bounded in [−1, +1] and positive for a divergent turn, negative for a "
  "convergent one (Figure B1). The conversation-, half-, and quartile-level stance values in the results "
  "are means of b(t) over the relevant assistant or user turns.")
P("The instrument was calibrated against a stratified 60-turn sample that the first author hand-rated on a "
  "five-point divergent–convergent scale; the classifier balance correlated with those ratings at Pearson "
  "r = 0.715. Per-marker contrasts between the treatment personas (Taylor vs. Alex) confirm the "
  "manipulation acted across the marker set rather than resting on one or two markers (Table B1).")
FIGURE('figB1_stance.png',
       'Figure B1. How one turn is scored. The turn is tested against ten stance hypotheses by a zero-shot '
       'NLI model; each returns an entailment probability e(t, h). The five divergent and five convergent '
       'scores are averaged into D(t) and C(t), and their difference is the stance balance b(t) = D(t) − '
       'C(t), running from −1 (fully convergent) to +1 (fully divergent).')
CAP('Table B1. Per-marker stance contrast, treatment Taylor (divergent) vs. Alex (convergent); positive g = '
    'stronger in the divergent persona.')
table(['Stance marker', 'Taylor', 'Alex', 'g'],
 [['Asks open “what if” question (D)', '0.71', '0.20', '+1.45'],
  ['Reframes / invites interpretations (D)', '0.66', '0.28', '+1.15'],
  ['Proposes alternatives / new angles (D)', '0.71', '0.33', '+1.13'],
  ['Uses analogy or metaphor (D)', '0.54', '0.15', '+1.04'],
  ['Keeps options open (D)', '0.70', '0.48', '+0.75'],
  ['Articulates explicit criteria (C)', '0.14', '0.40', '−0.96'],
  ['Offers structured stepwise planning (C)', '0.29', '0.47', '−0.53'],
  ['Compares options and recommends (C)', '0.37', '0.25', '+0.32']])

H1('Appendix C. Idea-portfolio pipeline, normalisation, and dose-response')
P('Idea portfolios were extracted with Experiment 1’s multi-agent pipeline on user-only English text. Agent 1 '
  '(a local instruction-tuned model under the Experiment 1 extractor contract) extracts only user-originated, '
  'evidence-grounded ideas, filtering challenge-restatements, assistant-echoes, meta-questions, and reactions; '
  'Agent 2 consolidates near-duplicates within a conversation (fluency = canonical-idea count); Agent 4 '
  'categorises canonical ideas across the corpus (flexibility); Agent 5 computes idea-centroid originality. '
  'Because conversations vary widely in length, quantitative measures are reported as rates and peer-relative '
  'originality is recomputed within the sample. A within-treatment dose-response (Pearson r, no correction) '
  'relates interaction style (convergent share, switch rate, timing of convergent engagement) to each measure. '
  'Topic-residualised embedding originality (challenge labels, hand-validated) serves as a cross-check, and '
  'forward flow (Gray et al., 2019) and divergent semantic integration (Johnson et al., 2022), computed per '
  'user idea and tagged by persona, as validated semantic-creativity checks (no persona difference). Outputs: '
  'outputs/idea_portfolio_exp1.csv, long_portfolio.csv, dose_response_normalized.csv, '
  'originality_topic_controlled.csv.')
CAP('Table C1. Length-normalised idea-portfolio measures over the long, on-task sample (treatment vs control; '
    'Hedges’ g). Control n=5; all n.s.')
table(['Measure', 'Treatment', 'Control', 'g'],
 [['Idea-generation rate (ideas / user turn)', '0.50', '0.56', '−0.34'],
  ['Fluency (idea count, raw)', '8.23', '6.20', '+0.65'],
  ['Originality (topic-residualised)', '1.05', '1.05', '−0.07'],
  ['Originality (idea-centroid)', '0.19', '0.20', '−0.27'],
  ['Within-portfolio idea diversity', '0.32', '0.33', '−0.14']])
CAP('Table C2. Within-treatment dose-response (n=13; Pearson r): convergent share predicts no creativity '
    'measure.')
table(['Predictor → outcome', 'r', 'p'],
 [['Convergent share → idea rate', '−0.21', '.48'],
  ['Convergent share → originality (idea-centroid)', '+0.17', '.57'],
  ['Convergent share → originality (topic-residualised)', '−0.28', '.35'],
  ['Convergent share → within-portfolio diversity', '−0.44', '.13']])

if RUBRIC:
    _flag = ', '.join(f.replace('_', ' ') for f in RUBRIC['len_flags']) or 'none'
    H1('Appendix D. Regulated LLM-rubric process proxy (ported from Experiment 1)')
    P('This proxy reproduces Experiment 1’s regulated LLM-rubric (its os_pipeline/regulated module). Each of '
      'the 18 long, on-task conversations was masked (persona names → Assistant_A / Assistant_B; the condition '
      'is never shown) and scored as a whole on twelve anchored 0–4 process-regulation criteria: exploration '
      'opening, reframing quality, evaluative discipline, agency preservation, anchor management, co-regulation '
      'uptake, timing fit, implementation grounding, cognitive-load clarity, stance integrity, and two '
      'reverse-scored risks (premature-convergence, runaway-divergence). Two paraphrased scorers, A (strict, '
      '“transcript analysis instrument, not a creative judge”) and B (paraphrased), each return one bundled, '
      'evidence-grounded JSON object per conversation (every non-null score carries a verbatim quote; '
      'inapplicable criteria return null), reconciled per criterion by Experiment 1’s conservative adjudicator: '
      'agreement within one point is averaged; a gap of two or more takes the lower score and is flagged. Both '
      'scorers are gemini-3.1-flash-lite at temperature 0.15.')
    P('Pre-registered audits (Experiment 1’s stack). Evidence sufficiency: 100% of scored cells carry a '
      'verbatim evidence quote. Prompt robustness: the two paraphrased scorers agree closely (A–B Pearson '
      f'r = {RUBRIC["rab_lo"]:.2f}–{RUBRIC["rab_hi"]:.2f}; a gap of two or more points on a single cell). '
      'Length bias: a per-criterion regression of score on conversation word count leaves length explaining '
      f'under 30% of the variance on every criterion except {_flag} (just above the bound), so no headline '
      'rests on length. Positive controls: all ten design-implied directions hold (treatment higher on the '
      'expansion and evaluation criteria, lower on premature-convergence risk). Divergences from Experiment 1, '
      'by design: Experiment 1 ran scorers A and B on a local Qwen model with Gemini as a cross-MODEL Scorer C '
      'and applied FDR correction; here both scorers run on Gemini (so the cross-model leg is not reproduced), '
      'scoring is at the conversation level rather than per segmented episode, and no multiplicity correction '
      'is applied (the Experiment 3 analysis policy). Outputs: outputs/regulated_rubric_raw.csv, '
      'regulated_rubric_adjudicated.csv, regulated_rubric_audit.csv, regulated_rubric_contrast.csv.')

if URUBRIC:
    P('User-behaviour rubric (Experiment 1 “Option B”; §3.6, Table 4). The same masked transcripts are scored '
      'on six 0–4 user-only criteria (initiative, question richness, proposal specificity, yes-and uptake, '
      'reframing, and engagement depth), with the model instructed to rate only the participant’s turns, to '
      'quote user text only for every non-null score, and to treat echoes of the assistant as non-originated. '
      'Experiment 1 ran this as a single Gemini scorer; here it uses the same dual paraphrased scorers (A/B) '
      'and conservative adjudicator as the dialogic rubric, giving an added prompt-robustness check. The audit '
      f'is strong: 100% evidence-grounded, no high-disagreement cells, A–B agreement r={URUBRIC["rab_lo"]:.2f}–'
      f'{URUBRIC["rab_hi"]:.2f}, and no length confound (R²≤0.11). Outputs: outputs/user_rubric_raw.csv, '
      'user_rubric_adjudicated.csv, user_rubric_audit.csv, user_rubric_contrast.csv.')

# ---------------------------------------------------------------- References
def REF(t):
    p = doc.add_paragraph(t)
    pf = p.paragraph_format
    pf.left_indent = Inches(0.5); pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(4)
    return p

H1('References')
for r in [
 'Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. Qualitative Research in '
 'Psychology, 3(2), 77–101. https://doi.org/10.1191/1478088706qp063oa',
 'Dorst, K., & Cross, N. (2001). Creativity in the design process: Co-evolution of problem–solution. '
 'Design Studies, 22(5), 425–437. https://doi.org/10.1016/S0142-694X(01)00009-6',
 'Fusaroli, R., & Tylén, K. (2016). Investigating conversational dynamics: Interactive alignment, '
 'interpersonal synergy, and collective task performance. Cognitive Science, 40(1), 145–171. '
 'https://doi.org/10.1111/cogs.12251',
 'Gray, K., Anderson, S., Doyle, C. M., Hester, N., Schmitt, P., Vonasch, A. J., Allison, S. T., & '
 'Jackson, J. C. (2019). "Forward flow": A new measure of associative and creative thought. American '
 'Psychologist, 74(5), 539–554. https://doi.org/10.1037/amp0000391',
 'Hollan, J., Hutchins, E., & Kirsh, D. (2000). Distributed cognition: Toward a new foundation for '
 'human–computer interaction research. ACM Transactions on Computer–Human Interaction, 7(2), 174–196. '
 'https://doi.org/10.1145/353485.353487',
 'Jansson, D. G., & Smith, S. M. (1991). Design fixation. Design Studies, 12(1), 3–11. '
 'https://doi.org/10.1016/0142-694X(91)90003-F',
 'Järvelä, S., & Hadwin, A. F. (2013). New frontiers: Regulating learning in CSCL. Educational '
 'Psychologist, 48(1), 25–39. https://doi.org/10.1080/00461520.2012.748006',
 'Johnson, D. R., Kaufman, J. C., Baker, B. S., Patterson, J. D., Barbot, B., Green, A. E., van Hell, J., '
 'Kennedy, E., Sullivan, G. F., Taylor, C. L., Ward, T., & Beaty, R. E. (2022). Divergent semantic '
 'integration (DSI): Extracting creativity from narratives with distributional semantic modeling. '
 'Behavior Research Methods. https://doi.org/10.3758/s13428-022-01986-2',
 'Orne, M. T. (1962). On the social psychology of the psychological experiment: With particular reference '
 'to demand characteristics and their implications. American Psychologist, 17(11), 776–783. '
 'https://doi.org/10.1037/h0043424',
 'Reitter, D., & Moore, J. D. (2014). Alignment and task success in spoken dialogue. Journal of Memory '
 'and Language, 76, 29–46. https://doi.org/10.1016/j.jml.2014.05.008',
 'Sawyer, R. K. (2021). The iterative and improvisational nature of the creative process. Journal of '
 'Creativity, 31, 100002. https://doi.org/10.1016/j.yjoc.2021.100002',
 'Sowden, P. T., Pringle, A., & Gabora, L. (2015). The shifting sands of creative thinking: Connections '
 'to dual-process theory. Thinking & Reasoning, 21(1), 40–60. https://doi.org/10.1080/13546783.2014.885464',
 'Tseng, Y.-M., Huang, Y.-C., Hsiao, T.-Y., Chen, W.-L., Huang, C.-W., Meng, Y., & Chen, Y.-N. (2024). '
 'Two tales of persona in LLMs: A survey of role-playing and personalization. In Findings of the '
 'Association for Computational Linguistics: EMNLP 2024. Association for Computational Linguistics.',
 'White, P. R. R. (2003). Beyond modality and hedging: A dialogic view of the language of intersubjective '
 'stance. Text, 23(2), 259–284.',
]:
    REF(r)

try:
    doc.save(OUT)
    _saved = OUT
except PermissionError:
    _saved = OUT.replace('.docx', '_NEW.docx')
    doc.save(_saved)
    print('!! PRIMARY LOCKED (close it in Word). Saved fallback ->', _saved)
body = 0
for p in doc.paragraphs:
    if p.text.strip().startswith('Appendix A'): break
    body += len(p.text.split())
print('Saved:', _saved)
print('Total words:', sum(len(p.text.split()) for p in doc.paragraphs), '| Body (excl. appendix):', body,
      '| Tables:', len(doc.tables), '| Figures:', len(doc.inline_shapes))
