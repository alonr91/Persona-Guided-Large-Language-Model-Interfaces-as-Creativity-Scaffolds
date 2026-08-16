"""Generate the Gemini-Scorer-C version of the persona-family rubric radar
(matches the styling of the Qwen Scorer A radar) and embed it into both:

  - regulated_llm_reanalysis/figures/fig_persona_family_rubric_profiles_gemini.png
  - Experiment1_Results.docx                  (after the existing Figure R9)
  - Experiment1_StanceContracts_Paper_v4.docx (after the existing Figure 9 in 3.8)

Forbidden-language audit runs after each docx save.
"""
from __future__ import annotations
import os, sys
from pathlib import Path
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

from docx import Document
from docx.shared import Inches

ROOT = Path(r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1')
OUT = ROOT / 'regulated_llm_reanalysis'
FIG_DIR = OUT / 'figures'
FIG_DIR.mkdir(parents=True, exist_ok=True)
FIG_PATH = FIG_DIR / 'fig_persona_family_rubric_profiles_gemini.png'

CRITERIA_ORDER = (
    'exploration_opening', 'reframing_quality', 'evaluative_discipline',
    'agency_preservation', 'anchor_management', 'coregulation_uptake',
    'timing_fit', 'implementation_grounding', 'cognitive_load_clarity',
    'stance_integrity', 'premature_convergence_risk', 'runaway_divergence_risk',
)
FAMILY_ORDER = ('Divergent', 'Convergent', 'Rational', 'BoundedRational')
COLORS = {
    'GPT': '#6c6c6c',
    'Divergent': '#2a9d8f', 'Convergent': '#e76f51',
    'Rational': '#6f62b6', 'BoundedRational': '#e9c46a',
}

CAPTION = (
    "Figure 9b. Rubric profile per persona family (Gemini Scorer C, 0–4). "
    "Same axes and ordering as Figure 9 but scored by Gemini 3.1 Flash Lite "
    "Preview on the full 200-episode stratified sample (50 per family for "
    "Divergent/Convergent/GPT; 25 each for Rational/BoundedRational). All four "
    "persona families occupy a regulation envelope above the GPT baseline, "
    "with the Divergent persona reaching the highest reframing_quality and "
    "exploration_opening, consistent with the Qwen Scorer A pattern."
)


# -------------------- figure --------------------

def build_figure():
    plt.rcParams.update({'figure.dpi': 120, 'savefig.dpi': 160, 'font.size': 10})
    df = pd.read_csv(OUT / '04_episode_rubric_scores_raw_scorerC.csv')
    df = df[df.scorer == 'C'].copy()
    df['score_0_4'] = pd.to_numeric(df['score_0_4'], errors='coerce')
    df = df.dropna(subset=['score_0_4'])

    fams = ['GPT'] + list(FAMILY_ORDER)
    means = {fam: {} for fam in fams}
    for fam in fams:
        for crit in CRITERIA_ORDER:
            sub = df[(df.criterion == crit)
                     & (df.persona_family_original_hidden == fam)]
            means[fam][crit] = float(sub['score_0_4'].mean()) if len(sub) else np.nan

    angles = np.linspace(0, 2 * np.pi, len(CRITERIA_ORDER), endpoint=False).tolist()
    angles += angles[:1]

    fig = plt.figure(figsize=(9, 8))
    ax = fig.add_subplot(111, polar=True)
    for fam in fams:
        vals = [means[fam][c] for c in CRITERIA_ORDER]
        vals += vals[:1]
        ax.plot(angles, vals, marker='o', label=fam,
                color=COLORS.get(fam, 'black'), linewidth=1.5)
        ax.fill(angles, vals, alpha=0.08, color=COLORS.get(fam, 'black'))
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([c.replace('_', '\n') for c in CRITERIA_ORDER], fontsize=8)
    ax.set_ylim(0, 4)
    ax.set_title('Figure 9b — rubric profiles by persona family '
                 '(Gemini Scorer C, 0–4)', y=1.08)
    ax.legend(loc='lower right', bbox_to_anchor=(1.2, -0.05), fontsize=8)
    plt.tight_layout()
    plt.savefig(FIG_PATH, bbox_inches='tight')
    plt.close()
    print(f'wrote {FIG_PATH}')


# -------------------- helpers --------------------

def insert_paragraph_before(target_para, text: str, style: str):
    new_p = target_para._parent.add_paragraph(text, style=style)
    target_para._element.addprevious(new_p._element)
    return new_p


def insert_paragraph_after(target_para, text: str, style: str):
    new_p = target_para._parent.add_paragraph(text, style=style)
    target_para._element.addnext(new_p._element)
    return new_p


def insert_image_after(target_para, image_path: str, width_in: float = 5.5):
    """Insert a paragraph holding an inline image immediately after target_para."""
    parent = target_para._parent
    new_p = parent.add_paragraph()
    target_para._element.addnext(new_p._element)
    run = new_p.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    return new_p


def forbidden_audit(path):
    forbid = ['true creativity', 'ground-truth creativity',
              'validated creativity improvement', 'objectively more creative',
              'made participants more creative']
    d = Document(path)
    full = '\n'.join(p.text for p in d.paragraphs)
    hits = [t for t in forbid if t in full.lower()]
    print(f'  forbidden-language hits: {hits or "none"}')


# -------------------- Results.docx --------------------

def update_results_doc():
    path = ROOT / 'Experiment1_Results.docx'
    d = Document(str(path))
    paras = list(d.paragraphs)

    # Find the existing R9 figure caption (it begins with "Figure R9.")
    r9 = None
    for p in paras:
        if p.text.strip().startswith('Figure R9.') and 'profile' in p.text.lower():
            r9 = p
            break
    if r9 is None:
        # Fall back: insert at end of Cross-Model section
        for p in paras:
            if p.style.name.startswith('Heading 1') and 'Cross-Model Validation' in p.text:
                r9 = p
                break
    if r9 is None:
        raise RuntimeError('No anchor found in Results.docx')

    img_para = insert_image_after(r9, str(FIG_PATH), width_in=5.5)
    insert_paragraph_after(img_para, CAPTION, 'Normal')
    d.save(str(path))
    print(f'updated {path}')
    forbidden_audit(str(path))


# -------------------- Paper v4 --------------------

def update_paper_doc():
    path = ROOT / 'Experiment1_StanceContracts_Paper_v4.docx'
    d = Document(str(path))
    paras = list(d.paragraphs)

    # Find the existing Figure 9 caption inside section 3.8 (it begins
    # "Figure 9. Rubric profile per persona family (radar).").
    in_38 = False
    fig9 = None
    for p in paras:
        t = p.text.strip()
        if p.style.name.startswith('Heading 2') and t.startswith('3.8 '):
            in_38 = True; continue
        if in_38 and p.style.name.startswith('Heading 2') and t.startswith('3.9 '):
            break
        if in_38 and t.startswith('Figure 9.') and 'persona family' in t.lower():
            fig9 = p
            break
    if fig9 is None:
        raise RuntimeError('Figure 9 caption not found in section 3.8')

    img_para = insert_image_after(fig9, str(FIG_PATH), width_in=5.5)
    insert_paragraph_after(img_para, CAPTION, 'Normal')
    d.save(str(path))
    print(f'updated {path}')
    forbidden_audit(str(path))


# -------------------- main --------------------

if __name__ == '__main__':
    build_figure()
    try:
        update_results_doc()
    except Exception as e:
        print(f'  Results.docx update failed: {e}')
    try:
        update_paper_doc()
    except Exception as e:
        print(f'  Paper.docx update failed: {e}')
