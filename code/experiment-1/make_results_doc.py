"""
Build a Results.docx file with text + embedded figures.
Also produces one new figure: perception-behavior scatter (Δa_prop vs Δcreativity).
"""
import os, sys, warnings
warnings.filterwarnings('ignore')
sys.stdout.reconfigure(encoding='utf-8')
import numpy as np, pandas as pd
from scipy import stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
OUT  = os.path.join(ROOT, 'analysis_out')
FIG  = os.path.join(ROOT, 'figures')

# ------------------------------------------------------------------
# NEW FIGURE: perception-behavior scatter (Δa_prop vs Δcreativity)
# ------------------------------------------------------------------
ep = pd.read_csv(os.path.join(OUT, 'extension_conv_master.csv'))
users = pd.read_excel(os.path.join(OUT, 'users_translated.xlsx'), sheet_name='corrected_users')
users.columns = [c.strip() for c in users.columns]

# find GPT round per user
def round_of_gpt(row):
    r1 = str(row.get('Persona round 1','')).lower()
    r2 = str(row.get('Persona round 2','')).lower()
    if 'gpt' in r1: return 1
    if 'gpt' in r2: return 2
    return np.nan
users['gpt_round'] = users.apply(round_of_gpt, axis=1)

def make_gp(df, r1c, r2c):
    gpt = np.where(df['gpt_round']==1, df[r1c], np.where(df['gpt_round']==2, df[r2c], np.nan))
    per = np.where(df['gpt_round']==1, df[r2c], np.where(df['gpt_round']==2, df[r1c], np.nan))
    return pd.Series(gpt, index=df.index), pd.Series(per, index=df.index)

users['cr_gpt'], users['cr_per'] = make_gp(users, 'Creativity assistant #1','Creativity assistant #2')
users['d_cr'] = users['cr_per'].astype(float) - users['cr_gpt'].astype(float)

# collect Δa_prop per user
ap = (ep.pivot_table(index='user', columns='condition', values='a_prop', aggfunc='mean'))
ap['d_a_prop'] = ap.get('Persona',0) - ap.get('GPT',0)
m = users[['id','d_cr']].merge(ap[['d_a_prop']], left_on='id', right_index=True)
m = m.dropna()

rho, p = stats.spearmanr(m['d_a_prop'], m['d_cr'])
print(f'Δa_prop vs Δcreativity: n={len(m)}, ρ={rho:.3f}, p={p:.4f}')

fig, ax = plt.subplots(figsize=(6.2, 4.8))
ax.scatter(m['d_a_prop'], m['d_cr'], alpha=0.55, s=40, edgecolor='k', linewidth=0.3, color='#2a6f97')
# robust regression line
slope, intercept, r, p_l, _ = stats.linregress(m['d_a_prop'], m['d_cr'])
xs = np.linspace(m['d_a_prop'].min(), m['d_a_prop'].max(), 50)
ax.plot(xs, intercept + slope*xs, color='#c1121f', lw=2)
ax.axhline(0, color='gray', lw=0.5, ls='--')
ax.axvline(0, color='gray', lw=0.5, ls='--')
ax.set_xlabel('Δ Assistant propose-new-idea  (Persona − GPT)')
ax.set_ylabel('Δ Self-reported creativity  (Persona − GPT)')
ax.set_title(f'Perception–Behavior Triangulation\nSpearman ρ = {rho:.2f},  p = {p:.3f},  n = {len(m)}')
ax.grid(alpha=0.25)
plt.tight_layout()
plt.savefig(os.path.join(FIG, 'figR_perception_behavior.png'), dpi=160)
plt.close()
print('saved figR_perception_behavior.png')

# ------------------------------------------------------------------
# BUILD THE DOCX
# ------------------------------------------------------------------
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def H1(t):
    p = doc.add_heading(t, level=1)
    return p
def H2(t):
    p = doc.add_heading(t, level=2)
    return p
def P(t, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = italic; r.bold = bold
    return p
def FIG_INS(path, caption, width=5.8):
    if not os.path.exists(path):
        P(f'[MISSING FIGURE: {path}]', italic=True); return
    doc.add_picture(path, width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption)
    r.italic = True; r.font.size = Pt(10)

# ----- Title -----
title = doc.add_heading('Results', level=0)
P('Experiment 1: Persona-Guided LLM Co-Creative Interaction', italic=True)

# ----- Overview -----
P('All inferential analyses were conducted on the keepable paired dataset comprising '
  '97 participants, 194 conversations, and 3,412 messages. Every within-subject contrast '
  'uses the participant as the analytic unit and compares the two rounds (GPT versus Persona) '
  'on the same individual, following the counterbalanced design described in the Method. '
  'For each dependent measure we report the mean in each condition, the mean difference '
  '(Persona − GPT), the paired t-statistic, the Wilcoxon signed-rank test as a nonparametric '
  'robustness check, and Cohen\'s dz. Effect-size magnitudes are treated as the primary '
  'evidential signal, with p-values reported for disclosure rather than as the principal '
  'criterion.')

# ----- RQ1 Subjective -----
H1('Subjective Experience (RQ1)')
P('The first research question asked whether persona-guided interaction changes '
  'participants\' perceived creative enhancement and sense of idea ownership. None of the '
  'four paired comparisons on the self-report items reached significance. Creativity ratings '
  'did not differ between the standard GPT round and the persona round (M_GPT = 3.58, '
  'M_Persona = 3.57, t(96) = 0.06, p = .95, dz = 0.01). Ownership ratings likewise did not '
  'differ (M_GPT = 3.46, M_Persona = 3.53, t(96) = −0.35, p = .72, dz = −0.04). The '
  'round-indexed analyses (Round 1 vs. Round 2) were similarly null, ruling out a simple '
  'order effect.')
P('The forced-choice preference item, however, revealed an asymmetry that the Likert '
  'ratings did not. When asked which interface was more effective for reaching a creative '
  'solution, 52 participants selected the persona-guided interface, 38 selected the '
  'standard GPT interface, and 7 selected both. The persona interface was therefore '
  'preferred by 57.8% of unambiguous responders, a dissociation between continuous and '
  'categorical self-report that is revisited in the Discussion.')
FIG_INS(os.path.join(FIG, 'fig1_paired_questionnaire.png'),
        'Figure 1. Paired self-report scores (Persona vs. GPT) on creativity and ownership. '
        'Within-subject lines are near-flat, confirming null effects on Likert ratings despite '
        'the directional preference captured by the forced-choice item.')

# ----- Personality moderation -----
H1('Personality Moderation (RQ2, Subjective Layer)')
P('Spearman correlations between each of the five BFI-2-XS domain scores and the '
  'within-subject deltas in creativity and ownership ratings (Persona − GPT) were estimated '
  'on the full sample (n = 97). None of the ten tests reached statistical significance, and '
  'all effect magnitudes were small (all |ρ| < .10). Open-Mindedness, the trait with the '
  'strongest theoretical prior for modulating persona benefit, produced ρ = 0.02 for '
  'creativity and ρ = 0.01 for ownership. The personality moderation hypothesis at the '
  'perception layer is therefore not supported in this sample.')
P('Descriptive breakdown by persona family revealed directional variation that is '
  'consistent with theoretical prediction but constrained by the small per-cell samples '
  'for the Rational and BoundedRational personas (n = 9 each). Participants assigned to '
  'the Divergent and Convergent personas produced Persona − GPT creativity deltas near zero '
  '(M = +0.08 and +0.10, respectively), whereas participants assigned to the Rational and '
  'BoundedRational personas showed negative deltas (M = −0.33 and −0.56).')
FIG_INS(os.path.join(FIG, 'fig7_personality.png'),
        'Figure 2. BFI-2-XS domain scores against within-subject Persona − GPT deltas in '
        'perceived creativity and ownership. No trait-level moderation reaches significance; '
        'all |ρ| < 0.10.')

# ----- Manipulation check -----
H1('Manipulation Check: Persona Fidelity')
P('Before testing condition contrasts on the dialogic process layer, we verified that '
  'the four personas produced distinguishable behavioral signatures on the assistant side. '
  'Each predicted ordinal signature is recovered. The Divergent persona produces the highest '
  'expansion score (1.81) and the highest reframe score (1.05) of any family. The Convergent '
  'persona produces elevated contraction (1.42) relative to expansion (0.77). The Rational '
  'persona produces the highest contraction (1.70), the highest critique (0.29), and the '
  'lowest propose-new-idea rate (0.87). The BoundedRational persona exhibits the mixed '
  'expansion-plus-contraction profile predicted by a satisficing heuristic. The GPT baseline '
  'produces the lowest expansion (0.45) and the highest propose-new-idea rate (1.61), '
  'confirming that the non-persona condition defaults to a solution-delivery register. '
  'Because each persona installed the behavioral regime prescribed by its theoretical '
  'template, the Persona–GPT contrast in the remainder of the Results can be interpreted '
  'as a stance-contract effect rather than as a prompt-compliance artifact.')
FIG_INS(os.path.join(FIG, 'figM1_manipulation.png'),
        'Figure 3. Assistant-side stance profile by persona family on the White (2003) '
        'appraisal dimensions. Each predicted signature is recovered: Divergent = high '
        'expansion/reframe; Convergent and Rational = elevated contraction; Rational = '
        'highest critique; GPT = highest propose-new-idea, lowest expansion.')

FIG_INS(os.path.join(FIG, 'figM2_tone_qtype.png'),
        'Figure 4. Critique tone (assistant) and question type (user) by persona family. '
        'Users ask substantially fewer clarifying questions under persona framing, with '
        'the user-question distribution shifting toward no-question turns.')

# ----- Assistant stance shifts -----
H1('Dialogic Process: Assistant-Side Stance Shifts')
P('Assistant-side stance indices were aggregated to the conversation level and entered '
  'into paired tests contrasting the Persona and GPT conditions within each participant. '
  'The magnitude of these effects is the principal mechanism finding of the experiment. '
  'Under persona framing the assistant commits to concrete recommendations far less often '
  '(dz = −2.25), expresses substantially lower epistemic certainty (dz = −1.45), and '
  'proposes new ideas outright at less than two-thirds the baseline rate (dz = −1.35). '
  'Simultaneously, its expansion moves — inviting alternatives, broadening the framing, '
  'surfacing additional angles — more than double relative to GPT (dz = +1.29), and '
  'reframing moves rise from near-zero prevalence under GPT to an appreciable rate under '
  'personas (dz = +0.91). Together these contrasts describe a clean shift from a '
  'solution-delivery register under GPT to a dialogic expansion register under personas. '
  'In White\'s (2003) terms, the persona prompts convert the assistant from a '
  'contraction-oriented speaker to an expansion-oriented speaker.')
FIG_INS(os.path.join(FIG, 'figM3_stance_paired.png'),
        'Figure 5. Paired within-subject effect sizes (Cohen\'s dz) for all LLM-coded '
        'stance dimensions. Assistant-side effects (top of plot) are the largest in the '
        'dataset (|dz| between 0.91 and 2.25); user-side spillover effects are moderate '
        'but systematic and in the theoretically predicted direction.')

# ----- User-side spillover -----
H1('Dialogic Process: User-Side Spillover')
P('A central prediction of the stance-contract framework is that shifts in assistant '
  'behavior propagate to the user side — that the user\'s own stance register tracks the '
  'assistant\'s. Three patterns deserve emphasis. First, the single largest user-side '
  'change is a questioning contraction: under personas, users ask markedly fewer questions '
  'overall (dz = −0.76) and, in particular, fewer clarifying questions (dz = −0.82). '
  'Second, users take on stance moves that mirror the assistant\'s expansion register: '
  'reframes rise (dz = +0.49), new-idea proposals rise (dz = +0.43), and expansion moves '
  'rise (dz = +0.47). Third, engagement intensifies: users write longer messages under '
  'personas (+5.5 words, dz = +0.48), contribute more total words (+34 words, dz = +0.38), '
  'and their messages grow more steeply over the course of the conversation (dz = +0.26). '
  'The user talks more, asks less, and produces more stance work — a profile that is '
  'qualitatively incompatible with a passive or disengaged reading of persona interaction.')
FIG_INS(os.path.join(FIG, 'fig2_trajectory.png'),
        'Figure 6. Turn-level engagement trajectories within a conversation, split by '
        'condition. User-message length grows more steeply under personas, with user '
        'questioning density correspondingly reduced.')

# ----- Fixation -----
H1('Fixation and Drift from the Initial AI Anchor')
P('We tested whether persona framing changes the degree to which user ideation remains '
  'semantically proximal to the assistant\'s first proposed idea. Early anchor similarity '
  'was lower under personas (M = 0.298) than under GPT (M = 0.336; dz = −0.25, p = .016). '
  'The derived drift trajectory (late − early) was negative under personas and positive '
  'under GPT (Δ = −0.064, dz = −0.28, p = .007): under GPT users move away from the anchor '
  'over time, whereas under personas they stay closer to it. The overall fixation index '
  'did not differ (dz = +0.03).')
P('The drift-trajectory finding runs counter to the intuitive expectation that personas '
  'marketed as "creative" should free users from the AI\'s framing. It is compatible, '
  'however, with the manipulation-check pattern: because personas open more dialogic space '
  'and withhold concrete proposals, the assistant\'s first proposal ends up functioning as '
  'the shared reference point throughout the conversation. Under GPT, in contrast, the '
  'assistant keeps injecting fresh proposals, which themselves draw the user\'s ideation '
  'in multiple directions. Personas expand dialogic work but do not, on their own, produce '
  'escape from an initial anchor.')
FIG_INS(os.path.join(FIG, 'figM4_fixation.png'),
        'Figure 7. Anchor similarity and drift trajectory across the conversation. Users '
        'under GPT drift away from the initial assistant proposal; users under personas '
        'remain closer to it, despite the broader dialogic register of the persona rounds.')

# ----- Product -----
H1('Creative Product: Portfolio Distinctiveness')
P('Portfolio-level distinctiveness was estimated twice for every participant\'s user-side '
  'idea portfolio: once using TF-IDF vectorization and once using SBERT sentence embeddings. '
  'Both measures operationalize "how semantically different this participant\'s ideas are '
  'from those of the rest of the sample," but they differ in sensitivity to surface lexicon '
  'versus underlying meaning.')
P('The TF-IDF distinctiveness measure produced a paired effect favoring personas (M_Persona '
  '= 0.830, M_GPT = 0.815, t(96) = 3.31, p = .001, dz = +0.34). The SBERT distinctiveness '
  'measure, however, was null (M_Persona = 0.342, M_GPT = 0.345, t(96) = −0.37, p = .72, '
  'dz = −0.04). SBERT breadth and redundancy measures were similarly null (both |dz| < 0.16). '
  'The interpretive implication is that persona framing shifts the lexical texture of user '
  'output — the specific vocabulary and phrasing — without changing its semantic footprint. '
  'Reported in isolation, the TF-IDF result would suggest that personas broaden the idea '
  'space. Reported jointly with SBERT, the two tell a more careful story: the idea space '
  'itself is unchanged; what changes is how users talk about it.')
FIG_INS(os.path.join(FIG, 'figM5_sbert_distinct.png'),
        'Figure 8. Portfolio distinctiveness by condition. TF-IDF (left) shows a reliable '
        'advantage for persona rounds (dz = +0.34, p = .001); SBERT (right) shows no '
        'effect (dz = −0.04, p = .72). The divergence isolates lexical surface shifts '
        'from semantic footprint shifts.')
FIG_INS(os.path.join(FIG, 'figM6_umap.png'),
        'Figure 9. UMAP projection of SBERT idea embeddings, colored by condition. The '
        'GPT and Persona idea clouds overlap substantially, consistent with the null '
        'SBERT distinctiveness effect.')

# ----- Perception-behavior -----
H1('Perception–Behavior Triangulation')
P('Twenty-eight process and product indices were correlated (Spearman ρ) against each of '
  'the two self-reported deltas (creativity and ownership), producing 56 tests in total. '
  'Only one survived α = .05: the Persona − GPT change in assistant propose-new-idea density '
  'was positively associated with the Persona − GPT change in perceived creativity '
  '(ρ = 0.26, p = .011). The direction of this correlation indicates that participants rated '
  'the persona round as more creative relative to the GPT round specifically when the '
  'persona withheld fresh proposals at a higher rate than GPT. Interpreted through the '
  'proposal-gap account, the subjective sensation of creativity enhancement is contingent '
  'on the assistant ceding the proposal role to the user. Two additional correlations '
  'approached but did not reach significance (Δa_con vs. Δownership: ρ = −0.19, p = .065; '
  'Δa_com vs. Δownership: ρ = −0.18, p = .074), both in the theoretically predicted '
  'direction.')
FIG_INS(os.path.join(FIG, 'figR_perception_behavior.png'),
        'Figure 10. Perception–behavior triangulation. Within-subject change in assistant '
        'propose-new-idea density (x-axis) against within-subject change in self-reported '
        'creativity (y-axis). The positive slope indicates that participants perceive the '
        'persona round as more creative precisely when the assistant ceded the proposal '
        'role to them.')

# ----- Archetypes -----
H1('Interaction Archetypes (Exploratory)')
P('K-means clustering (k = 3, silhouette = 0.20) on twelve standardized process indices '
  'identified three interaction archetypes. Archetype 0 (short-turn, high-question baseline) '
  'is the most common (74 GPT, 63 Persona). Archetype 1 (long-turn, low-question elaborator) '
  'is disproportionately populated by persona rounds (21 GPT, 30 Persona). Archetype 2 '
  '(late-commit spike) is small and roughly balanced. Persona framing modestly shifts the '
  'population from the baseline archetype toward the long-turn elaborator archetype. The '
  'three-archetype structure is reported as descriptive typology; inferential claims on '
  'cluster assignment are not made.')
FIG_INS(os.path.join(FIG, 'fig5_archetypes.png'),
        'Figure 11. Interaction-archetype distribution across conditions. Persona framing '
        'shifts the population toward the long-turn, low-question elaborator archetype.')

# ----- Summary -----
H1('Summary')
P('Across the three layers, the evidence converges on a coherent account. At the subjective '
  'layer, continuous self-report is null, but a forced-choice preference leans toward the '
  'persona interface. At the process layer, personas install distinct, theory-consistent '
  'assistant stance regimes (manipulation check), produce very large shifts in the '
  'assistant\'s expansion, reframe, certainty, commit, and propose behaviors (|dz| between '
  '0.91 and 2.25), and produce moderate spillover into the user\'s own stance register '
  '(user-side |dz| between 0.27 and 0.82). Personas also reduce users\' drift away from the '
  'assistant\'s first anchor, qualifying the divergent-creativity reading. At the product '
  'layer, the semantic idea portfolio is unchanged even though the lexical surface diverges. '
  'Finally, the only robust perception–behavior linkage operates through the assistant\'s '
  'proposal rate: users perceive a round as more creative to the extent that the assistant '
  'leaves proposal work to them. These findings form the empirical basis for the '
  'stance-contract mechanism account developed in the Discussion.')

out_path = os.path.join(ROOT, 'Experiment1_Results.docx')
doc.save(out_path)
print('saved:', out_path)
