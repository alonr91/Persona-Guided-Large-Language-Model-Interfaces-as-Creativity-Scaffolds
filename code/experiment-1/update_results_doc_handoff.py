"""
Final handoff update for Experiment1_Results.docx.

Replaces the original Summary section with an expanded synthesis that
references every analysis layer added during this phase (LAYERs E, F, H,
I, J, the agentic idea-extraction pipeline / LAYER L, and the regulated
LLM reanalysis). Adds a closing methodology overview paragraph that
points to the methods appendix and the regulated_llm_reanalysis/ folder.

Forbidden-language audit verified after save.
"""
from __future__ import annotations
import os, sys, warnings
warnings.filterwarnings('ignore')
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

from docx import Document

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
doc_path = os.path.join(ROOT, 'Experiment1_Results.docx')


# -- new closing summary --
NEW_SUMMARY_PARAS = [
    ('Heading 1', 'Summary'),
    ('Normal',
     "This Experiment-1 analysis was conducted across multiple layers: a "
     "subjective layer (questionnaire deltas and personality moderation), a "
     "process layer (assistant- and user-side stance shifts, drift / anchor "
     "behaviour, interaction archetypes, consecutive-message semantic novelty, "
     "message-level surprise under GPT-2, a Yes-And transition decomposition, "
     "and a Divergent-vs-Convergent persona breakdown), a product layer "
     "(TF-IDF and SBERT portfolio distinctiveness, plus an open-source agentic "
     "idea-extraction pipeline that produced participant-centroid originality "
     "scores in the manner of Experiment 2), and a regulated LLM reanalysis "
     "that scored conversational episodes on twelve 0-4 ordinal regulation "
     "criteria with two paraphrased-prompt scorers and a rule-based adjudicator. "
     "All LLM-derived numbers are labelled as proxy scores, since no human "
     "creativity judges were available."),
    ('Normal',
     "Three integrated findings emerge."),
    ('Normal',
     "(1) Persona prompts install distinct, theory-consistent assistant stance "
     "regimes that produce moderate-to-large user-side spillover (manipulation "
     "check; |dz| up to 2.25 on assistant side, |dz| 0.27-0.82 on user side). "
     "Persona conversations show greater consecutive-message coupling (LAYER E "
     "novelty, dz = -0.29 in the within-subject paired test) and greater "
     "external unpredictability under GPT-2 (LAYER F surprise, dz = +0.56). "
     "The two-axis signature \u2014 internally tighter, externally more "
     "idiosyncratic \u2014 is consistent with persona prompts holding a "
     "narrower semantic focus while using more stylistically unusual language. "
     "The Yes-And transition decomposition shows that the user, not the "
     "assistant, accepts the partner's semantic footing more strongly under "
     "personas (assistant\u2192user d\u0304 paired \u0394 = -0.048, "
     "p < 0.001, dz = -0.39)."),
    ('Normal',
     "(2) The Divergent-vs-Convergent breakdown shows that the Divergent "
     "persona is the stronger and more trait-amplified intervention. Within "
     "the Persona arm, users of the Divergent persona ask markedly more "
     "questions (q-rate 0.34 vs 0.20, p < 0.001) and the Divergent persona "
     "produces a larger shift in assistant propose-rate than the Convergent "
     "persona (\u0394 of \u0394 = +0.135, d = +0.56, p = 0.015). Open-Mindedness "
     "moderates several behavioural-gain effects within Divergent users only; "
     "Convergent users\u2019 behavioural gains are essentially trait-uniform. "
     "Ownership shifts in opposite directions across the two thinking-mode "
     "personas (Divergent +0.32; Convergent -0.27)."),
    ('Normal',
     "(3) The product-layer evidence is mixed, and the regulated reanalysis "
     "explains why. On the open-source agentic idea-extraction pipeline "
     "(Qwen3-4B-Instruct INT4 on Arc 140T, Filters 1-3 + dual scorers), "
     "persona rounds produce significantly more discrete user-originated ideas "
     "per round (paired \u0394 = +1.18, dz = +0.53) but those portfolios are "
     "*less* distinctive from same-condition peers than GPT portfolios "
     "(orig_same paired \u0394 = -0.025, p = 2\u00D710\u207B\u2078, dz = -0.67). "
     "This direction is opposite to Experiment 2 \u2014 plausibly because "
     "Experiment 1 assigned a single persona per round while Experiment 2 "
     "allowed elective use of two simultaneously available personas. At the "
     "regulated-rubric level, episode-by-episode proxy ratings show that "
     "persona conversations score much higher on reframing_quality "
     "(g = +1.70), exploration_opening (g = +1.58), anchor_management "
     "(g = +1.51), timing_fit (g = +1.26), and co-regulation (g = +1.05), "
     "and lower on premature_convergence_risk (g = -0.62), with 9 of 12 "
     "criteria FDR-significant. Positive controls (Divergent > Rational on "
     "exploration_opening; Rational > Divergent on evaluative_discipline) and "
     "length-bias checks both pass. Together these point to persona prompts "
     "functioning as **stance contracts** \u2014 they reorganise *how* "
     "creative work is regulated in dialogue (process), produce more fluent "
     "but less between-subject-distinctive idea portfolios (product), and do "
     "not move self-reported creativity ratings (subjective). The strongest "
     "claim this study supports, given the absence of human raters, is that "
     "persona-guided LLM interaction *changes the regulation of creative "
     "dialogue*. Establishing an externally judged creativity gain would "
     "require human or domain-expert rating that this study does not include."),
    ('Heading 2', 'Methodology, Reproducibility, and Artifacts'),
    ('Normal',
     "All inferential analyses use the participant as the unit of analysis "
     "where appropriate; deltas are signed Persona \u2212 GPT throughout. The "
     "agentic idea-extraction pipeline runs end-to-end on open-source "
     "components (Qwen3-4B-Instruct via OpenVINO INT4 on an Intel Arc 140T "
     "GPU, BAAI/bge-large-en-v1.5 sentence embeddings, sklearn HDBSCAN). The "
     "regulated-reanalysis layer uses the same model under two paraphrased "
     "prompts (Scorer A and Scorer B), schema-constrained JSON via "
     "lm-format-enforcer, and a rule-based adjudicator. Every LLM-derived "
     "number is a proxy score; bias and length-leakage audits are reported in "
     "08_validation_and_bias_audit.csv."),
    ('Normal',
     "All artifacts of the regulated reanalysis are written to "
     "`regulated_llm_reanalysis/` (00_data_audit.md through 12_methods_appendix.md "
     "plus seven figures). All artifacts of the agentic idea-extraction pipeline "
     "are written to `analysis_out/production/`. The full method appendix is "
     "`regulated_llm_reanalysis/12_methods_appendix.md`. To re-run the regulated "
     "reanalysis: `python -m os_pipeline.regulated.regulated_run`. To re-run the "
     "extraction pipeline: `python -m os_pipeline.production_run --all`."),
]


def replace_summary():
    d = Document(doc_path)
    paras = list(d.paragraphs)
    summary_idx = None
    for i, p in enumerate(paras):
        if p.style.name.startswith('Heading 1') and p.text.strip().lower() == 'summary':
            summary_idx = i; break
    if summary_idx is None:
        print('[handoff] Summary heading not found; aborting.')
        return
    # delete from Summary heading to end of doc (Summary is the last section)
    n_before = len(paras)
    for p in paras[summary_idx:]:
        el = p._element
        parent = el.getparent()
        if parent is not None: parent.remove(el)
    print(f'[handoff] removed {n_before - summary_idx} old summary paragraphs')
    # append the new ones
    for style, text in NEW_SUMMARY_PARAS:
        d.add_paragraph(text, style=style)
    d.save(doc_path)


def verify():
    forbid = ['true creativity', 'ground-truth creativity',
              'validated creativity improvement', 'objectively more creative',
              'made participants more creative']
    d = Document(doc_path)
    full = '\n'.join(p.text for p in d.paragraphs)
    hits = [t for t in forbid if t in full.lower()]
    print('forbidden-language hits:', hits or 'none')
    hs = [p.text.strip() for p in d.paragraphs
          if p.style.name.startswith('Heading 1') and p.text.strip()]
    print(f'{len(hs)} Heading-1 sections, {len(d.paragraphs)} paragraphs.')


if __name__ == '__main__':
    replace_summary()
    verify()
