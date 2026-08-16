"""
Export Experiment1_Results.docx → Experiment1_Results.md.

Walks paragraphs in order, maps Heading X → markdown #s, and EXTRACTS
embedded images from the docx (rather than guessing which file in figures/
matches each Figure label). Extracted images go to `md_figures/`.
"""
from __future__ import annotations
import os, re, sys
from docx import Document
from docx.oxml.ns import qn

try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
DOC = os.path.join(ROOT, 'Experiment1_Results.docx')
OUT = os.path.join(ROOT, 'Experiment1_Results.md')
MD_FIG_DIR = os.path.join(ROOT, 'md_figures')
os.makedirs(MD_FIG_DIR, exist_ok=True)

# Map figure label (case-insensitive prefix that appears in caption) → file path.
# The captions in the document take forms like "Figure 1.", "Figure R5.", etc.
FIG_MAP = {
    'figure 1':  'fig1_paired_questionnaire.png',
    'figure 2':  'fig2_trajectory.png',
    'figure 3':  'fig3_family_trajectory.png',
    'figure 4':  'fig4_ownership_gap.png',
    'figure 5':  'fig5_archetypes.png',
    'figure 6':  'fig6_distinctiveness.png',
    'figure 7':  'fig7_personality.png',
    # the make_results_doc.py outputs use figM* for the manipulation/extension layer
    'figure r1': 'figR_novelty_surprise_two_axis.png',
    'figure r2': 'figR_yes_and_transitions.png',
    'figure r3': 'figR_DvC_manipulation.png',
    'figure r4': 'figR_DvC_personality.png',
    'figure r5': 'figR_agentic_workflow.png',
    'figure r6': 'figR_originality_same_by_family.png',
    'figure r7': 'figR_fluency_by_family.png',
    'figure r8': 'figR_originality_big5.png',
    'figure r9': 'fig_persona_family_rubric_profiles.png',
    'figure r10': 'fig_rubric_condition_effects.png',
    'figure r11': 'fig_regulation_trajectory_by_condition.png',
    'figure r12': 'fig_validation_disagreement_heatmap.png',
}
# regulated-reanalysis figures live under regulated_llm_reanalysis/figures
REG_FIGS = {
    'fig_persona_family_rubric_profiles.png',
    'fig_rubric_condition_effects.png',
    'fig_regulation_trajectory_by_condition.png',
    'fig_validation_disagreement_heatmap.png',
}


def figure_path(fname: str) -> str:
    if fname in REG_FIGS:
        return f'regulated_llm_reanalysis/figures/{fname}'
    return f'figures/{fname}'


_caption_rx = re.compile(r'^figure\s+r?\d+', re.I)


def has_picture(p) -> bool:
    return bool(p._element.findall('.//' + qn('w:drawing')))


def extract_pictures_in_para(p, doc) -> list[str]:
    """Find embedded image relationship IDs in paragraph p; resolve to image
    parts and write each to md_figures/. Returns list of relative paths."""
    paths = []
    # blip elements carry the image relationship id
    a_blip = p._element.findall('.//' + qn('a:blip'))
    if not a_blip:
        # also check for {drawingml namespace}
        a_blip = p._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    for blip in a_blip:
        rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not rid:
            continue
        try:
            part = doc.part.related_parts[rid]
        except KeyError:
            continue
        # part.partname like '/word/media/image5.png'
        ext = os.path.splitext(part.partname)[1] or '.png'
        # use a deterministic name based on global counter so re-runs are stable
        n = len(os.listdir(MD_FIG_DIR)) + 1
        out_name = f'figure_{n:02d}{ext}'
        out_path = os.path.join(MD_FIG_DIR, out_name)
        # only write if not already exists (or content differs)
        with open(out_path, 'wb') as fh:
            fh.write(part.blob)
        paths.append(f'md_figures/{out_name}')
    return paths


def heading_level(style_name: str) -> int | None:
    if not style_name: return None
    s = style_name.lower()
    if s == 'title': return 1
    m = re.match(r'heading\s+(\d+)', s)
    return int(m.group(1)) if m else None


def main():
    # the original is sometimes locked by Word / OneDrive sync. Python open()
    # requests exclusive access; the Windows shell copy works around this.
    src = DOC
    try:
        d = Document(src)
    except Exception:
        # try multiple copy strategies — mingw-bash cp opens with FILE_SHARE_READ
        # which works around Word/OneDrive locks
        import tempfile, subprocess
        tmp = os.path.join(tempfile.gettempdir(), 'exp1_results_export.docx')
        copied = False
        for cmd in (
            ['bash', '-c', f'cp "{src}" "{tmp.replace(chr(92), "/")}"'],
            ['robocopy', os.path.dirname(src), tempfile.gettempdir(),
             os.path.basename(src), '/np', '/njh', '/njs'],
        ):
            try:
                r = subprocess.run(cmd, capture_output=True, timeout=30)
                if os.path.exists(tmp) and os.path.getsize(tmp) > 100_000:
                    copied = True
                    break
                # robocopy uses different output filename: src basename
                rb_path = os.path.join(tempfile.gettempdir(), os.path.basename(src))
                if os.path.exists(rb_path) and os.path.getsize(rb_path) > 100_000:
                    import shutil as _s
                    _s.move(rb_path, tmp)
                    copied = True
                    break
            except Exception:
                continue
        if not copied:
            raise RuntimeError(
                f'Cannot copy locked docx {src!r}. Close Word/OneDrive and retry.')
        print(f'[export] original locked; reading from {tmp}')
        d = Document(tmp)
    out: list[str] = []
    out.append('# Experiment 1: Persona-Guided LLM Co-Creative Interaction\n')
    out.append('')
    out.append('*This file is auto-converted from `Experiment1_Results.docx`. '
               'For figure files see `figures/` and `regulated_llm_reanalysis/figures/`.*')
    out.append('')

    # wipe and recreate md_figures so numbering is deterministic
    for f in os.listdir(MD_FIG_DIR):
        try: os.remove(os.path.join(MD_FIG_DIR, f))
        except: pass

    paras = list(d.paragraphs)
    pending_paths: list[str] = []
    for i, p in enumerate(paras):
        text = (p.text or '').strip()
        lvl = heading_level(p.style.name)

        # extract any embedded images in this paragraph
        if has_picture(p):
            pending_paths.extend(extract_pictures_in_para(p, d))
            # if the paragraph has both an image AND text, fall through to write text
            if not text:
                continue

        # if pending images and this paragraph is a caption — emit image + caption
        if pending_paths and _caption_rx.match(text):
            # extract figure label like "Figure R5" for the alt text
            m = re.match(r'^(figure\s+\S+?)\.', text, re.I)
            label = m.group(1) if m else 'figure'
            for fp in pending_paths:
                out.append('')
                out.append(f'![{label}]({fp})')
            out.append('')
            # caption text styled as bold-leading
            first_period = text.find('.')
            if first_period > 0:
                label = text[:first_period+1]
                rest = text[first_period+1:].strip()
                out.append(f'**{label}** {rest}')
            else:
                out.append(f'**{text}**')
            out.append('')
            pending_paths = []
            continue
        elif pending_paths:
            # image but next paragraph wasn't a caption — emit images alone
            for fp in pending_paths:
                out.append('')
                out.append(f'![figure]({fp})')
            out.append('')
            pending_paths = []

        if not text:
            out.append('')
            continue

        if lvl is not None:
            # we already emitted the top-level title in the file header; skip the
            # docx Title, and map Heading 1/2 to ## / ###.
            if p.style.name.lower() == 'title' and text.lower() == 'results':
                continue
            md_level = min(max(lvl, 1), 4)
            if p.style.name.lower() == 'title':
                md_level = 1
            out.append('')
            out.append('#' * (md_level + 1) + ' ' + text)  # +1 keeps doc title at H1
            out.append('')
        else:
            # Bold-leading "Figure X." caption (no image preceded — orphan caption)
            if _caption_rx.match(text):
                first_period = text.find('.')
                if first_period > 0:
                    label = text[:first_period+1]
                    rest = text[first_period+1:].strip()
                    out.append(f'*{label}* {rest}')
                else:
                    out.append(f'*{text}*')
            else:
                out.append(text)
            out.append('')

    # collapse 3+ blank lines to 2
    md = '\n'.join(out)
    md = re.sub(r'\n{3,}', '\n\n', md)
    with open(OUT, 'w', encoding='utf-8') as fh:
        fh.write(md)
    print(f'wrote {OUT} ({len(md)} chars)')


if __name__ == '__main__':
    main()
