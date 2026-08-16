"""
Convert Experiment1_StanceContracts_Paper.md → .docx with embedded figures,
tables, and basic markdown formatting (bold/italic, headings, blockquotes,
inline code, ordered/unordered lists).
"""
from __future__ import annotations
import os, re, sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
SRC = os.path.join(ROOT, 'Experiment1_StanceContracts_Paper.md')
OUT = os.path.join(ROOT, 'Experiment1_StanceContracts_Paper.docx')


# inline-formatting regex: bold, italic, inline code, bold-italic
_INLINE_RX = re.compile(
    r'(\*\*\*[^*]+?\*\*\*'                         # ***bi***
    r'|\*\*[^*]+?\*\*'                             # **b**
    r'|__[^_]+?__'                                 # __b__
    r'|\*[^*]+?\*'                                 # *i*
    r'|_[^_]+?_'                                   # _i_
    r'|`[^`]+?`'                                   # `code`
    r')'
)

_IMG_RX = re.compile(r'!\[([^\]]*?)\]\(([^)]+?)\)')


def add_inline_runs(p, text: str):
    """Append runs to paragraph p, parsing inline ** _ * ` markers."""
    parts = _INLINE_RX.split(text)
    for tok in parts:
        if not tok:
            continue
        if tok.startswith('***') and tok.endswith('***'):
            r = p.add_run(tok[3:-3])
            r.bold = True; r.italic = True
        elif (tok.startswith('**') and tok.endswith('**')) or \
             (tok.startswith('__') and tok.endswith('__')):
            r = p.add_run(tok[2:-2])
            r.bold = True
        elif (tok.startswith('*') and tok.endswith('*')) or \
             (tok.startswith('_')  and tok.endswith('_')):
            r = p.add_run(tok[1:-1])
            r.italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = p.add_run(tok[1:-1])
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        else:
            p.add_run(tok)


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith('|') and s.endswith('|') and s.count('|') >= 2


def is_table_separator(line: str) -> bool:
    s = line.strip()
    if not (s.startswith('|') and s.endswith('|')): return False
    inner = s.strip('|')
    cells = [c.strip() for c in inner.split('|')]
    return all(re.fullmatch(r':?-+:?', c) for c in cells if c)


def split_table_row(line: str) -> list[str]:
    inner = line.strip().strip('|')
    return [c.strip() for c in inner.split('|')]


def add_table(doc, rows: list[list[str]]):
    if not rows: return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            text = row[ci] if ci < len(row) else ''
            cell.text = ''   # clear default
            p = cell.paragraphs[0]
            add_inline_runs(p, text)
            if ri == 0:
                for r in p.runs:
                    r.bold = True


def resolve_image_path(rel: str) -> str:
    """Try the path as-is, then relative to ROOT."""
    if os.path.isabs(rel) and os.path.exists(rel):
        return rel
    candidate = os.path.join(ROOT, rel)
    return candidate if os.path.exists(candidate) else rel


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    with open(SRC, encoding='utf-8') as fh:
        lines = fh.read().split('\n')

    i = 0
    pending_table: list[list[str]] = []

    def flush_table():
        nonlocal pending_table
        if pending_table:
            add_table(doc, pending_table)
            pending_table = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ---- table accumulation ----
        if is_table_row(stripped):
            if is_table_separator(stripped):
                # skip the |---|---| separator after consuming it
                i += 1
                continue
            pending_table.append(split_table_row(stripped))
            i += 1
            continue
        else:
            flush_table()

        # ---- empty line ----
        if not stripped:
            i += 1
            continue

        # ---- horizontal rule ----
        if stripped in ('---', '***', '___'):
            doc.add_paragraph().add_run('').add_break()
            i += 1
            continue

        # ---- headings ----
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # title (single # at file top) becomes Title style
            if level == 1:
                p = doc.add_paragraph(style='Title')
                add_inline_runs(p, text)
            else:
                # Word's Heading X styles map nicely; cap at 3
                level = min(level - 1, 3)
                p = doc.add_paragraph(style=f'Heading {level}')
                add_inline_runs(p, text)
            i += 1
            continue

        # ---- image-only line ----
        img_match = _IMG_RX.match(stripped)
        if img_match and stripped == img_match.group(0):
            alt = img_match.group(1)
            rel = img_match.group(2)
            path = resolve_image_path(rel)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=Inches(6.0))
            else:
                p = doc.add_paragraph()
                p.add_run(f'[missing image: {rel}]').italic = True
            i += 1
            continue

        # ---- blockquote ----
        if stripped.startswith('>'):
            # consume contiguous blockquote lines
            buf = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                buf.append(lines[i].strip().lstrip('>').strip())
                i += 1
            text = ' '.join(buf)
            p = doc.add_paragraph(style='Intense Quote')
            add_inline_runs(p, text)
            continue

        # ---- list items (- * + or 1.) ----
        if re.match(r'^(\-|\*|\+)\s+', stripped) or re.match(r'^\d+\.\s+', stripped):
            buf_items = []
            ordered = bool(re.match(r'^\d+\.', stripped))
            while i < len(lines) and (
                re.match(r'^\s*(\-|\*|\+)\s+', lines[i])
                or re.match(r'^\s*\d+\.\s+', lines[i])
            ):
                m1 = re.match(r'^\s*(?:\-|\*|\+|\d+\.)\s+(.*)$', lines[i])
                if m1: buf_items.append(m1.group(1))
                i += 1
            for item in buf_items:
                style = 'List Number' if ordered else 'List Bullet'
                try:
                    p = doc.add_paragraph(style=style)
                except Exception:
                    p = doc.add_paragraph(); p.add_run('• ')
                add_inline_runs(p, item)
            continue

        # ---- regular paragraph (may include inline images) ----
        # consume contiguous non-empty non-special lines
        buf = [stripped]
        i += 1
        while i < len(lines):
            s = lines[i].strip()
            if not s: break
            if (re.match(r'^#{1,6}\s', s)
                or s.startswith('>')
                or re.match(r'^(\-|\*|\+)\s', s)
                or re.match(r'^\d+\.\s', s)
                or is_table_row(s)
                or _IMG_RX.fullmatch(s)
                or s in ('---','***','___')):
                break
            buf.append(s)
            i += 1
        text = ' '.join(buf)
        p = doc.add_paragraph()
        # if line contains an inline image, split around it
        last = 0
        for im in _IMG_RX.finditer(text):
            if im.start() > last:
                add_inline_runs(p, text[last:im.start()])
            ipath = resolve_image_path(im.group(2))
            if os.path.exists(ipath):
                p.add_run().add_picture(ipath, width=Inches(5.0))
            else:
                p.add_run(f'[missing image: {im.group(2)}]').italic = True
            last = im.end()
        if last < len(text):
            add_inline_runs(p, text[last:])

    flush_table()
    doc.save(OUT)
    print(f'wrote {OUT}')
    # report
    n_imgs = sum(1 for _ in re.finditer(r'shape', open(OUT,'rb').read().decode('latin-1', 'ignore')))


if __name__ == '__main__':
    main()
