"""Insert cross-model validation (Qwen Scorer A vs Gemini Scorer C) into:

  1. Experiment1_Results.docx       — new Heading-1 section before Summary
  2. Experiment1_StanceContracts_Paper_v4.docx
                                    — new section 3.10 in Results, plus
                                      addendum to 4.6 Limitations

Source numbers come from regulated_llm_reanalysis/14_*, 15_*, 16_*.
Forbidden-language audit runs after each save.
"""
from __future__ import annotations
import os
import sys
import warnings
warnings.filterwarnings('ignore')
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
OUT = os.path.join(ROOT, 'regulated_llm_reanalysis')
RESULTS_DOC = os.path.join(ROOT, 'Experiment1_Results.docx')
PAPER_DOC = os.path.join(ROOT, 'Experiment1_StanceContracts_Paper_v4.docx')


# ---------- helpers ----------

def insert_paragraph_before(target_para, text: str, style: str):
    """Insert a new paragraph immediately before target_para, preserving style.
    python-docx's add_paragraph appends at end, so we manipulate XML directly."""
    new_p = target_para._parent.add_paragraph(text, style=style)
    target_el = target_para._element
    new_el = new_p._element
    target_el.addprevious(new_el)
    # The add_paragraph call left a stray element at end; remove it.
    # (Actually add_paragraph only emits one element; addprevious moves it.)
    return new_p


def add_table_after(para, headers: list[str], rows: list[list[str]]):
    """Add a docx table immediately after `para`. Returns the table."""
    doc = para._parent
    if hasattr(doc, 'part'):
        d = doc
    else:
        # paragraph parent in a Document is _Body; need its document
        d = para._parent._parent if hasattr(para._parent, '_parent') else para._parent
    # Use the document's add_table at end, then move after `para`.
    # We get Document via the paragraph element's owning document.
    body = para._element.getparent()
    # Build the table by appending then moving.
    # Find the Document via paragraph traversal:
    return None  # Tables handled differently below.


def append_table(d: Document, headers: list[str], rows: list[list[str]]):
    """Append a small table at end of doc with a Light-Grid look."""
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        t.style = 'Light Grid Accent 1'
    except KeyError:
        try:
            t.style = 'Table Grid'
        except KeyError:
            pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = h
        for r in c.paragraphs:
            for run in r.runs:
                run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    return t


def move_table_after_paragraph(t, para):
    """Move docx table element to immediately follow `para`."""
    para._element.addnext(t._element)


def forbidden_audit(doc_path):
    forbid = ['true creativity', 'ground-truth creativity',
              'validated creativity improvement', 'objectively more creative',
              'made participants more creative']
    d = Document(doc_path)
    full = '\n'.join(p.text for p in d.paragraphs)
    hits = [t for t in forbid if t in full.lower()]
    print(f'  forbidden-language hits: {hits or "none"}')


# ---------- load source numbers ----------

def load_summary():
    side = pd.read_csv(os.path.join(OUT, '16_qwen_vs_gemini_sidebyside.csv'))
    between = pd.read_csv(os.path.join(OUT, '14_between_group_family_vs_gpt.csv'))
    within = pd.read_csv(os.path.join(OUT, '15_within_arm_pairwise.csv'))
    agree = pd.read_csv(os.path.join(OUT, '13_cross_model_agreement.csv'))
    return side, between, within, agree


def family_vs_gpt_summary(side: pd.DataFrame) -> list[list[str]]:
    rows = []
    for fam in ('Divergent', 'Convergent', 'Rational', 'BoundedRational'):
        s = side[side.contrast == f'{fam} vs GPT']
        n_rows = len(s)
        sm = int(s['sign_match'].sum())
        bs = int(s['both_sig'].sum())
        sq = int(s['sig_Qwen'].sum())
        sg = int(s['sig_Gemini'].sum())
        rows.append([
            f'{fam} vs GPT', f'{sm}/{n_rows}', str(bs),
            str(sq), str(sg),
        ])
    return rows


def family_vs_gpt_g_table(side: pd.DataFrame) -> list[list[str]]:
    """Per-criterion Hedges' g for each family vs GPT, side by side."""
    headline_crit = (
        'reframing_quality', 'exploration_opening', 'anchor_management',
        'timing_fit', 'stance_integrity', 'coregulation_uptake',
        'agency_preservation', 'evaluative_discipline', 'cognitive_load_clarity',
        'implementation_grounding',
        'premature_convergence_risk', 'runaway_divergence_risk',
    )
    rows = []
    for crit in headline_crit:
        cells = [crit]
        for fam in ('Divergent', 'Convergent', 'Rational', 'BoundedRational'):
            r = side[(side.criterion == crit)
                     & (side.contrast == f'{fam} vs GPT')]
            if len(r) == 0:
                cells.extend(['—', '—'])
                continue
            r = r.iloc[0]
            qg = f"{r.g_Qwen:+.2f}{'*' if r.sig_Qwen else ''}"
            gg = f"{r.g_Gemini:+.2f}{'*' if r.sig_Gemini else ''}"
            cells.extend([qg, gg])
        rows.append(cells)
    return rows


def divergent_vs_convergent_table(within: pd.DataFrame) -> list[list[str]]:
    sub = within[within.contrast == 'Divergent vs Convergent']
    rows = []
    for crit in sorted(sub.criterion.unique()):
        a = sub[(sub.criterion == crit) & (sub.scorer == 'A')]
        c = sub[(sub.criterion == crit) & (sub.scorer == 'C')]
        if len(a) == 0 or len(c) == 0:
            continue
        a = a.iloc[0]; c = c.iloc[0]
        rows.append([
            crit,
            f"{a.hedges_g:+.2f}{'*' if a.sig_05 else ''}",
            f"{c.hedges_g:+.2f}{'*' if c.sig_05 else ''}",
            '✓' if (a.hedges_g > 0) == (c.hedges_g > 0) else '✗',
        ])
    return rows


# ---------- Results.docx update ----------

RESULTS_HEADING = 'Cross-Model Validation: Gemini Scorer C'


def update_results_doc():
    side, between, within, agree = load_summary()
    d = Document(RESULTS_DOC)
    # find Summary heading position
    paras = list(d.paragraphs)
    summary_para = None
    for p in paras:
        if p.style.name.startswith('Heading 1') and p.text.strip().lower() == 'summary':
            summary_para = p; break
    if summary_para is None:
        raise RuntimeError('Summary heading not found in Results doc')

    # text content
    H1 = RESULTS_HEADING
    intro = (
        "To address the single-model limitation noted in the regulated "
        "reanalysis (the original Scorer A and Scorer B both ran on Qwen3-4B "
        "with paraphrased prompts), the 12-criterion rubric was rerun on the "
        "full 200-episode stratified sample with Gemini 3.1 Flash Lite "
        "Preview as Scorer C, holding Scorer A's prompt template constant. "
        "A vs B isolates prompt variance; A vs C isolates model variance. "
        "Scorer C is informational and does not modify the published "
        "adjudicated scores."
    )
    p1 = (
        "Pooled across 458 dual-scored rows, Qwen and Gemini agree at "
        "quadratic-weighted Cohen's κ = 0.69 and Spearman ρ = 0.77 (p < 1e-89). "
        "By Landis-Koch, this is substantial agreement on ordinal rubric ratings."
    )
    p2 = (
        "Direction agreement on the headline between-group contrasts (each "
        "persona family vs GPT) is 47/48 (97.9%) across the 12 criteria × 4 "
        "families. The single discordance is implementation_grounding for "
        "Divergent vs GPT (Qwen g = +0.12, Gemini g = -0.10); both effects are "
        "near-zero and neither reaches p < 0.05 in either scorer. On joint "
        "significance (both p < 0.05), 34/48 contrasts replicate. Significance "
        "is reported at raw p < 0.05; no multiple-comparison correction is "
        "applied."
    )
    p3 = (
        "Gemini also closes the GPT-cell coverage asymmetry recorded in the "
        "original reanalysis. Scorer A produced 369 GPT criterion-rows out of "
        "an expected 600 (61.5% coverage); Scorer C produced 600/600. This is "
        "consistent with the asymmetry being a Qwen-side bundle-validity "
        "artifact (HANDOFF Limitation 4) rather than a property of the data."
    )
    p4_h2 = 'Per-family direction-match summary'
    p5 = (
        "Each row counts how many of the 12 criteria show direction agreement, "
        "joint significance at raw p < 0.05, and one-model-only significance "
        "for that family-vs-GPT contrast."
    )
    p6_h2 = "Per-criterion effect sizes (Hedges' g): each persona family vs GPT"
    p7 = (
        "Asterisk (*) marks raw p < 0.05 (no multiple-comparison correction). "
        "Positive g means the persona family scored higher than GPT on that "
        "criterion. For the two RISK criteria (premature_convergence_risk, "
        "runaway_divergence_risk) positive g means MORE risk; negative is the "
        "favourable direction."
    )
    p8_h2 = 'Within Persona arm: Divergent vs Convergent'
    p9 = (
        "The Divergent-Convergent contrast is the strongest within-arm "
        "comparison in both scorers, distinguishing 7 (Qwen) and 7 (Gemini) of "
        "12 criteria at p < 0.05. The trade-off signature reproduces in both "
        "models: Divergent shows higher exploration_opening, reframing_quality, "
        "stance_integrity, and runaway_divergence_risk; lower "
        "premature_convergence_risk and cognitive_load_clarity. The remaining "
        "pairwise within-arm contrasts (Rational, BoundedRational) show few or "
        "no differences reaching p < 0.05 in either scorer."
    )
    p10 = (
        "Together these results upgrade the original prompt-paraphrase "
        "robustness check into a true cross-model robustness check. The "
        "regulated-rubric findings reported earlier in this document are not "
        "an artifact of the Qwen scorer: an independent commercial model "
        "agrees on direction, on significance at p < 0.05 for the dominant "
        "criteria, and on the magnitude ranking among persona families."
    )

    # build paragraphs in reverse insertion order (each insert_before brings the
    # element right above summary_para; later inserts continue to push them up)
    # so we accumulate in order, inserting before summary, while tracking the
    # last-inserted paragraph as the new "anchor" for any tables.

    paragraphs_to_insert = [
        ('Heading 1', H1),
        ('Normal', intro),
        ('Normal', p1),
        ('Normal', p2),
        ('Normal', p3),
        ('Heading 2', p4_h2),
        ('Normal', p5),
    ]

    last_inserted = None
    for style, text in paragraphs_to_insert:
        np_ = insert_paragraph_before(summary_para, text, style)
        last_inserted = np_

    # Table 1: per-family summary
    rows1 = family_vs_gpt_summary(side)
    headers1 = ['contrast', 'sign-match', 'both sig (p<0.05)',
                'Qwen-A sig (p<0.05)', 'Gemini-C sig (p<0.05)']
    t1 = append_table(d, headers1, rows1)
    move_table_after_paragraph(t1, last_inserted)

    # following paragraphs and table 2
    for style, text in [('Heading 2', p6_h2), ('Normal', p7)]:
        np_ = insert_paragraph_before(summary_para, text, style)
        last_inserted = np_

    rows2 = family_vs_gpt_g_table(side)
    headers2 = ['criterion',
                'Div Qwen', 'Div Gem',
                'Con Qwen', 'Con Gem',
                'Rat Qwen', 'Rat Gem',
                'BR Qwen', 'BR Gem']
    t2 = append_table(d, headers2, rows2)
    move_table_after_paragraph(t2, last_inserted)

    for style, text in [('Heading 2', p8_h2), ('Normal', p9)]:
        np_ = insert_paragraph_before(summary_para, text, style)
        last_inserted = np_

    rows3 = divergent_vs_convergent_table(within)
    headers3 = ['criterion', 'Qwen-A g', 'Gemini-C g', 'sign match']
    t3 = append_table(d, headers3, rows3)
    move_table_after_paragraph(t3, last_inserted)

    insert_paragraph_before(summary_para, p10, 'Normal')

    d.save(RESULTS_DOC)
    print(f'updated {RESULTS_DOC}')
    forbidden_audit(RESULTS_DOC)


# ---------- Paper v4 update ----------

PAPER_NEW_SECTION_HEAD = '3.10 Cross-model validation of the regulated rubric'

PAPER_NEW_SECTION_BODY = [
    'A reviewer concern about the regulated reanalysis is that Scorer A and '
    'Scorer B both ran on the same local model (Qwen3-4B-Instruct INT4); the '
    'A/B comparison probes prompt-paraphrase robustness but not model '
    'robustness. We added a third scorer using a different model family — '
    'Gemini 3.1 Flash Lite Preview, accessed via API — holding the Scorer A '
    'prompt template constant so that Qwen-A vs Gemini-C isolates model '
    'variance. The two scorers were applied to the same 200-episode '
    'stratified sample. Scorer C is informational and does not modify the '
    'adjudicated scores in 05_episode_rubric_scores_adjudicated.csv.',

    'Pooled across 458 dual-scored rows, Qwen and Gemini agree at '
    "quadratic-weighted Cohen's κ = 0.69 and Spearman ρ = 0.77, "
    'p < 10^-89 — substantial ordinal agreement by Landis-Koch.',

    'On the headline between-group contrasts (each persona family vs GPT, '
    '12 criteria × 4 families = 48 cells), the two scorers agree on the sign '
    'of the effect in 47/48 cases (97.9%). The single sign disagreement is '
    'implementation_grounding for Divergent vs GPT (g_Qwen = +0.12, '
    'g_Gemini = -0.10); both effects are near zero and neither is '
    'FDR-significant. Joint FDR-significance (both q < 0.05) holds for '
    '33/48 cells. The dominant within-arm contrast — Divergent vs '
    'Convergent — distinguishes 6 (Qwen) and 7 (Gemini) of 12 criteria at '
    'q < 0.05, with the same trade-off signature: Divergent has higher '
    'exploration_opening, reframing_quality, and runaway_divergence_risk, '
    'and lower premature_convergence_risk.',

    'A secondary observation: Scorer A produced bundle-valid output for '
    '369 of 600 GPT criterion-rows (61.5%), while Scorer C produced 600/600. '
    'The original GPT-vs-Persona invalid-rate asymmetry (HANDOFF Limitation '
    '4) is therefore a Qwen-side bundle-validity artifact, not a property of '
    'the data. Persona-vs-GPT effects under Scorer C use balanced cell sizes; '
    'they reproduce the Qwen-A direction throughout.',

    'Net effect on the paper\'s claim structure: the regulated-rubric '
    'findings replicate across two independent scorers, one local and one '
    'commercial, on bounded process constructs. The proxy-bounded language '
    'is unchanged — no human creativity raters were involved in either '
    'scoring run — but the multi-model agreement removes the single-model '
    'caveat from the reanalysis layer.',
]

PAPER_LIMITATION_ADDENDUM = (
    'Update from cross-model validation: the single-model limitation has '
    'been partly addressed by adding a third scorer (Gemini 3.1 Flash Lite '
    'Preview) on the same 200-episode stratified sample, holding the '
    'Scorer A prompt template constant. Qwen-A and Gemini-C agree at '
    "quadratic-weighted κ = 0.69 / Spearman ρ = 0.77 pooled, with sign "
    'agreement on 47/48 family-vs-GPT × criterion contrasts. The remaining '
    'limitation — that no human raters scored the rubric — is unchanged.'
)


def update_paper_doc():
    side, between, within, agree = load_summary()
    d = Document(PAPER_DOC)

    # 1) Insert section 3.10 before "4. Discussion" (Heading 1)
    paras = list(d.paragraphs)
    discussion_para = None
    for p in paras:
        if p.style.name.startswith('Heading 1') and p.text.strip().startswith('4. Discussion'):
            discussion_para = p; break
    if discussion_para is None:
        raise RuntimeError('Could not find "4. Discussion" heading')

    last_inserted = None
    for style, text in [('Heading 2', PAPER_NEW_SECTION_HEAD)] + [('Normal', t) for t in PAPER_NEW_SECTION_BODY]:
        np_ = insert_paragraph_before(discussion_para, text, style)
        last_inserted = np_

    # Append a small per-family summary table just after the section.
    rows1 = family_vs_gpt_summary(side)
    headers1 = ['contrast', 'sign-match', 'both FDR-sig',
                'Qwen-A FDR-sig (of 12)', 'Gemini-C FDR-sig (of 12)']
    t1 = append_table(d, headers1, rows1)
    move_table_after_paragraph(t1, last_inserted)

    # 2) Append addendum to 4.6 Limitations
    paras = list(d.paragraphs)  # refresh
    lim_para = None
    for p in paras:
        if p.style.name.startswith('Heading 2') and p.text.strip().startswith('4.6 Limitations'):
            lim_para = p; break
    if lim_para is None:
        raise RuntimeError('Could not find "4.6 Limitations" heading')

    # Insert the addendum BEFORE the next Heading-2 (i.e., at end of 4.6).
    next_h2 = None
    paras_after = paras[paras.index(lim_para)+1:]
    for p in paras_after:
        if p.style.name.startswith('Heading 2'):
            next_h2 = p; break
    if next_h2 is None:
        d.add_paragraph(PAPER_LIMITATION_ADDENDUM, style='Normal')
    else:
        insert_paragraph_before(next_h2, PAPER_LIMITATION_ADDENDUM, 'Normal')

    d.save(PAPER_DOC)
    print(f'updated {PAPER_DOC}')
    forbidden_audit(PAPER_DOC)


# ---------- main ----------

if __name__ == '__main__':
    update_results_doc()
    # NOTE: update_paper_doc() is intentionally NOT called.
    # The paper integrates cross-model results into §3.8 via
    # update_paper_38_crossmodel.py instead of a standalone 3.10.
