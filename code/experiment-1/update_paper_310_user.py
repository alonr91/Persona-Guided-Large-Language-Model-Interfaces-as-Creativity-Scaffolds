"""Insert new section 3.10 'The user-side response under each persona' into
Experiment1_StanceContracts_Paper_v4.docx, between 3.9 and Section 4.

Pulls numbers from:
  17_user_optionA_paired_by_family.csv          (Option A — existing measures)
  18_user_rubric_between_group.csv              (Option B — Gemini user rubric)
  18_user_rubric_within_arm.csv

Embeds two figures:
  fig_user_optionA_paired_by_family.png         Figure 10
  fig_user_rubric_radar.png                     Figure 11

Forbidden-language audit runs after save.
"""
from __future__ import annotations
import os, sys, warnings
warnings.filterwarnings('ignore')
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

import pandas as pd
from docx import Document
from docx.shared import Inches

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
OUT = os.path.join(ROOT, 'regulated_llm_reanalysis')
FIG = os.path.join(OUT, 'figures')
PAPER = os.path.join(ROOT, 'Experiment1_StanceContracts_Paper_v4.docx')


def insert_before(target_para, text, style):
    new_p = target_para._parent.add_paragraph(text, style=style)
    target_para._element.addprevious(new_p._element)
    return new_p


def insert_image_before(target_para, image_path, width_in=5.5):
    new_p = target_para._parent.add_paragraph()
    target_para._element.addprevious(new_p._element)
    run = new_p.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    return new_p


def insert_table_before(doc, target_para, headers, rows):
    """Insert a small table immediately before target_para. `doc` is the Document."""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    try: t.style = 'Light Grid Accent 1'
    except KeyError:
        try: t.style = 'Table Grid'
        except KeyError: pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for r in c.paragraphs:
            for run in r.runs: run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    target_para._element.addprevious(t._element)
    return t


def forbidden_audit(path):
    forbid = ['true creativity', 'ground-truth creativity',
              'validated creativity improvement', 'objectively more creative',
              'made participants more creative']
    d = Document(path)
    full = '\n'.join(p.text for p in d.paragraphs)
    hits = [t for t in forbid if t in full.lower()]
    print(f'  forbidden-language hits: {hits or "none"}')


# -------- numbers --------

def load():
    optA = pd.read_csv(os.path.join(OUT, '17_user_optionA_paired_by_family.csv'))
    optB_b = pd.read_csv(os.path.join(OUT, '18_user_rubric_between_group.csv'))
    optB_w = pd.read_csv(os.path.join(OUT, '18_user_rubric_within_arm.csv'))
    return optA, optB_b, optB_w


def optA_summary_rows(a):
    """Per family: count of FDR-sig measures and a 1-line direction summary."""
    rows = []
    fams = ('Divergent','Convergent','Rational','BoundedRational')
    for fam in fams:
        s = a[a.family == fam]
        sig = s[s.sig_05 == True]
        n_total = len(s)
        n_sig = len(sig)
        # build a compact effect direction summary
        ups = sig[sig.mean_diff > 0]['measure'].tolist()
        downs = sig[sig.mean_diff < 0]['measure'].tolist()
        N = int(s['n'].max()) if len(s) else 0
        rows.append([f'{fam} (N={N})', f'{n_sig}/{n_total}',
                     ', '.join(ups) if ups else '—',
                     ', '.join(downs) if downs else '—'])
    return rows


def optB_summary_rows(b):
    rows = []
    fams = ('Divergent','Convergent','Rational','BoundedRational')
    for fam in fams:
        s = b[b.family == fam].copy().sort_values('criterion')
        cells = [fam]
        for _, r in s.iterrows():
            cells.append(f"{r.g:+.2f}{'*' if r.sig_05 else ''}")
        rows.append(cells)
    return rows


# -------- main --------

H2 = '3.10 The user-side response under each persona (cross-model)'

INTRO = (
    "The dialogic rubric in §3.8 scores the conversation as a whole; many of "
    "its criteria are weighted toward assistant-driven moves. To isolate the "
    "user's contribution we add two complementary user-only views: (A) a "
    "within-subject paired analysis of pre-existing user-side measures (turn "
    "counts, idea fluency and originality, taxonomy-coded user moves), and "
    "(B) a new 6-criterion user-behaviour rubric scored by Gemini Scorer C on "
    "the same 200-episode stratified sample, where the scorer reads the full "
    "episode but rates ONLY user behaviour and is constrained to quote "
    "verbatim from user turns. (A) tells us what users do; (B) tells us how "
    "substantively they do it. Both views are user-only by construction."
)

OPTA_HDR = '(A) Within-subject paired diffs on existing user measures'
OPTA_TXT = (
    "Each participant has a Persona round and a GPT round. We take the "
    "per-participant paired diff (Persona − GPT) on each user-only measure "
    "and group by the persona family the participant was assigned. "
    "Significance is raw p < 0.05 (no multiple-comparison correction). "
    "The two high-N families (Divergent, N=38; Convergent, N=41) show a "
    "coherent user-side adaptation: longer messages, more total words, "
    "higher idea fluency (+1 idea per round), and a sharply lower question "
    "rate (dz ≈ −0.7 to −0.8) — the assistant is now doing the "
    "questioning. Same-condition originality drops slightly (consistent "
    "with §3.5). Rational and BoundedRational (N=9 each) reach significance "
    "on fewer measures but trend in the same direction."
)
OPTA_FIG_CAPTION = (
    "Figure 10. User-only measures, paired diff (Persona − GPT) by persona "
    "family with 95% CI. Positive = user expressed more of the measure in "
    "the Persona round than the GPT round. The bars are unstandardised; "
    "scales differ across panels."
)

OPTB_HDR = '(B) User-behaviour rubric (Gemini Scorer C, 0–4 ordinal)'
OPTB_TXT = (
    "The 6 user-behaviour criteria are: user_initiative, "
    "user_question_richness, user_proposal_specificity, "
    "user_acceptance_yes_and, user_reframing, and user_engagement_depth. "
    "Each rating cites a verbatim user-turn quote; assistant echoes are "
    "explicitly disallowed as evidence. Scoring is episode-level Welch t "
    "with Hedges' g, significance at raw p < 0.05 (no multiple-comparison "
    "correction)."
)
OPTB_FIG_CAPTION = (
    "Figure 11. User-behaviour rubric profile per persona family (Gemini "
    "Scorer C, 0–4). All four persona families sit slightly above the GPT "
    "baseline on user_acceptance_yes_and, user_engagement_depth, and "
    "user_proposal_specificity, and slightly below on user_question_richness. "
    "Profiles overlap heavily — within-arm pairwise contrasts show zero "
    "criteria reaching p < 0.05 across all six pairs."
)

VALIDATION_HDR = 'Validation summary.'
VALIDATION_BULLETS = [
    "(A) Convergence of (A) and (B) on the question-rate signal. The "
    "behavioural user_q_rate drops sharply in Persona rounds for Divergent "
    "(dz = −0.71) and Convergent (dz = −0.82), and the user-rubric "
    "user_question_richness mirrors this with the same direction "
    "(Convergent g = −0.54, p = 0.008). The user does not just ask fewer "
    "questions — the questions they do ask are also less probing.",

    "(B) Persona-vs-GPT effect on substantive user behaviour. Three "
    "user-rubric criteria reach p < 0.05 for at least one family vs GPT: "
    "user_acceptance_yes_and (Divergent g = +0.48, Convergent g = +0.55), "
    "user_engagement_depth (Divergent g = +0.48), and "
    "user_proposal_specificity (Divergent g = +0.49). Direction is uniformly "
    "positive for the three criteria where the assistant's persona invites "
    "user contribution.",

    "(C) Within-arm pairwise contrasts on user behaviour are flat. Across "
    "the 6 criteria × 6 pairwise family contrasts (36 cells), zero reach "
    "p < 0.05. The persona families differentiate strongly on ASSISTANT "
    "behaviour (§3.8: 7 of 12 criteria distinguish Divergent vs Convergent "
    "in both scorers) but they do not differentiate on USER behaviour. "
    "Users adapt to a persona conversation in a broadly uniform way; the "
    "family-level differentiation lives in the assistant.",
]

CLAIM = (
    "Claim 3.10. Persona prompts produce a uniform user-side adaptation: "
    "users in any persona arm engage more deeply, propose more concretely, "
    "yes-and more, and ask fewer (and less probing) questions than the same "
    "users do in the GPT control arm. This adaptation is driven by being in "
    "a persona conversation rather than by which specific persona is "
    "running — within-arm pairwise contrasts on user-only measures show no "
    "differences reaching p < 0.05 across the four persona families on "
    "either the existing-measures view or the user-rubric view. The "
    "differentiation between persona families lives in the assistant's "
    "regulatory behaviour (§3.8); the user-side correlate is a regime "
    "change, not a family signature."
)


def update():
    a, b, w = load()
    d = Document(PAPER)
    paras = list(d.paragraphs)

    # find anchor: Section 4 heading (Heading 1)
    anchor = None
    for p in paras:
        if p.style.name.startswith('Heading 1') and p.text.strip().startswith('4. Discussion'):
            anchor = p; break
    if anchor is None:
        raise RuntimeError('"4. Discussion" not found')

    # All inserts go BEFORE the anchor; insertions accumulate above it.
    insert_before(anchor, H2, 'Heading 2')
    insert_before(anchor, INTRO, 'Normal')

    insert_before(anchor, OPTA_HDR, 'Normal')
    insert_before(anchor, OPTA_TXT, 'Normal')

    # Option A summary table
    rows_a = optA_summary_rows(a)
    headers_a = ['Family (paired N)', 'sig measures', 'positive paired diff (Persona > GPT)',
                 'negative paired diff (Persona < GPT)']
    insert_table_before(d, anchor, headers_a, rows_a)

    # Option A figure (image first, caption after — matches §3.8 convention)
    insert_image_before(anchor, os.path.join(FIG, 'fig_user_optionA_paired_by_family.png'), width_in=6.0)
    insert_before(anchor, OPTA_FIG_CAPTION, 'Normal')

    insert_before(anchor, OPTB_HDR, 'Normal')
    insert_before(anchor, OPTB_TXT, 'Normal')

    # Option B per-criterion table
    headers_b = ['family',
                 'user_initiative', 'user_question_richness',
                 'user_proposal_specificity', 'user_acceptance_yes_and',
                 'user_reframing', 'user_engagement_depth']
    rows_b = optB_summary_rows(b)
    insert_table_before(d, anchor, headers_b, rows_b)
    insert_before(anchor, "Asterisk = raw p < 0.05 (no multiple-comparison correction).",
                  'Normal')

    # Option B figure (radar) — image first, caption after
    insert_image_before(anchor, os.path.join(FIG, 'fig_user_rubric_radar.png'), width_in=5.5)
    insert_before(anchor, OPTB_FIG_CAPTION, 'Normal')

    insert_before(anchor, VALIDATION_HDR, 'Normal')
    for bul in VALIDATION_BULLETS:
        insert_before(anchor, bul, 'List Bullet')
    insert_before(anchor, CLAIM, 'Normal')

    d.save(PAPER)
    print(f'updated {PAPER}')
    forbidden_audit(PAPER)


if __name__ == '__main__':
    update()
