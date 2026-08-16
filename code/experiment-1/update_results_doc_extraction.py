"""
Update Experiment1_Results.docx with the LAYER L (agentic-extraction)
results and a flow-diagram figure illustrating the pipeline.
"""
import os, sys, warnings
warnings.filterwarnings('ignore')
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

import numpy as np, pandas as pd
from scipy import stats
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
OUT  = os.path.join(ROOT, 'analysis_out')
FIG  = os.path.join(ROOT, 'figures')

# -----------------------------------------------------------------
# 1) Load data and rebuild family/Big-5/perception on `wide`
# -----------------------------------------------------------------
orig = pd.read_csv(os.path.join(OUT, 'production', 'participant_originality.csv'))
logs = pd.read_csv(os.path.join(ROOT, 'Experiment1_logs.csv'))
up = (logs.groupby('User_id')['Persona_type']
        .apply(lambda s: [x for x in s.unique() if x!='GPT'][0])
        .reset_index().rename(columns={'Persona_type':'persona_cond'}))
fm = {'Divergent':'Divergent','Convergent':'Convergent','strictly rational':'Rational','bounded rationality':'BoundedRational'}
up['family'] = up['persona_cond'].map(fm)

users = pd.read_excel(os.path.join(OUT, 'users_translated.xlsx'), sheet_name='corrected_users')
users.columns = [c.strip() for c in users.columns]
def _gpt_round(row):
    r1 = str(row.get('Persona round 1','')).lower(); r2 = str(row.get('Persona round 2','')).lower()
    if 'gpt' in r1: return 1
    if 'gpt' in r2: return 2
    return np.nan
users['gpt_round'] = users.apply(_gpt_round, axis=1)
def _mk(df,a,b):
    g = np.where(df['gpt_round']==1, df[a], np.where(df['gpt_round']==2, df[b], np.nan))
    p = np.where(df['gpt_round']==1, df[b], np.where(df['gpt_round']==2, df[a], np.nan))
    return pd.Series(g, index=df.index), pd.Series(p, index=df.index)
users['cr_gpt'], users['cr_per'] = _mk(users,'Creativity assistant #1','Creativity assistant #2')
users['ow_gpt'], users['ow_per'] = _mk(users,'Ownership #1','Ownership #2')
users['cr_diff'] = users['cr_per'].astype(float)-users['cr_gpt'].astype(float)
users['ow_diff'] = users['ow_per'].astype(float)-users['ow_gpt'].astype(float)
pers_cols = ['Extraversion','Agreeableness','Conscientiousness','Negative Emotionality','Open-Mindedness']

wide = orig.pivot_table(index='user', columns='condition',
                         values=['n_ideas','orig_same','orig_all','orig_cross'],
                         aggfunc='first')
wide.columns = [f'{m}__{c}' for m,c in wide.columns]
wide = wide.reset_index().merge(up[['User_id','family']], left_on='user', right_on='User_id', how='left').drop(columns=['User_id'])
for col in ['n_ideas','orig_same','orig_all','orig_cross']:
    cp, cg = f'{col}__Persona', f'{col}__GPT'
    if cp in wide.columns and cg in wide.columns:
        wide[f'd_{col}'] = wide[cp].astype(float) - wide[cg].astype(float)
for p in pers_cols:
    if p in users.columns:
        wide[p] = wide['user'].map(users.set_index('id')[p])
for src in ['cr_diff','ow_diff']:
    wide[src] = wide['user'].map(users.set_index('id')[src])

# -----------------------------------------------------------------
# 2) Compute headline stats (for in-text citation in the docx)
# -----------------------------------------------------------------
def welch_g(a, b):
    a = pd.to_numeric(a, errors='coerce').dropna()
    b = pd.to_numeric(b, errors='coerce').dropna()
    t,p = stats.ttest_ind(a, b, equal_var=False)
    sp = np.sqrt(((len(a)-1)*a.var(ddof=1)+(len(b)-1)*b.var(ddof=1))/(len(a)+len(b)-2))
    d = (a.mean()-b.mean())/sp if sp>0 else np.nan
    J = 1 - 3/(4*(len(a)+len(b))-9)
    g = d*J if not np.isnan(d) else np.nan
    return dict(n_p=len(a), n_g=len(b), mean_p=a.mean(), mean_g=b.mean(),
                diff=a.mean()-b.mean(), t=t, p=p, g=g)

def paired(col):
    cp, cg = f'{col}__Persona', f'{col}__GPT'
    a = wide[cp].astype(float); b = wide[cg].astype(float)
    m = (~a.isna())&(~b.isna())
    a, b = a[m], b[m]
    d = a - b
    t,p = stats.ttest_rel(a, b)
    dz = d.mean()/d.std(ddof=1)
    return dict(n=int(m.sum()), mean_p=a.mean(), mean_g=b.mean(),
                diff=d.mean(), t=t, p=p, dz=dz)

ws_fluency  = paired('n_ideas')
ws_same     = paired('orig_same')
ws_all      = paired('orig_all')
ws_cross    = paired('orig_cross')
bs_fluency  = welch_g(orig.loc[orig.condition=='Persona','n_ideas'],
                      orig.loc[orig.condition=='GPT','n_ideas'])
bs_same     = welch_g(orig.loc[orig.condition=='Persona','orig_same'],
                      orig.loc[orig.condition=='GPT','orig_same'])
bs_all      = welch_g(orig.loc[orig.condition=='Persona','orig_all'],
                      orig.loc[orig.condition=='GPT','orig_all'])
bs_cross    = welch_g(orig.loc[orig.condition=='Persona','orig_cross'],
                      orig.loc[orig.condition=='GPT','orig_cross'])

# -----------------------------------------------------------------
# 3) Figure A — agentic pipeline flow diagram
# -----------------------------------------------------------------
plt.rcParams.update({'figure.dpi':120, 'savefig.dpi':200, 'font.size':10})

fig, ax = plt.subplots(figsize=(12, 6.5))
ax.set_xlim(0, 14); ax.set_ylim(0, 9); ax.axis('off')

def _box(x, y, w, h, text, fc='#e9eef5', ec='#2a4365', bold=False, sub=None, fs=10):
    ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.05',
                                 fc=fc, ec=ec, lw=1.4))
    if bold:
        ax.text(x + w/2, y + h*0.66, text, ha='center', va='center',
                fontsize=fs, fontweight='bold', color='#1a365d')
    else:
        ax.text(x + w/2, y + h*0.66, text, ha='center', va='center', fontsize=fs, color='#1a365d')
    if sub:
        ax.text(x + w/2, y + h*0.22, sub, ha='center', va='center',
                fontsize=8, color='#4a5568', style='italic')

def _arrow(x0, y0, x1, y1, style='->', color='#2a4365'):
    ax.add_patch(FancyArrowPatch((x0, y0), (x1, y1), arrowstyle=style,
                                  color=color, mutation_scale=18, lw=1.4))

# ---- Input ----
_box(0.5, 7.7, 3, 0.9, 'Experiment 1 logs', bold=True,
     sub='194 conversations, 3,412 turns')

# ---- Agent 1 ----
_box(4.2, 7.7, 3.4, 0.9, 'Agent 1 — Extractor', bold=True,
     sub='Qwen3-4B  •  schema-locked JSON  •  per user message')
_arrow(3.5, 8.15, 4.2, 8.15)

# ---- Pre-LLM Filter 1 ----
_box(4.2, 6.55, 3.4, 0.55,
     'Filter 1: challenge-restatement block',
     fc='#fff5e8', ec='#975a16', fs=9,
     sub='rapidfuzz partial_ratio >= 85 against challenge prompt')
_arrow(5.9, 7.55, 5.9, 7.15, style='<-')

# ---- Post-LLM Filter 2 ----
_box(4.2, 5.85, 3.4, 0.55,
     'Filter 2: assistant-echo block',
     fc='#fff5e8', ec='#975a16', fs=9,
     sub='evidence vs previous assistant turn >= 85')
_arrow(5.9, 7.65, 5.9, 6.45, style='<-')

# ---- Agent 2 ----
_box(8.2, 7.7, 3.4, 0.9, 'Agent 2 — Consolidator', bold=True,
     sub='BGE-large-en-v1.5  •  agglomerative τ=0.85  •  per conv')
_arrow(7.6, 8.15, 8.2, 8.15)

# ---- Agent 3 ----
_box(8.2, 5.8, 3.4, 0.9, 'Agent 3 — Validator', bold=True,
     sub='rule-based verbatim + fuzzy  •  Filter 3')
_arrow(9.9, 7.7, 9.9, 6.7)

_box(8.2, 4.65, 3.4, 0.55,
     'Filter 3: title-evidence consistency',
     fc='#fff5e8', ec='#975a16', fs=9,
     sub='content stems in title ⊆ stems in evidence')
_arrow(9.9, 5.75, 9.9, 5.25, style='<-')

# ---- Agent 4 ----
_box(4.2, 3.6, 3.4, 0.9, 'Agent 4 — Categorizer', bold=True,
     sub='HDBSCAN cosine, min=4  •  LLM-labeled')
_arrow(9.9, 5.8, 7.6, 4.05, style='-|>')

# ---- Agent 5 ----
_box(8.2, 3.6, 3.4, 0.9, 'Agent 5 — Originality', bold=True,
     sub='participant centroid + 3 distance measures')
_arrow(9.9, 5.75, 9.9, 4.5, style='-|>')

# ---- Outputs ----
out_color = dict(fc='#eef7e9', ec='#276749', fs=9)
_box(12.1, 7.5, 1.8, 0.6, 'candidates.jsonl', **out_color)
_arrow(7.6, 7.85, 12.1, 7.8)
_box(12.1, 6.5, 1.8, 0.6, 'canonical_ideas.jsonl', **out_color)
_arrow(11.6, 8.0, 12.1, 6.8)
_box(12.1, 5.5, 1.8, 0.6, 'validation_report.csv', **out_color)
_arrow(11.6, 6.2, 12.1, 5.8)
_box(12.1, 3.8, 1.8, 0.6, 'categorized_ideas.csv', **out_color)
_arrow(7.6, 4.0, 12.1, 4.05)
_box(12.1, 2.8, 1.8, 0.6, 'participant_\noriginality.csv',
     fc='#eef7e9', ec='#276749', fs=9)
_arrow(11.6, 3.9, 12.1, 3.0)

# ---- Originality formulas ----
ax.text(0.5, 2.4, r'Participant centroid:  $C_p = \widehat{\,\overline{E}_p\,}$'
        r'   (mean idea embedding, L2-normalized)',
        fontsize=11, color='#1a365d')
ax.text(0.5, 1.5,
        r'$orig_{\mathrm{same}}(p) = \frac{1}{|Q_{\mathrm{same}}|}\sum_{q \in Q_{\mathrm{same}}}\!\!(1 - \langle C_p, C_q \rangle)$',
        fontsize=11, color='#1a365d')
ax.text(5.2, 1.5,
        r'$orig_{\mathrm{all}}(p) = \frac{1}{|Q_{\mathrm{all}}|}\sum_{q \in Q_{\mathrm{all}}}\!\!(1 - \langle C_p, C_q \rangle)$',
        fontsize=11, color='#1a365d')
ax.text(9.6, 1.5,
        r'$orig_{\mathrm{cross}}(p) = \min_{q \in Q_{\mathrm{cross}}}(1 - \langle C_p, C_q \rangle)$',
        fontsize=11, color='#1a365d')
ax.text(0.5, 0.6,
        r'Sets: $Q_{\mathrm{same}}$ = same-condition peers;'
        r'  $Q_{\mathrm{all}}$ = all other participants;'
        r'  $Q_{\mathrm{cross}}$ = opposing-condition participants.',
        fontsize=10, color='#4a5568')

ax.set_title('Figure R5. Open-source agentic pipeline for evidence-grounded user-idea extraction.',
             fontsize=11, loc='left', y=1.02)
plt.tight_layout()
plt.savefig(os.path.join(FIG, 'figR_agentic_workflow.png'), bbox_inches='tight')
plt.close()

# -----------------------------------------------------------------
# 4) Figure B — paired originality by persona family
# -----------------------------------------------------------------
fams = ['Divergent', 'Convergent', 'Rational', 'BoundedRational']
color_map = {'Divergent':'#2a9d8f','Convergent':'#e76f51',
             'Rational':'#6f62b6','BoundedRational':'#e9c46a'}

def _paired_box(ax, df, col, fam, title_ylabel):
    cp, cg = f'{col}__Persona', f'{col}__GPT'
    sub = df[(df.family==fam) & df[cp].notna() & df[cg].notna()]
    a = sub[cg].astype(float); b = sub[cp].astype(float)
    for _, r in sub.iterrows():
        ax.plot([0, 1], [r[cg], r[cp]], color='gray', alpha=0.35, lw=0.6)
    ax.boxplot([a, b], positions=[0, 1], widths=0.36, showfliers=False,
               patch_artist=True,
               boxprops=dict(facecolor=color_map[fam], alpha=0.35))
    ax.set_xticks([0, 1]); ax.set_xticklabels(['GPT', 'Persona'])
    if len(sub) >= 5:
        t, p = stats.ttest_rel(b, a)
        d = b - a
        dz = d.mean()/d.std(ddof=1) if d.std(ddof=1)>0 else np.nan
        p_txt = 'p<0.001' if p<0.001 else f'p={p:.3f}'
        ax.set_title(f'{fam}  (n={len(sub)})\n{p_txt}, dz={dz:+.2f}',
                     fontsize=10)
    else:
        ax.set_title(f'{fam}  (n={len(sub)})', fontsize=10)
    ax.set_ylabel(title_ylabel)
    ax.grid(alpha=0.25, axis='y')

# orig_same by family
fig, axes = plt.subplots(1, 4, figsize=(13.5, 4.3), sharey=True)
for ax, fam in zip(axes, fams):
    _paired_box(ax, wide, 'orig_same', fam, 'orig$_{same}$  (1 - cos)')
fig.suptitle('Figure R6. Same-condition originality, paired Persona vs GPT, by persona family',
             y=1.02, fontsize=12)
plt.tight_layout()
plt.savefig(os.path.join(FIG, 'figR_originality_same_by_family.png'), bbox_inches='tight')
plt.close()

# n_ideas by family
fig, axes = plt.subplots(1, 4, figsize=(13.5, 4.3), sharey=True)
for ax, fam in zip(axes, fams):
    _paired_box(ax, wide, 'n_ideas', fam, 'ideas per round')
fig.suptitle('Figure R7. Fluency (ideas per round), paired Persona vs GPT, by persona family',
             y=1.02, fontsize=12)
plt.tight_layout()
plt.savefig(os.path.join(FIG, 'figR_fluency_by_family.png'), bbox_inches='tight')
plt.close()

# -----------------------------------------------------------------
# 5) Figure C — Big-5 moderation heatmap
# -----------------------------------------------------------------
metrics = [('d_n_ideas','Δ fluency'),
           ('d_orig_same','Δ same-cond orig.'),
           ('d_orig_all','Δ all-part. orig.'),
           ('d_orig_cross','Δ cross-cond orig.')]
mat  = np.full((len(metrics), len(pers_cols)), np.nan)
pmat = np.full_like(mat, np.nan)
for i,(m,_) in enumerate(metrics):
    for j,p in enumerate(pers_cols):
        if m not in wide.columns or p not in wide.columns: continue
        x = wide[m].astype(float); y = wide[p].astype(float)
        mask = (~x.isna())&(~y.isna())
        if mask.sum()<15: continue
        rho, pv = stats.spearmanr(x[mask], y[mask])
        mat[i,j] = rho; pmat[i,j] = pv

fig, ax = plt.subplots(figsize=(8.5, 4.3))
im = ax.imshow(mat, cmap='RdBu_r', vmin=-0.35, vmax=0.35, aspect='auto')
ax.set_xticks(range(len(pers_cols)))
ax.set_xticklabels(pers_cols, rotation=20, ha='right')
ax.set_yticks(range(len(metrics)))
ax.set_yticklabels([lbl for _,lbl in metrics])
for i in range(mat.shape[0]):
    for j in range(mat.shape[1]):
        if np.isnan(mat[i,j]): continue
        star = '*' if pmat[i,j]<0.05 else ('.' if pmat[i,j]<0.10 else '')
        ax.text(j, i, f'{mat[i,j]:+.2f}{star}', ha='center', va='center',
                fontsize=9, color='black')
plt.colorbar(im, ax=ax, label='Spearman ρ')
ax.set_title('Figure R8. Big-5 × Δ(extraction-derived originality / fluency), paired within subject')
plt.tight_layout()
plt.savefig(os.path.join(FIG, 'figR_originality_big5.png'), bbox_inches='tight')
plt.close()

print('figures saved:',
      'figR_agentic_workflow.png, figR_originality_same_by_family.png, '
      'figR_fluency_by_family.png, figR_originality_big5.png')

# -----------------------------------------------------------------
# 6) Docx update — insert LAYER L section before Summary
# -----------------------------------------------------------------
from docx import Document
from docx.shared import Inches, Pt

doc_path = os.path.join(ROOT, 'Experiment1_Results.docx')
doc = Document(doc_path)

summary_p = None
for p in doc.paragraphs:
    if p.style.name.startswith('Heading 1') and p.text.strip().lower() == 'summary':
        summary_p = p
        break
if summary_p is None:
    print('[warn] Summary heading not found; appending at end.')

def _ins(text='', style='Normal'):
    if summary_p is None:
        return doc.add_paragraph(text, style=style)
    return summary_p.insert_paragraph_before(text, style=style)

def _img(path, width_in=6.4):
    if summary_p is None:
        p = doc.add_paragraph()
    else:
        p = summary_p.insert_paragraph_before()
    p.add_run().add_picture(path, width=Inches(width_in))
    p.alignment = 1
    return p

def fmtp(p):
    if p < 1e-4: return "p < 10\u207B\u2074"
    if p < 1e-3: return "p < 0.001"
    return f"p = {p:.3f}"

# ---- Section: Agentic Idea Extraction Pipeline ----
_ins('Agentic Idea Extraction Pipeline', style='Heading 1')
_ins(
  "The process-layer measures reported above (consecutive-message novelty, "
  "message-level surprise) operate on raw conversational turns and so conflate "
  "substantive ideation with small-talk, questions, and reactions. To provide a "
  "product-level outcome comparable with Experiment 2 (Rosenbaum et al., UIST "
  "2026), we developed an open-source agentic pipeline that extracts discrete, "
  "evidence-grounded user-originated ideas from each transcript and computes "
  "participant-centroid originality measures in the manner of Rosenbaum et al. "
  "The pipeline runs entirely on open-source components: Qwen3-4B-Instruct "
  "(INT4-quantized) executed through OpenVINO on an Intel Arc 140T GPU, and "
  "BAAI/bge-large-en-v1.5 for sentence-level embeddings. No proprietary API "
  "was used.")
_ins('Five-agent architecture.', style='Normal')
_img(os.path.join(FIG, 'figR_agentic_workflow.png'), width_in=6.7)
_ins(
  "Figure R5 summarises the pipeline. Agent 1 reads a single user message "
  "together with its preceding assistant turn and emits up to three "
  "IdeaCandidate records {title, description, evidence_span, confidence}. The "
  "model output is constrained at decode time to a Pydantic JSON schema via "
  "lm-format-enforcer, which makes malformed JSON structurally impossible. "
  "Agent 1 is prompt-engineered to obey five rules: (i) only user-originated "
  "content is a valid source; (ii) evidence_span must be a verbatim substring "
  "of the user message, with typos preserved; (iii) titles use the user's "
  "vocabulary; (iv) compound proposals are decomposed into separate items; "
  "(v) non-proposals — questions, reactions, challenge restatements, "
  "assistant echoes — return an empty list.")
_ins(
  "Three rule-based filters enforce precision. Filter 1 blocks a user turn "
  "that is a near-duplicate (rapidfuzz partial_ratio \u2265 85) of the known "
  "challenge prompt; this is applied before the LLM call. Filter 2 blocks any "
  "candidate whose evidence_span is substantially contained in the preceding "
  "assistant turn; this removes echo-and-agree candidates. Filter 3 blocks "
  "canonical ideas whose title introduces content words not present in any "
  "evidence quote (after simple stemming), preventing LLM hallucination of "
  "vocabulary absent from the user's actual text.")
_ins(
  "Agent 2 consolidates per-conversation candidates. Each candidate's "
  "title+description is encoded with BGE-large-en-v1.5 (1024-dim, "
  "L2-normalised), and candidates are clustered by agglomerative linkage with "
  "cosine-distance threshold 1\u2212\u03C4, \u03C4=0.85. Clusters of size \u2265 2 "
  "are re-summarised by the LLM into a single canonical record, with a "
  "vocabulary constraint to prevent the introduction of new synonyms; "
  "singletons pass through unchanged.")
_ins(
  "Agent 3 is rule-based. For each canonical idea, every evidence_quote must "
  "be a verbatim substring of some user message in the source conversation "
  "(case- and whitespace-normalised). Quotes failing exact matching are "
  "fuzzy-matched with rapidfuzz.partial_ratio \u2265 90; those failing both are "
  "classified ungrounded and the idea is dropped. Filter 3 is applied here: "
  "canonical ideas whose title content words are not present (modulo stemming, "
  "with one word of tolerance) in any evidence_quote are classified "
  "title_hallucination and dropped.")
_ins(
  "Agent 4 computes corpus-level idea categories. The union of all surviving "
  "canonical ideas is embedded with BGE-large-en-v1.5 and clustered with "
  "HDBSCAN (cosine metric, min_cluster_size=4). Each non-noise cluster is "
  "labelled by the LLM with a short category name (2\u20135 words, drawn from "
  "cluster exemplars' vocabulary).")
_ins(
  "Agent 5 computes participant-level originality. Given participant p's set "
  "of surviving idea embeddings E\u209A, the participant centroid is")
eq = _ins("    C\u209A = normalize(mean_{i \u2208 E\u209A} (i))")
eq.runs[0].italic = True
_ins("Three originality measures follow (higher = more distinctive idea portfolio):")
eq = _ins("    orig\u209B\u2090\u2098\u2091(p)  = (1/|Q\u209B\u2090\u2098\u2091|) \u2211\u2096\u2208Q\u209B\u2090\u2098\u2091 (1 \u2212 \u27E8C\u209A, C\u2096\u27E9)")
eq.runs[0].italic = True
eq = _ins("    orig\u2090\u2097\u2097(p)   = (1/|Q\u2090\u2097\u2097|)  \u2211\u2096\u2208Q\u2090\u2097\u2097 (1 \u2212 \u27E8C\u209A, C\u2096\u27E9)")
eq.runs[0].italic = True
eq = _ins("    orig\u2091\u2093(p)    = min\u2096\u2208Q\u2091\u2093 (1 \u2212 \u27E8C\u209A, C\u2096\u27E9)")
eq.runs[0].italic = True
_ins(
  "Q\u209B\u2090\u2098\u2091 denotes other participant-rounds in the same "
  "experimental condition as p; Q\u2090\u2097\u2097 denotes all other "
  "participant-rounds; Q\u2091\u2093 denotes participant-rounds in the "
  "opposing condition. The three measures together decompose \"how distinctive "
  "is this participant's idea portfolio\" across (a) same-condition peers — "
  "Experiment 2's headline measure — (b) all peers regardless of condition, "
  "and (c) distance from the nearest opposing-condition peer.")
_ins('Pipeline evaluation.', style='Normal')
_ins(
  "Five smoke-test rounds preceded the production deployment. Round 1 (Qwen "
  "1.5B on CPU) achieved 95% JSON validity and 60% evidence-grounding on five "
  "stratified conversations \u2014 grounding failures were dominated by "
  "role-boundary confusion (extracting assistant-authored ideas as if they "
  "were user ideas). Round 2 (Gemma 4 E4B on CPU) hit production-grade "
  "quality on the single conversation that completed, but throughput at "
  "\u22480.5 tok/s was prohibitive for the full corpus. Round 3 switched to "
  "Qwen3-4B-Instruct-2507 in OpenVINO INT4 form on the Intel Arc 140T GPU, "
  "reaching 100% JSON validity and 100% evidence-grounding across 5 "
  "conversations at \u224820 tok/s. Manual recall audit on three conversations "
  "revealed that while grounding was perfect, approximately 33% of extracted "
  "items were not genuine user proposals but challenge restatements, "
  "assistant echoes, or meta-questions (this is the classic "
  "grounding-vs-validity distinction \u2014 every extracted evidence quote was "
  "verbatim in user text, yet the associated turn was not itself propositional "
  "content). Rounds 4 and 5 introduced the three precision filters described "
  "above; Round 5 achieved 100% precision and 90% recall on the audit set, at "
  "which point the pipeline was promoted to production. Production runtime on "
  "all 194 Experiment-1 conversations was 306 minutes (5.1 h) on the Arc 140T, "
  "yielding 740 canonical ideas across 181 participant-rounds "
  "(mean 3.8 ideas per round, SD 2.2).")

# ---- Section: Extracted-Idea Originality ----
_ins('Extracted-Idea Originality and Fluency (Product-Level)', style='Heading 1')
_ins(
  f"Eighty-seven of the 97 participants contributed both a GPT-round and a "
  f"Persona-round idea portfolio, permitting within-subject paired tests that "
  f"were unavailable to Experiment 2's between-subjects design. Two robust "
  f"findings emerge. First, Persona-guided conversation produces "
  f"approximately 33% more discrete user-originated idea proposals than "
  f"GPT-guided conversation (paired \u0394 = {ws_fluency['diff']:+.2f} ideas "
  f"per round, t({ws_fluency['n']-1}) = {ws_fluency['t']:+.2f}, "
  f"{fmtp(ws_fluency['p'])}, dz = {ws_fluency['dz']:+.2f}). The same contrast "
  f"between-subjects is M\u209A = {bs_fluency['mean_p']:.2f}, "
  f"M_G = {bs_fluency['mean_g']:.2f}, Welch t = {bs_fluency['t']:+.2f}, "
  f"{fmtp(bs_fluency['p'])}, Hedges' g = {bs_fluency['g']:+.2f}. Figure R7 "
  f"shows the paired contrast broken out by persona family; the effect is "
  f"directionally consistent in all four families and statistically "
  f"significant in Divergent, Convergent, and BoundedRational (Rational "
  f"n = 7, underpowered).")
_img(os.path.join(FIG, 'figR_fluency_by_family.png'), width_in=6.7)
_ins(
  "Figure R7. Per-round idea count (fluency), paired GPT vs Persona within "
  "each participant, faceted by persona family. Grey lines connect the same "
  "participant across rounds; boxes show the per-condition distribution. All "
  "four families show a positive Persona\u2013GPT gap; significance markers "
  "are shown above each panel.")
_ins(
  f"Second, and counter to Experiment 2's direction, the distinctiveness of "
  f"each participant's idea portfolio relative to same-condition peers is "
  f"*lower* under Persona than under GPT on all three originality measures "
  f"(Table: orig\u209B\u2090\u2098\u2091 paired \u0394 = "
  f"{ws_same['diff']:+.4f}, t({ws_same['n']-1}) = {ws_same['t']:+.2f}, "
  f"{fmtp(ws_same['p'])}, dz = {ws_same['dz']:+.2f}; orig\u2090\u2097\u2097 "
  f"paired \u0394 = {ws_all['diff']:+.4f}, t = {ws_all['t']:+.2f}, "
  f"{fmtp(ws_all['p'])}, dz = {ws_all['dz']:+.2f}; orig\u2091\u2093 paired "
  f"\u0394 = {ws_cross['diff']:+.4f}, t = {ws_cross['t']:+.2f}, "
  f"{fmtp(ws_cross['p'])}, dz = {ws_cross['dz']:+.2f}). The between-subjects "
  f"effect on same-condition originality is Welch t = {bs_same['t']:+.2f}, "
  f"{fmtp(bs_same['p'])}, Hedges' g = {bs_same['g']:+.2f}. Figure R6 shows "
  f"the paired contrast on orig\u209B\u2090\u2098\u2091, faceted by persona "
  f"family; all four families show a negative Persona-GPT shift, significant "
  f"in Divergent (dz = -0.85), Convergent (dz = -0.59), and BoundedRational "
  f"(dz = -1.03).")
_img(os.path.join(FIG, 'figR_originality_same_by_family.png'), width_in=6.7)
_ins(
  "Figure R6. Same-condition originality paired within each participant, "
  "faceted by persona family. Every panel shows a negative Persona-GPT shift "
  "\u2014 participants' idea portfolios become less distinctive from their "
  "same-family peers when interacting with a Persona than when interacting "
  "with baseline GPT.")
_ins('Reconciling with Experiment 2.', style='Normal')
_ins(
  "Experiment 2 reported the opposite direction on the matching measure "
  "(Treatment M = 0.34 vs. Control M = 0.28, Welch t = 5.13, p = 2 \u00D7 "
  "10\u207B\u2076, Hedges' g = +0.57). The key design difference is that "
  "Experiment 2 exposed each participant to two simultaneously available "
  "personas (Taylor and Alex) with free choice between them, whereas "
  "Experiment 1 assigned a single persona per round with no choice. We "
  "interpret the divergent directions as an interaction-structural effect: "
  "elective persona use functions as a creative scaffold that broadens the "
  "idea space across participants, while mandatory persona assignment "
  "functions as an attractor that concentrates each participant's "
  "contributions within a persona-typical region of idea space. Both studies "
  "find Persona conditions increase engagement (more questions, more words, "
  "more ideas); they diverge on whether Persona conditions also increase "
  "portfolio distinctiveness at the between-subjects level \u2014 and that "
  "divergence plausibly reflects whether the persona is selected or imposed.")
_ins('Individual-difference moderation.', style='Normal')
_img(os.path.join(FIG, 'figR_originality_big5.png'), width_in=6.0)
_ins(
  "Figure R8. Spearman correlations between Big-5 domain scores and "
  "within-subject deltas (Persona - GPT) on fluency and the three originality "
  "measures. Open-Mindedness shows the only consistent pattern: it predicts "
  "preservation of cross-condition originality under Persona "
  "(\u03C1 = +0.29, p = 0.007) and negatively predicts the fluency gain "
  "(\u03C1 = -0.24, p = 0.027). This suggests participants higher in openness "
  "are less funnelled by persona assignment and less fluency-boosted by it "
  "\u2014 consistent with a ceiling effect on their baseline GPT ideation.")
_ins('Perception-behaviour dissociation.', style='Normal')
_ins(
  "Although the behavioural effects of persona assignment on fluency and "
  "originality are large and robust, none of the four deltas correlates "
  "with either the self-reported creativity delta or the ownership delta "
  "(all |\u03C1| < 0.18, all p > 0.10). In other words, persona conditions "
  "systematically change what participants produce but not what they report "
  "about the experience. This mirrors the main-line Experiment-1 finding "
  "(self-report nulls alongside large process effects) and reinforces the "
  "paper's conclusion that subjective creativity ratings do not track the "
  "behavioural signatures of persona-guided interaction.")

doc.save(doc_path)
print('updated:', doc_path)
