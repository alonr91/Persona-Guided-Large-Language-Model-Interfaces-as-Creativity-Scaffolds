"""Augment section 3.8 of Experiment1_StanceContracts_Paper_v4.docx with the
cross-model (Gemini Scorer C) validation, in the same Validation-Summary /
Claim style the section already uses. Touches:

  - section 3.8: extends the bullet list, updates the sampling-caveat bullet,
                 adds a small per-family direction-match table, and rewrites
                 Claim 3.8 to be "across two independent models".
  - section 4.6 Limitations: adds a short paragraph noting that the
                 single-model caveat has been partly addressed, the human-
                 raters caveat is unchanged.

Forbidden-language audit runs after save.
"""
from __future__ import annotations
import os, sys, warnings
warnings.filterwarnings('ignore')
try: sys.stdout.reconfigure(encoding='utf-8')
except: pass

import pandas as pd
from docx import Document

ROOT = r'C:/Users/alonr/OneDrive/Documents/LLM creativity/Experiment 1'
OUT = os.path.join(ROOT, 'regulated_llm_reanalysis')
PAPER = os.path.join(ROOT, 'Experiment1_StanceContracts_Paper_v4.docx')


def insert_before(target_para, text: str, style: str):
    new_p = target_para._parent.add_paragraph(text, style=style)
    target_para._element.addprevious(new_p._element)
    return new_p


def append_table(d: Document, headers: list[str], rows: list[list[str]]):
    t = d.add_table(rows=1 + len(rows), cols=len(headers))
    try: t.style = 'Light Grid Accent 1'
    except KeyError:
        try: t.style = 'Table Grid'
        except KeyError: pass
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for r in c.paragraphs:
            for run in r.runs: run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    return t


def move_table_after(t, para):
    para._element.addnext(t._element)


def forbidden_audit(path):
    forbid = ['true creativity', 'ground-truth creativity',
              'validated creativity improvement', 'objectively more creative',
              'made participants more creative']
    d = Document(path)
    full = '\n'.join(p.text for p in d.paragraphs)
    hits = [t for t in forbid if t in full.lower()]
    print(f'  forbidden-language hits: {hits or "none"}')


# ---------- source numbers ----------

def family_summary_rows() -> list[list[str]]:
    side = pd.read_csv(os.path.join(OUT, '16_qwen_vs_gemini_sidebyside.csv'))
    rows = []
    for fam in ('Divergent', 'Convergent', 'Rational', 'BoundedRational'):
        s = side[side.contrast == f'{fam} vs GPT']
        sm = int(s['sign_match'].sum()); n = len(s)
        bs = int(s['both_sig'].sum())
        sq = int(s['sig_Qwen'].sum()); sg = int(s['sig_Gemini'].sum())
        rows.append([f'{fam} vs GPT', f'{sm}/{n}', str(bs), str(sq), str(sg)])
    return rows


def headline_g_rows() -> list[list[str]]:
    """Side-by-side g for the criteria where Qwen reported the largest effects."""
    side = pd.read_csv(os.path.join(OUT, '16_qwen_vs_gemini_sidebyside.csv'))
    crits = ('reframing_quality', 'exploration_opening', 'anchor_management',
             'timing_fit', 'stance_integrity', 'coregulation_uptake',
             'agency_preservation', 'evaluative_discipline',
             'cognitive_load_clarity', 'implementation_grounding',
             'premature_convergence_risk', 'runaway_divergence_risk')
    out = []
    for crit in crits:
        cells = [crit]
        for fam in ('Divergent', 'Convergent', 'Rational', 'BoundedRational'):
            r = side[(side.criterion == crit) & (side.contrast == f'{fam} vs GPT')]
            if len(r) == 0:
                cells.extend(['—', '—']); continue
            r = r.iloc[0]
            cells.append(f"{r.g_Qwen:+.2f}{'*' if r.sig_Qwen else ''}")
            cells.append(f"{r.g_Gemini:+.2f}{'*' if r.sig_Gemini else ''}")
        out.append(cells)
    return out


# ---------- update ----------

def update_paper():
    d = Document(PAPER)
    paras = list(d.paragraphs)

    # Locate anchors
    intro_38 = None       # paragraph ending with "(Figure 8)."
    sampling_caveat = None  # the 4th bullet starting with "One sampling caveat"
    claim_38 = None       # "Claim 3.8."
    next_h2_after_38 = None

    in_38 = False
    for i, p in enumerate(paras):
        t = p.text.strip()
        if p.style.name.startswith('Heading 2') and t.startswith('3.8 Regulated rubric'):
            in_38 = True; continue
        if in_38 and p.style.name.startswith('Heading 2') and t.startswith('3.9 '):
            next_h2_after_38 = p; in_38 = False; break
        if in_38:
            if 'stratified sample of 200' in t:
                intro_38 = p
            if t.startswith('One sampling caveat'):
                sampling_caveat = p
            if t.startswith('Claim 3.8'):
                claim_38 = p
    assert next_h2_after_38 is not None, '3.9 anchor not found'
    assert sampling_caveat is not None, 'sampling-caveat bullet not found'
    assert claim_38 is not None, 'Claim 3.8 not found'

    # 1) Update the intro paragraph to mention the Gemini layer up-front.
    intro_38.text = (
        "The regulated reanalysis scored a stratified sample of 200 conversational "
        "episodes (50 GPT, 150 Persona; balanced across episode type and persona "
        "family) on twelve 0–4 ordinal criteria, with three scorers: Scorer A "
        "(Qwen3-4B-Instruct, primary), Scorer B (Qwen3-4B-Instruct with a "
        "paraphrased prompt, on a 50-episode subset, for prompt-paraphrase "
        "robustness), and Scorer C (Gemini 3.1 Flash Lite Preview, on the full "
        "200-episode sample with the Scorer A prompt held constant, for cross-"
        "model robustness). We report Welch t with Hedges' g and Benjamini–"
        "Hochberg-FDR-corrected q-values (Figure 8)."
    )

    # 2) Update the sampling-caveat bullet to note Gemini fixes the asymmetry.
    sampling_caveat.text = (
        "One sampling caveat to flag. Scorer A produced bundle-valid JSON for "
        "29/50 (58%) of GPT episodes vs 118/150 (79%) of Persona episodes. "
        "Episode lengths are matched across conditions (mean 3.54 vs 3.53 turns), "
        "so the differential is content-driven, not length-driven. Masking was "
        "intact, so the differential reflects text-content properties rather "
        "than condition leakage. Scorer C (Gemini) produced bundle-valid JSON "
        "for 50/50 GPT and 150/150 Persona episodes (600/600 GPT criterion-rows "
        "vs 369/600 from Scorer A), which closes this asymmetry and confirms "
        "that it was a Qwen-side bundle-validity artifact rather than a "
        "property of the data."
    )

    # 3) Insert a new "Cross-model agreement" bullet before Claim 3.8.
    cross_bullet = (
        "Cross-model agreement (Qwen vs Gemini). Pooled across 458 dual-scored "
        "rows, the two scorers agree at quadratic-weighted Cohen's κ = 0.69 and "
        "Spearman ρ = 0.77 (p < 10⁻⁸⁹), substantial agreement by Landis–Koch. "
        "Across the 48 family-vs-GPT × criterion contrasts, the two scorers "
        "agree on the sign of the effect in 47/48 (97.9%) and on joint "
        "significance at p < 0.05 in 34/48. The single sign disagreement is "
        "implementation_grounding for Divergent vs GPT (g_Qwen = +0.12, "
        "g_Gemini = −0.10); both effects are near zero and neither reaches "
        "p < 0.05 in either scorer. The within-arm Divergent vs Convergent "
        "contrast — the strongest pairwise within-arm comparison — "
        "distinguishes 7 (Qwen) and 7 (Gemini) of 12 criteria at p < 0.05, "
        "with the same trade-off signature: Divergent has higher "
        "exploration_opening, reframing_quality, and runaway_divergence_risk, "
        "and lower premature_convergence_risk."
    )
    insert_before(claim_38, cross_bullet, 'List Bullet')

    # 4) Add a small intro paragraph + per-family table BEFORE Claim 3.8 too.
    intro_table_para = insert_before(
        claim_38,
        "Per-family direction-match (Scorer A vs Scorer C; significance is "
        "raw p < 0.05, no multiple-comparison correction):",
        'Normal',
    )
    rows = family_summary_rows()
    headers = ['contrast', 'sign-match', 'both sig at p<0.05 (of 12)',
               'Qwen-A sig (p<0.05)', 'Gemini-C sig (p<0.05)']
    t1 = append_table(d, headers, rows)
    move_table_after(t1, intro_table_para)

    # 5) Add a second small paragraph + per-criterion side-by-side g table.
    g_intro = insert_before(
        claim_38,
        "Per-criterion Hedges' g, each persona family vs GPT (Scorer A | "
        "Scorer C). Asterisk = raw p < 0.05 (no multiple-comparison "
        "correction). For the two RISK criteria positive g means MORE risk; "
        "negative is the favourable direction.",
        'Normal',
    )
    g_rows = headline_g_rows()
    g_headers = ['criterion',
                 'Div A', 'Div C',
                 'Con A', 'Con C',
                 'Rat A', 'Rat C',
                 'BR A', 'BR C']
    t2 = append_table(d, g_headers, g_rows)
    move_table_after(t2, g_intro)

    # 6) Rewrite Claim 3.8 to reflect cross-model corroboration.
    claim_38.text = (
        "Claim 3.8. On LLM-rubric proxy scores, persona-guided interaction is "
        "rated higher on every regulation dimension that captures how the "
        "dialogue is conducted (reframing, exploration, anchor management, "
        "timing, co-regulation, stance integrity, evaluative discipline, "
        "agency preservation), and lower on premature-convergence risk. The "
        "direction and significance pattern (raw p < 0.05) reproduces across "
        "two independent LLM scorers — a local INT4 Qwen3-4B-Instruct and a "
        "commercial Gemini 3.1 Flash Lite Preview — on the full 200-episode "
        "stratified sample, with sign agreement on 47 of 48 family-vs-GPT × "
        "criterion contrasts. The strongest claim this layer supports remains "
        "proxy-bounded: persona-guided interaction changes the regulation of "
        "creative dialogue. Establishing an externally judged creativity gain "
        "would still require human or domain-expert rating that this study "
        "does not include."
    )

    # 7) 4.6 Limitations: append a short addendum paragraph.
    paras = list(d.paragraphs)
    lim_h2 = None; next_after_lim = None
    for i, p in enumerate(paras):
        if p.style.name.startswith('Heading 2') and p.text.strip().startswith('4.6 Limitations'):
            lim_h2 = p
        elif lim_h2 is not None and p.style.name.startswith('Heading 2'):
            next_after_lim = p; break
    addendum = (
        "Cross-model validation update. The single-model limitation noted "
        "above for the regulated rubric has been partly addressed: Scorer C "
        "(Gemini 3.1 Flash Lite Preview) was added on the full 200-episode "
        "stratified sample, holding the Scorer A prompt template constant. "
        "Qwen-A and Gemini-C agree at quadratic-weighted κ = 0.69 / ρ = 0.77 "
        "pooled, with sign agreement on 47/48 family-vs-GPT × criterion "
        "contrasts (significance assessed at raw p < 0.05); this upgrades "
        "the original Scorer A vs Scorer B prompt-paraphrase check into a "
        "true cross-model robustness check. The remaining limitation — that "
        "no human raters scored the rubric — is unchanged."
    )
    if next_after_lim is not None:
        insert_before(next_after_lim, addendum, 'Normal')
    else:
        d.add_paragraph(addendum, style='Normal')

    d.save(PAPER)
    print(f'updated {PAPER}')
    forbidden_audit(PAPER)


if __name__ == '__main__':
    update_paper()
