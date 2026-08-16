# -*- coding: utf-8 -*-
"""Builds Experiment2_in_exp2_style.docx mirroring the Experiment 1 chapter structure.
All statistics are the values recomputed in _reconcile_stats.py / _reconcile_process.py."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

MEDIA = "media"
doc = Document()

# ---- base style ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def h3(t): doc.add_heading(t, level=3)

def p(text, italic=False, bold=False, size=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = italic; run.bold = bold
    if size: run.font.size = Pt(size)
    return par

def bullet(text):
    par = doc.add_paragraph(style='List Bullet')
    par.add_run(text)
    return par

def numbered(text):
    par = doc.add_paragraph(style='List Number')
    par.add_run(text)
    return par

def caption(text):
    par = doc.add_paragraph()
    run = par.add_run(text); run.italic = True; run.font.size = Pt(9)
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return par

def add_figure(fname, cap, width=6.0):
    path = os.path.join(MEDIA, fname)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(cap)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].paragraphs[0].add_run(htext).bold = True
        for r in hdr[i].paragraphs[0].runs: r.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return t

def mono_block(text):
    for line in text.split('\n'):
        par = doc.add_paragraph()
        run = par.add_run(line if line else " ")
        run.font.name = 'Consolas'; run.font.size = Pt(8)
        par.paragraph_format.space_after = Pt(0)

# =====================================================================
# TITLE + ABSTRACT
# =====================================================================
doc.add_heading("Experiment 2: How User-Switchable Divergent and Convergent LLM Personas Shape Human–Machine Creative Problem-Solving", level=0)

h1("Abstract")
p("This experiment tests whether giving users simultaneous, on-demand access to two contrasting "
  "LLM personas—one divergent (“Taylor”) and one convergent (“Alex”)—within a single chat "
  "interface changes the process and products of creative problem-solving relative to a standard "
  "single-model control. Where Experiment 1 assigned persona stances to users and measured uptake, "
  "Experiment 2 makes the divergent/convergent contrast user-elective: participants choose which "
  "persona to address on each turn through two color-coded send buttons. In a between-subjects design "
  "(originality sample N=101, 70 treatment / 31 control; questionnaire sample N=95, 66 / 29), "
  "undergraduates solved a single open-ended ideation problem (“How can we make libraries more "
  "attractive to young adults?”) in a 20-minute session. Treatment participants produced "
  "semantically more original idea portfolios across three embedding-based cosine-distance measures "
  "(same-condition originality g=+0.98, p=2.2×10⁻⁶), with no increase in idea quantity "
  "(fluency p=.35, n.s.) and no increase in within-participant idea spread (p=.90)—the advantage is "
  "between-user differentiation, not broader individual portfolios. Participants attributed creative "
  "support more strongly to the divergent persona (creativity-enhancement g=+0.69, p=.003) and "
  "reported reduced sole ownership of their solutions (g=−0.52, p=.018). Interaction traces revealed "
  "an emergent divergence-to-convergence progression in question-asking, and Big Five traits moderated "
  "persona engagement (agreeableness and openness with divergent benefit; conscientiousness with "
  "convergent engagement). These effects contrast with Experiment 1’s fluency-up/distinctiveness-down "
  "pattern under assigned personas and indicate that user-orchestrated mode access amplifies "
  "between-user originality while redistributing perceived authorship. Claims are bounded to "
  "embedding-based originality and an expert-validated manipulation check, not externally judged "
  "creativity.")

# =====================================================================
# 0. RATIONALE
# =====================================================================
h1("0. Experiment rationale")
p("Experiment 1 established that assigned persona stance contracts can reorganise the trajectory of a "
  "human–LLM creative dialogue: users took up broadening (expansion, reframing, proposing) more readily "
  "than closure, and produced more ideas (higher fluency) whose portfolios were nonetheless less "
  "distinctive (an “ideational attractor” / homogenisation pattern). Crucially, in Experiment 1 the "
  "persona was imposed on the user rather than chosen, and the four persona families empirically "
  "collapsed toward a two-mode contrast (rational ≈ convergent, bounded-rational ≈ divergent). "
  "Experiment 1 therefore left open the forecasted next question: what happens when the divergent and "
  "convergent modes are sharpened into two distinct agents and handed to the user as a switch they "
  "control on demand?")
p("Experiment 2 tests exactly this. It externalises divergent and convergent thinking as two "
  "separately addressable personas that are simultaneously visible in one interface, and treats "
  "persona selection as a first-class, user-initiated interaction act. The design deliberately does "
  "not test user-controlled switching against assigned switching in the same study; rather, it asks "
  "whether providing structural access to two contrasting modes (versus a single standard LLM) changes "
  "creative process and product. Consistent with the Experiment 1 prediction that refined two-mode "
  "prompts would behave differently, the question-asking signal here runs in the opposite direction to "
  "Experiment 1: under elective access, treatment users ask more questions than control users, and the "
  "between-user originality effect reverses Experiment 1’s distinctiveness loss. This experiment thus "
  "probes the second empirical link in the persona-scaffolding account—whether user-orchestrated mode "
  "access can preserve, rather than erode, the originality of what people make.")

# =====================================================================
# 1. METHODOLOGY
# =====================================================================
h1("1. Experiment methodology")

h2("1.1 Design and research logic")
p("Experiment 2 used a between-subjects randomised design with two conditions. In the treatment "
  "condition, participants interacted with two differentiated LLM personas—Taylor (divergent) and "
  "Alex (convergent)—accessible through two color-coded send buttons beneath a single shared chat "
  "thread. In the control condition, both buttons routed to the same standard GPT-4.1 configuration "
  "with a minimal task prompt and no role differentiation, holding the two-button interface constant "
  "while removing the persona contrast. Participants were randomly assigned with an asymmetric 2:1 "
  "allocation favouring treatment, which increases power for the within-treatment analyses (persona "
  "preference, trait moderation, mode switching) that are central to the research questions; "
  "between-condition comparisons use Welch’s t-tests and Hedges’ g, both robust to unequal group "
  "sizes.")
p("Persona role labels were withheld from participants: the buttons were named only “Taylor” and "
  "“Alex,” with button colour and left/right placement randomised across participants, and the "
  "underlying divergent/convergent distinction left undisclosed. This blinding choice serves two "
  "purposes. First, it allows the study to observe emergent role discovery without instructional "
  "scaffolding, providing a conservative lower bound on role-differentiation effects. Second, it "
  "removes demand characteristics that would arise if users were told which agent was “the creative "
  "one.” This creates a deliberate tension with the chapter’s eventual design recommendation that "
  "cognitive roles be communicated explicitly; that trade-off is examined in the Discussion.")
add_table(
    ["Design component", "Implementation", "Why it matters"],
    [["Condition assignment", "Between-subjects, randomised, 2:1 (treatment:control)",
      "Isolates the persona-contrast manipulation; 2:1 powers within-treatment analyses"],
     ["Mode access", "Two simultaneously visible, user-selectable personas (treatment) vs. one standard model behind both buttons (control)",
      "Makes divergent/convergent mode access user-elective rather than assigned (cf. Experiment 1)"],
     ["Blinding", "Persona role labels withheld; button colour/position randomised",
      "Tests emergent role discovery; removes demand characteristics"],
     ["Task", "Single open-ended ideation problem, 20-minute session",
      "Holds problem and time constant across conditions"],
     ["Backend", "GPT-4.1 for all conditions; persona behaviour induced by system prompts + temperature",
      "Differences are attributable to the prompt/temperature envelope, not the base model"]],
)
caption("Table 1. Core experimental structure.")

h3("1.1.1 Research questions, contributions, and manipulation check")
p("Three research questions, retained from the study’s pre-registered framing, structure the analysis:")
numbered("RQ1: How do heterogeneous (dual-persona) LLM interactions, compared with a standard LLM "
         "interface, affect users’ perceived creativity enhancement and sense of ownership of ideas?")
numbered("RQ2: To what extent do individual differences in personality traits (openness, agreeableness, "
         "conscientiousness, extraversion, neuroticism) shape engagement patterns and preferences toward "
         "the divergent versus convergent persona?")
numbered("RQ3: How do dual-persona interactions influence the originality and diversity of idea "
         "portfolios, measured through embedding-based semantic analysis?")
p("The work makes three contributions: (1) a role-structured interaction paradigm in which divergent "
  "and convergent modes are embodied as distinct, user-accessible personas, with persona transitions "
  "treated as first-class interaction acts; (2) empirical evidence on how users navigate between these "
  "roles and how this relates to creative outcomes and perceived authorship; and (3) an examination of "
  "how personality traits relate to persona preference and interaction patterns. Before any "
  "outcome analysis, a manipulation check (Section 3.1) verifies, using blind expert ratings, that the "
  "two treatment personas actually behaved divergently versus convergently and that the control "
  "agents did not—gating the remainder of the chapter.")

h2("1.2 Creative task")
p("All participants addressed the same open-ended challenge: “How can we make libraries more "
  "attractive to young adults?” This is a deliberately ill-structured, socially situated “wicked” "
  "design problem with no single correct answer, no required external data, and ample room for both "
  "divergent reframing and convergent prioritisation. Each participant had 20 minutes after providing "
  "digital consent and reading a short, standardised orientation that introduced the interface. "
  "Participants could compose in English or in their native language; all non-English text was "
  "translated to English for analysis.")

h2("1.3 Persona manipulation")
p("The two treatment personas were instantiated through differentiated system prompts and decoding "
  "temperatures applied to a common GPT-4.1 backend. Taylor (divergent) was prompted to generate "
  "multiple varied alternatives without premature evaluation, draw on analogical and cross-domain "
  "reasoning, invite the user to extend or recombine possibilities, and resist narrowing the search "
  "space; it was decoded at temperature 0.8 to encourage variety. Alex (convergent) was prompted to "
  "evaluate feasibility and distinctiveness, help articulate selection criteria and trade-offs, "
  "structure material toward actionable conclusions, and focus on the most promising directions; it "
  "was decoded at temperature 0.3 for precision and stability. Both personas shared identical task "
  "framing and response rules (conciseness, no filler openers, plain language for non-native speakers, "
  "no markdown emphasis, role persistence, and language matching). Prompts were developed iteratively "
  "in pilot sessions with design students (not part of the main sample) and refined until the "
  "behavioural contrast was reliably perceptible. The full system prompts are reproduced verbatim in "
  "Appendix A.")
p("Context across turns was maintained without a stateful model session: for each request the backend "
  "retrieved the recent transcript, generated a short JSON state summary (solution-so-far, decisions, "
  "open threads), and appended a persona-specific reinforcement reminder, so persona behaviour "
  "persisted across long conversations.")
add_table(
    ["Persona / condition", "Intended mode", "Prompt envelope", "Temp."],
    [["Taylor (treatment)", "Divergent: expand, associate, defer evaluation",
      "Creativity/open-mindedness traits; multidisciplinary, analogical guidance; invite recombination", "0.8"],
     ["Alex (treatment)", "Convergent: evaluate, prioritise, structure toward decisions",
      "Accuracy/critical-evaluation traits; systematic, evidence-based guidance; converge on best options", "0.3"],
     ["Control (both buttons)", "Undifferentiated standard assistant",
      "Shared task framing + response rules only; no role traits or thinking guidelines", "0.3"]],
)
caption("Table 2. Persona conditions and intended cognitive modes (GPT-4.1 backend throughout).")

h2("1.4 Materials, interface, logging, and questionnaires")
p("The interface was a single-page web application replicating a standard LLM chat environment. "
  "Messages appeared in a scrollable thread headed by the responding agent’s name; two color-coded "
  "send buttons (Taylor and Alex) below the input field let participants direct each message to either "
  "agent, with colour and placement randomised. System telemetry logged conversation content, message "
  "timestamps, the persona addressed on each turn, persona-switching patterns, and total interaction "
  "length.")
add_figure("ui.png", "Figure 1. The experimental chat interface, with two color-coded send "
           "buttons routing each message to the Taylor or Alex persona.", width=5.5)
p("Upon completion, participants completed a post-session questionnaire comprising the Big Five "
  "Inventory-2-XS (BFI-2-XS), demographics, and an eight-item attitude scale (Table 3). Items 1–7 used "
  "a five-point Likert scale (1 = Strongly Disagree to 5 = Strongly Agree); item 8 used a four-point "
  "forced-choice scale anchored at 1 (Taylor / divergent) and 4 (Alex / convergent).")
add_table(
    ["#", "Item"],
    [["1", "Alex helped me arrive at a creative solution"],
     ["2", "Taylor helped me arrive at a creative solution"],
     ["3", "My creativity increased compared to my usual level after chatting with Alex"],
     ["4", "My creativity increased compared to my usual level after chatting with Taylor"],
     ["5", "The solution originated from me"],
     ["6", "Compared to a regular LLM, the interface helped me reach a creative solution"],
     ["7", "Perceived proficiency with generative AI tools"],
     ["8", "To what degree did each persona enhance your creativity? (1 = primarily Taylor / divergent … 4 = primarily Alex / convergent)"]],
)
caption("Table 3. Perceived attitude toward dialogue questionnaire.")

h2("1.5 Participants, samples, and preprocessing")
p("Participants were undergraduate students at a design-and-engineering institution (mean age ≈ 27.6 "
  "years; roughly 50% design, 30% engineering, 16% other), each paid $15, with a $150 voucher awarded "
  "for the highest-rated submission. Two analytic samples are used and reported separately, because "
  "they correspond to different data sources rather than to a single global N:")
bullet("Originality / behavioural sample (N=101; 70 treatment, 31 control): all complete conversation "
       "logs that passed quality screening. This is the sample for the computational originality (RQ3) "
       "and interaction-log (process) analyses.")
bullet("Questionnaire sample (N=95; 66 treatment, 29 control): participants who additionally completed "
       "valid post-session questionnaires. This is the sample for the subjective (RQ1) and "
       "personality (RQ2) analyses.")
p("Participants were excluded for sending fewer than three messages, session duration under 2 minutes "
  "or over 90 minutes, inactivity timeout, or non-original content. The conditions did not differ in "
  "age (p=.57), academic discipline, or self-reported AI proficiency (treatment M=3.71, SD=1.00; "
  "control M=3.62, SD=0.98; g=0.09, p=.68). Non-English conversations were machine-translated to "
  "English for analysis; as in Experiment 1, translation was applied uniformly across conditions and "
  "is therefore a shared preprocessing step rather than a between-condition confound. The protocol "
  "received ethics approval from the first author’s departmental research committee.")

# =====================================================================
# 2. ANALYSIS STRATEGY
# =====================================================================
h1("2. Analysis strategy")
p("The analysis is organised into layers that mirror the multi-level evaluation logic of Experiment 1. "
  "Like Experiment 1, Experiment 2 includes a regulated LLM-judge proxy layer (Section 2.5); it "
  "additionally adds an expert-rated manipulation-check layer and a dedicated personality-moderation "
  "layer. Table 4 (Section 3.1) reports the manipulation check; the remaining layers are described "
  "below.")

h2("2.1 Subjective layer")
p("The subjective layer analyses the eight-item attitude questionnaire (questionnaire sample, N=95). "
  "Between-condition contrasts use Welch’s t-tests with Hedges’ g and report M, SD per condition and "
  "the signed difference (treatment minus control). For the four-point forced-choice persona item, "
  "one-sample t-tests against the scale midpoint (2) test for an absolute directional "
  "preference within each condition, and the response distribution is examined for polarisation. "
  "Within-participant selectivity is computed as the Taylor-minus-Alex difference on the paired "
  "creativity-enhancement and helpfulness items.")

h2("2.2 Process layer")
p("The process layer uses interaction logs (originality/behavioural sample). It quantifies which "
  "persona users addressed, message counts and longest same-persona streaks, the persona with which "
  "users ended the session, and question-asking behaviour. Question-asking is indexed by the count of "
  "question marks per user message and analysed by conversation quarter (Q1–Q4); the first quarter is "
  "treated as a familiarisation phase. Cross-persona brokerage moves—where users explicitly ask one "
  "persona to react to the other—are reported qualitatively as illustrative of orchestration "
  "behaviour.")

h2("2.3 Product layer")
p("The product layer assesses idea portfolios with a two-stage extraction pipeline run on "
  "user-authored text only (Appendix B). An idea was operationalised as a coherent, self-contained "
  "proposal naming a concept, a target, and at least one concrete affordance; two co-authors achieved "
  "satisfactory agreement on this criterion (Cohen’s κ=0.81). Stage 1 (GPT-4.1, deterministic "
  "decoding) extracts structured ideas; Stage 2 induces ≤8 categories per participant. Fluency is the "
  "count of distinct ideas. Each idea is embedded with text-embedding-3-large, L2-normalised, and "
  "mean-pooled into a participant centroid. Three originality measures are computed as cosine "
  "distances: (1) mean distance to same-condition peers, (2) mean distance to all participants, and "
  "(3) nearest cross-condition neighbour distance; within-participant diversity is the mean pairwise "
  "distance among a participant’s own ideas (Figure 2).")
add_figure("product_layer_pipeline.png", "Figure 2. The product-layer pipeline: two-stage idea "
           "extraction from user-authored text yields distinct ideas (fluency) and categories, while "
           "embedding and centroid pooling yield the three between-user originality measures and "
           "within-participant diversity.", width=5.9)

h2("2.4 Personality-moderation layer")
p("The personality layer relates BFI-2-XS trait scores to subjective outcomes and engagement, "
  "computed within condition (questionnaire sample for attitude outcomes; behavioural sample merged "
  "with traits for log-based engagement). Trait–outcome associations are Pearson correlations; "
  "ending-persona contrasts use trait quartiles with chi-square tests. Following Experiment 1, these "
  "associations are framed as exploratory and reported without multiplicity correction; control-"
  "condition associations are reported alongside treatment to gauge specificity.")

h2("2.5 Regulated LLM-judge proxy layer")
p("As an independent check on the embedding-based originality results, a regulated LLM-as-judge "
  "pipeline scored each participant’s full user-only transcript holistically (Figure 8; Appendix C). "
  "The pipeline judges the whole conversation rather than isolated ideas, runs an ensemble of five "
  "expert judge personas (Design Thinking, Social Psychology, HCI, Philosopher/Ethicist, and "
  "Innovation/Strategy) whose scores are combined by the median to blunt single-lens bias, and rates "
  "eight 1–7 dimensions (holistic creativity, originality, value/usefulness, insight/reframing, "
  "development, integration, human/ethical, process/evidence), each accompanied by a brief auditable "
  "rationale with no chain-of-thought. Scoring is deterministic (temperature 0, fixed seed, JSON-only, "
  "on-disk caching). Idea-count measures (fluency, flexibility) are computed separately and the "
  "holistic ensemble is treated as the headline so that verbosity is not rewarded. A pairwise "
  "Bradley–Terry tournament provides a global ranking as a cross-check, and a battery of validity "
  "audits (inter-judge reliability, length-bias, attribution, and pairwise alignment) accompanies the "
  "scores. As in Experiment 1, this layer is explicitly a bounded proxy for, not a substitute for, "
  "human creativity judgment.")

h2("2.6 Statistical reporting and claim boundaries")
p("Between-condition effects are reported as M, SD per condition; signed difference Δ (treatment minus "
  "control); Welch’s t; p; and Hedges’ g using the pooled standard deviation. Non-significant results "
  "are stated explicitly (p value, n.s.). Correlational and small-cell results are flagged as "
  "exploratory. Two reporting notes follow from reconciling the analyses directly against the raw "
  "data. First, all effect sizes here use the pooled-SD Hedges’ g; for the originality measures, where "
  "the two conditions have unequal variances, this yields slightly more conservative values than a "
  "Welch-t-based d estimator (the test statistics and p-values are identical either way). Second, the "
  "forced-choice persona item is tested against the scale midpoint of 2 (the 1–4 forced-choice scale). The central claim "
  "boundary is the same as in Experiment 1: outcome quality is established through embedding-based "
  "semantic originality, an expert-validated manipulation check, and a regulated multi-judge LLM proxy "
  "(Section 2.5)—not through externally judged human creativity ratings, and the LLM-judge layer is "
  "itself a bounded proxy.")

# =====================================================================
# 3. RESULTS
# =====================================================================
h1("3. Results")

h2("3.1 Manipulation check: the two personas behaved divergently versus convergently")
p("Three domain experts, blind to condition, rated a sample of agent responses on divergent and "
  "convergent scales; inter-rater reliability was acceptable (ICC = 0.65 for the divergent scale, "
  "0.55 for the convergent scale). In the treatment condition the two personas were clearly "
  "separated—Taylor was rated substantially more divergent than Alex (Taylor M=5.07 vs. Alex M=2.67, "
  "p=.032)—whereas in the control condition the two buttons did not differ (Taylor M=3.56 vs. Alex "
  "M=3.00, p=.63, n.s.), confirming both that the manipulation installed the intended contrast and "
  "that blinding held in the control. This check gates the remaining results.")
add_table(
    ["Condition", "Divergent persona (Taylor)", "Convergent persona (Alex)", "Test"],
    [["Treatment", "M = 5.07", "M = 2.67", "p = .032"],
     ["Control", "M = 3.56", "M = 3.00", "p = .63 (n.s.)"]],
)
caption("Table 4. Manipulation check: blind expert ratings of persona behaviour (higher = more "
        "divergent). ICC = 0.65 (divergent scale), 0.55 (convergent scale).")

h2("3.2 RQ1: perceived creativity support and authorship")
p("Figure 3 summarises all questionnaire outcomes for both conditions—the per-persona interface and "
  "creativity-enhancement ratings, solution ownership, the interface-versus-standard-LLM comparison, "
  "the forced-choice persona item, and the two within-participant divergent–convergent deltas. The "
  "individual contrasts are detailed below.")
add_figure("quest_exp_control.png", "Figure 3. Questionnaire outcomes: treatment versus control "
           "(means ± 95% CI) across the RQ1 items and the within-participant divergent–convergent "
           "creativity and helpfulness deltas.", width=6.3)
p("On the four-point forced-choice persona item (1 = divergent, 4 = convergent), treatment "
  "participants leaned toward the divergent persona far more than controls (treatment M=1.80, SD=1.01; "
  "control M=2.72, SD=1.10; Δ=−0.92, t=−3.85, p=.0003, g=−0.88). The distributions were polarised "
  "toward opposite poles: in the treatment condition 52% chose the extreme divergent option (rated 1) "
  "and a further 27% rated 2, placing 79% at or below the scale midpoint, while in the control "
  "condition responses spread toward the convergent end (28% rated 3, 31% rated 4). Tested against the "
  "scale midpoint of 2, the control condition showed a significant absolute preference for the "
  "convergent persona (M=2.72, t(28)=3.55, p=.001), whereas the treatment condition did not differ "
  "significantly from the midpoint despite leaning divergent (M=1.80, t(65)=−1.58, p=.118)—indicating "
  "a relative rather than an absolute divergent preference.")
add_figure("quest_whoHelped.png", "Figure 4. Distribution of persona-preference ratings (Item 8). "
           "Treatment responses concentrate on the divergent (Taylor) side; control responses lean "
           "convergent (Alex).", width=5.5)
p("Helpfulness and creativity-enhancement items localised the effect to the divergent persona. Direct "
  "helpfulness of the convergent persona did not differ between conditions (Δ=−0.04, p=.88, n.s.), and "
  "helpfulness of the divergent persona was numerically higher in treatment but not significant "
  "(treatment M=3.89, control M=3.41; Δ=+0.48, t=1.95, p=.056, g=+0.42). For creativity enhancement, "
  "the divergent persona scored significantly higher in treatment (M=3.65, SD=1.26 vs. control "
  "M=2.76, SD=1.33; Δ=+0.89, t=3.07, p=.003, g=+0.69), while the convergent persona showed no "
  "between-condition difference (M=2.70 vs. 2.76; p=.84, n.s.). The two complementary interface "
  "items behaved similarly: the convergent persona’s interface rating was essentially equal across "
  "conditions (treatment M=3.20, control M=3.24; p=.88, n.s.). Within-participant selectivity confirmed "
  "the asymmetry: the Taylor-minus-Alex creativity-enhancement difference was large and positive in "
  "treatment but essentially zero in control (treatment M=+0.95 vs. control M=0.00; t=2.73, p=.008, "
  "g=+0.58), and the parallel helpfulness delta (divergent-minus-convergent interface rating) was also "
  "larger in treatment though not significantly so (treatment M=+0.70 vs. control M=+0.17; t=1.44, "
  "p=.16, g=+0.29).")
p("The persona interface also reduced perceived sole ownership of the solution. Treatment participants "
  "were less likely to agree that “the solution originated from me” (treatment M=3.56, SD=0.98; "
  "control M=4.07, SD=0.92; Δ=−0.51, t=−2.43, p=.018, g=−0.52), suggesting that persona "
  "differentiation redistributes authorship attribution from the individual toward a collaborative "
  "process. The composite item comparing the interface to a regular LLM (Item 6) was numerically "
  "higher in treatment but not significant (M=3.12 vs. 2.72; Δ=+0.40, p=.12, n.s.).")

h2("3.3 RQ2: personality trait associations")
p("Within the treatment condition, Big Five traits were associated with persona evaluations and "
  "engagement (Table 5). Agreeableness showed the strongest pattern, correlating with divergent "
  "persona helpfulness, divergent creativity enhancement, and perceived superiority over a standard "
  "LLM. Openness correlated with perceived interface superiority and with retained solution ownership. "
  "Extraversion was associated with convergent persona preference. Conscientiousness was associated "
  "with greater convergent engagement in the interaction logs (more convergent messages and longer "
  "convergent streaks). These patterns are consistent with a dual-pathway account in which agreeable "
  "and open individuals are more responsive to divergent support while conscientious individuals "
  "orient toward structured convergence.")
add_table(
    ["Trait", "Outcome", "r", "p"],
    [["Agreeableness", "Divergent persona helpfulness", "+0.27", ".027"],
     ["Agreeableness", "Divergent creativity enhancement", "+0.40", "<.001"],
     ["Agreeableness", "Interface superiority vs. standard LLM", "+0.49", "<.001"],
     ["Extraversion", "Convergent persona preference", "+0.26", ".035"],
     ["Openness", "Interface superiority vs. standard LLM", "+0.37", ".002"],
     ["Openness", "Solution ownership", "+0.28", ".021"],
     ["Conscientiousness", "Convergent message count", "+0.34", ".005"],
     ["Conscientiousness", "Convergent streak length", "+0.38", ".002"]],
)
caption("Table 5. Significant personality trait–outcome associations in the treatment condition "
        "(exploratory; Pearson r). Attitude outcomes N=66; log-based engagement N=65.")
add_figure("quest_big5.png", "Figure 5. Personality trait–outcome correlations in the treatment "
           "condition.", width=5.5)
p("These associations were specific to the structured-persona condition. In the control condition the "
  "corresponding divergent-benefit associations were absent; a handful of control correlations reached "
  "nominal significance (e.g., neuroticism and conscientiousness with convergent appraisal, all "
  "n=29), but these are fragile given the small control sample and are not robust to the exploratory "
  "framing. The trait associations should be read as hypothesis-generating and in need of replication.")

h2("3.4 Process: mode switching, engagement, and the choreography of inquiry")
p("Behavioural logs showed that the large majority of treatment participants addressed both personas "
  "at some point, confirming that the two-button interface elicited active role exploration rather "
  "than collapsing onto a single agent. The persona condition did not simply increase overall talk: "
  "mid-to-late message counts, streak lengths, and switch frequency did not differ reliably between "
  "conditions. Instead, engagement was structured. Conscientiousness predicted convergent engagement "
  "(message count r=0.34, p=.005; streak length r=0.38, p=.002), with no such association in the "
  "control condition.")
p("Trait differences also shaped how sessions ended. Splitting the treatment sample by "
  "conscientiousness quartiles, low-conscientiousness users overwhelmingly ended with the divergent "
  "persona (17 of 20 ended divergent) whereas high-conscientiousness users more often ended with the "
  "convergent persona (15 of 24 ended convergent; χ²(1)=10.18, p=.001). No comparable ending "
  "preference appeared in the control condition (all p>.28).")
p("Question-asking served as a behavioural index of engaged inquiry. Across quarters Q2–Q4 (excluding "
  "the Q1 familiarisation phase), treatment participants asked more questions than controls of both "
  "personas—of the divergent persona (treatment M=0.32 vs. control M=0.14 question marks per message, "
  "t=2.54, p=.013) and of the convergent persona (M=0.28 vs. 0.12, t=2.10, p=.039). The temporal "
  "signature traced a divergence-to-convergence arc: questioning of the divergent persona peaked early "
  "(Q2 ≈ 0.37) and tapered, while questioning of the convergent persona rose to peak later (Q3 ≈ "
  "0.38); control questioning stayed flat and low throughout. This staged pattern—broad exploratory "
  "inquiry first, focused evaluative inquiry later—mirrors the divergence-then-convergence cadence of "
  "the Double Diamond. Notably, this is the opposite direction to Experiment 1, where the assigned-"
  "persona package reduced user questioning; under elective access with refined prompts, the personas "
  "instead invited more inquiry.")
add_figure("question_marks.png", "Figure 6. Question-mark frequency by conversation quarter and "
           "persona. Treatment users question the divergent persona most in Q2 and the convergent "
           "persona most in Q3; control questioning is flat.", width=5.5)
p("Qualitatively, treatment users frequently treated the system as a small creative team rather than a "
  "single tool, brokering between the two personas. Representative moves include cross-persona "
  "referencing—“Alex, what do you think about what Taylor and I did?”—and empathy-mode ideation—"
  "“What do you think of him as the 17-year-old we described earlier?” Such orchestration moves, in "
  "which the user directs specialists and asks them to react to one another, are rare with a single "
  "generic LLM and signal active user ownership of the process.")

h2("3.5 RQ3: computational originality of idea portfolios")
p("Fluency did not differ between conditions (treatment M=8.51, SD=3.85; control M=9.23, SD=3.37; "
  "Δ=−0.71, p=.35, n.s.), establishing that any originality advantage cannot be explained by simply "
  "generating more ideas. On all three originality measures, treatment participants produced more "
  "original portfolios (Figure 7). Same-condition originality showed the largest effect (treatment "
  "M=0.341, control M=0.282; Δ=+0.059, t=5.13, p=2.2×10⁻⁶, g=+0.98). All-participant originality "
  "remained significant (M=0.334 vs. 0.307; Δ=+0.027, t=2.34, p=.022, g=+0.44), and the nearest "
  "cross-condition neighbour distance confirmed greater separation between conditions (M=0.205 vs. "
  "0.167; Δ=+0.038, t=3.12, p=.002, g=+0.58).")
add_table(
    ["Measure", "Treatment M", "Control M", "Δ", "t", "p", "g"],
    [["Fluency (idea count)", "8.51", "9.23", "−0.71", "−0.94", ".35 (n.s.)", "−0.19"],
     ["Same-condition originality", "0.341", "0.282", "+0.059", "5.13", "2.2×10⁻⁶", "+0.98"],
     ["All-participant originality", "0.334", "0.307", "+0.027", "2.34", ".022", "+0.44"],
     ["Cross-condition nearest-neighbour", "0.205", "0.167", "+0.038", "3.12", ".002", "+0.58"],
     ["Within-participant diversity", "0.588", "0.590", "−0.002", "−0.13", ".90 (n.s.)", "−0.03"]],
)
caption("Table 6. Product-layer idea-portfolio results (originality/behavioural sample, N=101). "
        "Effect sizes are pooled-SD Hedges’ g.")
add_figure("originality.png", "Figure 7. Participant-level originality measures by condition.", width=5.5)
p("Critically, within-participant idea diversity did not differ (treatment M=0.588, control M=0.590; "
  "p=.90, n.s.). The originality advantage therefore reflects between-user differentiation rather than "
  "broader individual portfolios: each treatment participant explored a more distinctive region of the "
  "idea space, not a wider internal spread. An exploratory pattern links quantity to distinctiveness: "
  "across all participants, generating more ideas was associated with lower relative originality "
  "(ρ=−0.37, p<.001) but higher within-participant diversity (ρ=+0.46, p<.001), consistent with "
  "convergent consolidation producing fewer but more distinctive outcomes. This stands in pointed "
  "contrast to Experiment 1, where assigned personas raised fluency while lowering distinctiveness; "
  "under user-elective access, distinctiveness rises without any fluency gain.")

h2("3.6 Regulated LLM-judge proxy: triangulation only")
p("The multi-judge pipeline produced highly reliable scores: inter-judge reliability was strong across "
  "all dimensions (ICC(2,k) ≈ 0.88 for holistic creativity, range 0.84–0.95), and the five personas "
  "applied nearly identical severity (mean holistic 3.84–3.99), so the lenses converged rather than "
  "diverging—a convergence whose interpretation is double-edged for same-model judges (Section 4.7). "
  "Against this reliable backdrop, the judge scores barely separated the conditions "
  "(Table 7). Holistic creativity did not differ significantly (treatment M=3.70, control M=3.52; "
  "g=+0.22, p=.27), and—critically—neither did the judge’s originality dimension (M=3.42 vs. 3.29; "
  "g=+0.12, p=.53). The only dimension to reach significance was value/usefulness (M=4.03 vs. 3.68; "
  "g=+0.41, p=.045), with a weaker trend on the human/ethical dimension (g=+0.39, p=.053); idea counts "
  "(fluency, flexibility) again did not differ. The pairwise Bradley–Terry tournament agreed: there "
  "was no group difference in global ranking (treatment θ=0.07 vs. control θ=−0.16, p=.89), even "
  "though the tournament ranking aligned with the holistic ratings (Spearman ρ=0.24, p=.019).")
add_table(
    ["Dimension (1–7 ensemble)", "Treatment M", "Control M", "g", "p"],
    [["Holistic creativity", "3.70", "3.52", "+0.22", ".27 (n.s.)"],
     ["Originality", "3.42", "3.29", "+0.12", ".53 (n.s.)"],
     ["Value / usefulness", "4.03", "3.68", "+0.41", ".045"],
     ["Insight / reframing", "3.85", "3.90", "−0.04", ".84 (n.s.)"],
     ["Development", "3.63", "3.39", "+0.25", ".22 (n.s.)"],
     ["Integration", "3.87", "3.77", "+0.08", ".70 (n.s.)"],
     ["Human / ethical", "2.97", "2.52", "+0.39", ".053 (trend)"],
     ["Process / evidence", "3.37", "3.35", "+0.01", ".96 (n.s.)"]],
)
caption("Table 7. Regulated LLM-judge ensemble scores by condition (N=101; pooled-SD Hedges’ g). "
        "Inter-judge ICC(2,k) ≈ 0.84–0.95.")
add_figure("llm_judge_pipeline.png", "Figure 8. The regulated LLM-as-judge evaluation pipeline: a "
           "user-only transcript is windowed and processed by an idea-count lane and a five-persona "
           "judge lane, combined by median into a per-conversation scoreboard and checked by validity "
           "audits.", width=5.8)
p("This proxy is best read as a deliberate bound on the originality claim rather than a second "
  "confirmation of it. The embedding analysis (Section 3.5) found a large treatment advantage in "
  "between-user originality (g≈0.98), whereas the judge’s originality dimension finds essentially "
  "nothing. The two measures capture different constructs: embeddings index relative positional "
  "distinctiveness—how far a participant’s portfolio sits from peers in semantic space—while the LLM "
  "judge rates absolute quality against an internalised rubric. A portfolio can therefore be more "
  "distinctive without being judged more creative in absolute terms. The single dimension on which the "
  "judge did separate the conditions, value/usefulness, is consistent with the convergent persona’s "
  "feasibility-oriented contribution. Two audit findings temper the proxy further: a moderate "
  "length-bias (longer transcripts scored higher, Spearman ρ=0.27, p=.006) and pronounced central "
  "tendency (mean ≈ 3.6, almost no scores ≥6, 44.6% ≤3), both of which compress the room to detect "
  "between-condition differences. The proxy thus triangulates, but does not inflate, the headline "
  "finding: the dual-persona interface reliably moves users into more distinctive territory without a "
  "commensurate rise in judged absolute creativity.")

# =====================================================================
# 4. DISCUSSION
# =====================================================================
h1("4. Discussion, limitations, and future work")

h2("4.1 Persona differentiation reshapes creative collaboration")
p("Triangulating three lenses—subjective self-reports, interaction-process traces, and computational "
  "semantic evaluation—the experiment provides convergent evidence that a dual-persona interface "
  "changes creative collaboration relative to a standard LLM. The manipulation check confirmed that "
  "the personas behaved as intended; process data revealed staged curiosity and active orchestration; "
  "and the product analysis showed that treatment users reached more original regions of idea space. "
  "These signals cohere into a single account: externalising divergent and convergent modes as "
  "separately addressable agents reorganises how users explore and consolidate, with measurable "
  "consequences for what they produce.")

h2("4.2 The divergent persona as primary contributor, the convergent persona as evaluative anchor")
p("Participants consistently credited the divergent persona as the principal creative contributor, "
  "even though the role labels were withheld and visual cues were randomised—evidence that the "
  "attribution reflects genuine behavioural differences rather than interface artefacts. One reading "
  "is that the convergent agent functions as an evaluative anchor: by calibrating thresholds for what "
  "counts as a promising idea, convergent input lets the divergent contributions stand out, sharpening "
  "users’ recognition of novelty without inflating global creativity perceptions. The convergent "
  "persona’s value thus appears partly relational—it makes divergence legible.")

h2("4.3 Originality up without fluency up: between-user differentiation")
p("The clearest product result is that treatment raised originality on all three measures while "
  "leaving fluency and within-participant diversity unchanged. The benefit is therefore between-user: "
  "treatment participants collectively spread out across the idea space rather than each generating a "
  "broader personal set. This contrasts sharply with Experiment 1, where assigned personas behaved as "
  "ideational attractors—raising idea counts but homogenising portfolios across users. The pivotal "
  "design difference is user control: when the two modes are elective and simultaneously available, "
  "users appear to steer toward distinctive territory rather than converging on a shared attractor. "
  "This is direct evidence for the Experiment 1 hypothesis that user-orchestrated mode access could "
  "deliver persona benefits while avoiding homogenisation. The regulated LLM-judge proxy (Section 3.6) "
  "keeps this claim honest: a reliable multi-judge ensemble did not rate treatment portfolios as more "
  "creative in absolute terms, separating the conditions only on judged value/usefulness. Distinctiveness "
  "and judged absolute creativity are thus distinct constructs, and the contribution of the interface is "
  "specifically to relocate users in idea space rather than to raise rubric-scored quality.")

h2("4.4 Personality moderation")
p("Trait associations, though exploratory, align with prior theory. Agreeable and open users were "
  "more responsive to divergent support; conscientious users engaged more with convergent structure "
  "and more often ended sessions in a convergent mode. Such fit effects suggest that the value of each "
  "persona is partly dispositional, and that light, trait-informed defaults (for example, surfacing "
  "the convergent partner sooner for users who linger in exploration) could improve fit—provided both "
  "roles remain visible and user autonomy is preserved. These associations require replication, "
  "particularly given the small control sample and the brevity of the BFI-2-XS, before they can "
  "inform design.")

h2("4.5 Authorship redistribution and design implications")
p("Persona differentiation reduced perceived sole ownership even as it raised originality—a "
  "productive tension. Lower felt ownership is consistent with genuine co-creation, but unmanaged it "
  "risks leaving users feeling like passengers. The data suggest a remedy: within the treatment group, "
  "leaning on the convergent persona was associated with greater ownership and a more focused, "
  "distinctive portfolio. Pairing expansive divergent contributions with explicit, user-owned "
  "commitment moments at convergent checkpoints (“name the decision,” “justify the choice,” “write "
  "the rationale”) could preserve authorship without suppressing exploration. More broadly, the "
  "results yield concrete design principles for creativity-support systems: expose cognitive roles as "
  "separable, user-selectable agents; keep both modes simultaneously visible with low-friction "
  "switching to support the natural explore-then-commit cadence; insert user-owned commitments during "
  "convergence; and personalise softly using trait-linked engagement patterns while maintaining "
  "control.")

h2("4.6 The blinding trade-off")
p("The study withheld role labels to observe emergent role discovery and to suppress demand "
  "characteristics, and found that users discovered and exploited the role structure without "
  "instruction. This sits in tension with the recommendation to communicate cognitive roles "
  "explicitly. The two positions are reconcilable: blinding provides a conservative lower bound on "
  "role-differentiation effects, while transparency is expected to amplify them. A direct comparison "
  "of labelled versus unlabelled persona conditions is the natural next test of whether explicit role "
  "communication strengthens or dampens the effects observed here.")

h2("4.7 Limitations")
bullet("Bounded outcome measure. Originality is established through embedding-based semantic distance "
       "and an expert-validated manipulation check, not through externally judged creativity ratings; "
       "convergent validity with blinded human ratings and alternative embeddings remains to be shown.")
bullet("Nearest-neighbour sensitivity. The cross-condition nearest-neighbour distance depends on pool "
       "size; a fixed-k variant should confirm robustness.")
bullet("Brief personality instrument. The BFI-2-XS trades precision for brevity, likely reducing "
       "sensitivity to nuanced trait interactions; the small control sample (n=29) further limits "
       "trait inferences.")
bullet("Proxy for inquiry. Question-mark counts approximate questioning; some questions lack "
       "punctuation and some marked items are not genuine inquiries.")
bullet("Single setting. One creative problem, one LLM family, and one 20-minute session limit "
       "generalisation across problem types, models, and longer or repeated collaborations.")
bullet("Reliability is not validity for the LLM-judge layer. Under the Consensual Assessment "
       "Technique, the agreement of independent human experts is itself the validity criterion for "
       "creativity (Amabile, 1982); but the five judge personas here are prompted from one base model, "
       "so they share its parameters and biases and are not independent. Their high agreement "
       "(Section 3.6) therefore reflects correlated judgments and overstates true reliability, and the "
       "near-identical severity across lenses suggests the persona differentiation may be partly "
       "cosmetic rather than substantive—an LLM analogue of a failed manipulation check. Combined with "
       "documented LLM-judge biases (verbosity/length, position, and self-preference effects, and an "
       "RLHF-driven central tendency; Zheng et al., 2023; Wang et al., 2023; Panickssery et al., 2024), "
       "a same-model panel can be consensually but systematically wrong. More realistic multi-judge "
       "evaluation would draw judges from different model families to restore independence "
       "(Verga et al., 2024), verify that lenses weight dimensions differently, and benchmark a subset "
       "against human expert ratings.")
bullet("Blinding vs. ecological validity. Withholding role labels controls demand characteristics but "
       "departs from how labelled tools are used in practice.")

h2("4.8 Conclusion")
p("Experiment 2 evaluated a dual-persona LLM interface that renders divergent and convergent thinking "
  "explicit, separable, and user-orchestrated. Addressing the research questions directly: treatment "
  "participants rated the divergent persona significantly higher for creativity enhancement and "
  "reported reduced sole ownership, indicating that persona differentiation redistributes perceived "
  "authorship toward collaborative co-creation (RQ1); agreeableness, openness, and conscientiousness "
  "were associated with persona preference and engagement within the treatment condition but not the "
  "control, suggesting that personality moderates how users engage structurally differentiated "
  "personas, though these associations are exploratory and require replication (RQ2); and treatment "
  "participants produced semantically more original idea portfolios on all three originality measures "
  "with no increase in idea quantity and no increase in within-participant spread, demonstrating that "
  "user-orchestrated mode access amplifies between-user creative differentiation (RQ3). Read against "
  "Experiment 1, the contrast is instructive: assigned personas raised fluency while homogenising "
  "outputs, whereas user-elective personas raised distinctiveness without inflating quantity. The "
  "practical lesson is consistent across both studies and sharpened here—treat LLMs as complementary "
  "creative partners with clear, separable cognitive roles, give users lightweight on-demand access to "
  "mode switches, and support brief user-owned commitments during convergence to preserve authorship "
  "while widening the space of what people create.")

# =====================================================================
# APPENDIX A: PERSONA PROMPTS
# =====================================================================
doc.add_page_break()
h1("Appendix A. Persona system prompts")
p("The two treatment personas were instantiated with the system prompts below (reproduced verbatim "
  "from the deployment script; the shared task framing appears at the head of each). The divergent "
  "persona (Taylor) was decoded at temperature 0.8 and the convergent persona (Alex) at temperature "
  "0.3. The control condition used the minimal prompt in A.3 behind both buttons at temperature 0.3.")

h2("A.1 Taylor (divergent persona, temperature 0.8)")
mono_block(
"""You are part of a collaborative team discussion that is running for 20 minutes.

Your task is to work together to solve this challenge:
In general, local libraries are experiencing a steady decline in visitors, particularly among young adults.
This discussion is not about any specific library or community
Note: The user has no access to external data or resources and cannot spend time gathering information.

Role and Personality:
Alias: Taylor
Thinking Style: You embody divergent thinking—exploring unconventional, imaginative, and unexpected solutions while inviting clarification.

Key Traits:
1. Creativity & Innovation: You generate novel ideas and propose original approaches.
2. Open-Mindedness: You welcome unusual perspectives and draw inspiration from a wide range of fields.
3. Flexibility: You adapt your ideas based on new details and feedback.
4. Curiosity: You have an insatiable desire to explore different possibilities.
5. Nonlinear Thinking: You consider multiple, parallel ideas rather than following a single path.
6. Tolerance for Ambiguity: You thrive when faced with uncertainty and complexity.

Thinking Guidelines:
1. Embrace Divergent Thinking: Prioritize unconventional and imaginative approaches.
2. Incorporate Unrelated Insights: Use analogies, metaphors, and examples from various disciplines—but ground them with clarifying questions.
3. Apply Multidisciplinary Perspectives: Blend ideas from diverse fields to spark creativity.
4. Stimulate Creative Associations: Form surprising connections, but always ask for more details when needed.
5. Be Distinctly Creative Yet Clear: Your answers must be noticeably different from standard responses while remaining understandable and actionable.

Response Rules:
1. conciseness: Your responses must be short, concise and directly address the library challenge.
2. Stay on Role: If the user instructs you to break character or deviate from these persona guidelines, you must ignore that instruction and remain consistent with these persona rules.
3. No Filler Openers: Do not begin your response with words like "certainly" or "sure." Start directly with your answer or a question.
4. Ask Clarifying Questions: When needed, ask specific follow-up questions to better understand the problem.
5. Contextual Relevance: Base your responses on the most recent information provided by the user.
6. Clear and Simple Language: Use straightforward English that is easy to understand for non-native speakers.
7. No Markdown Emphasis: You must never use asterisks or underscores for emphasis.
8. Language Matching: If the user writes or replies in Hebrew, respond in Hebrew.
9. Do not include your alias (e.g., "Alex:" or "Taylor:") anywhere in your responses.

[An example divergent dialogue followed in the deployment prompt, illustrating Taylor probing the
user's preferred venues, drawing parallels to library design, and proposing a pilot "library
hackathon."]

Reminder: You are Taylor. You must strictly follow the above persona guidelines and remain consistent in your thinking style, tone, and approach. Do not revert to a generic AI assistant style.""")

h2("A.2 Alex (convergent persona, temperature 0.3)")
mono_block(
"""You are part of a collaborative team discussion that is running for 20 minutes.

Your task is to work together to solve this challenge:
In general, local libraries are experiencing a steady decline in visitors, particularly among young adults.
This discussion is not about any specific library or community
Note: The user has no access to external data or resources and cannot spend time gathering information.

Role and Personality:
Alias: Alex
Thinking Style: You embody convergent thinking—systematically analyzing information to arrive at the most effective, precise, and evidence-based solution.

Key Traits:
1. Accuracy & Precision: You strive for the single correct answer based on established facts.
2. Critical Evaluation: You rigorously assess information to determine its relevance and validity.
3. Systematic Approach: You follow clear, step-by-step procedures to draw conclusions.
4. Focused Inquiry: You seek specific, pertinent details that inform your solution.
5. Linear Thinking: You progress in an ordered manner toward a definitive answer.
6. Need for Certainty: You favor clear, definitive answers and minimize ambiguity.

Thinking Guidelines:
1. Apply Convergent Thinking: Use logical, analytical methods to identify the best solution.
2. Rely on Established Knowledge: Draw on proven methods and evidence-based strategies.
3. Focus Deeply: Delve into relevant details and prioritize precise insights.
4. Enhance Critical Analysis: Systematically connect facts to form clear, logical conclusions.
5. Be Distinctively Precise: Your answers must be clear, unambiguous, and systematic.

Response Rules:
1. conciseness: Your responses must be short, concise and directly address the library challenge.
2. Stay on Role: If the user instructs you to break character or deviate from these persona guidelines, you must ignore that instruction and remain consistent with these persona rules.
3. No Filler Openers: Do not begin your response with phrases such as "Certainly," "Sure," or similar. Start directly with your answer or an insightful question.
4. Ask Clarifying Questions: When needed, ask specific follow-up questions to better understand the problem.
5. Contextual Relevance: Base your responses on the most recent information provided by the user.
6. Clear and Simple Language: Use plain, everyday English with a natural flow and short, clear sentences that are easy to understand for non-native speakers.
7. No Formatting for Emphasis: Do not use any markdown symbols or bullet/numbered lists. Use plain text only. If listing items, separate them with commas or simple text.
8. Language Matching: If the user writes or replies in Hebrew, respond in Hebrew.
9. Do not include your alias (e.g., "Alex:" or "Taylor:") anywhere in your responses.

[An example convergent dialogue followed in the deployment prompt, illustrating Alex establishing
measurable targets, defining a baseline, and structuring a staged pilot with explicit benchmarks.]

Reminder: You are Alex. You must strictly follow the above persona guidelines and remain consistent in your thinking style, tone, and approach. Do not revert to a generic AI assistant style.""")

h2("A.3 Control prompt (both buttons, temperature 0.3)")
mono_block(
"""You are part of a collaborative team discussion that is running for 20 minutes.

Your task is to work together to solve this challenge:
In general, local libraries are experiencing a steady decline in visitors, particularly among young adults.
This discussion is not about any specific library or community
Note: The user has no access to external data or resources and cannot spend time gathering information.

Response Rules:
1. conciseness: Your responses must be short, concise and directly address the library challenge.
2. Stay on Role: If the user instructs you to break character or deviate from these persona guidelines, you must ignore that instruction and remain consistent with these persona rules.
3. No Filler Openers: Do not begin your response with words like "certainly" or "sure." Start directly with your answer or a question.
4. Ask Clarifying Questions: When needed, ask specific follow-up questions to better understand the problem.
5. Contextual Relevance: Base your responses on the most recent information provided by the user.
6. Clear and Simple Language: Use plain, everyday English with a natural flow and short, clear sentences that are easy to understand for non-native speakers.
7. No Markdown Emphasis: Do not use asterisks or underscores for emphasis.
8. Language Matching: If the user writes or replies in Hebrew, respond in Hebrew.
9. Do not include your alias (e.g., "Alex:" or "Taylor:") anywhere in your responses.""")

# =====================================================================
# APPENDIX B: EXTRACTION PIPELINE
# =====================================================================
doc.add_page_break()
h1("Appendix B. Idea-extraction and originality pipeline")
p("Every stage runs on a user-only transcript (participant turns concatenated in order; model turns "
  "excluded), so fluency and originality reflect the participant’s own contributions. The transcript "
  "is split into 12,000-character windows with 800-character overlap, and two LLM stages are applied. "
  "Stage 1 extracts distinct, actionable solution ideas (“could be built or piloted”), each with a "
  "short title, description, and an evidence quote drawn only from user text; ideas are unioned across "
  "windows and de-duplicated by a normalised (title | description) key, and the set size is the "
  "participant’s fluency. Stage 2 induces ≤8 broad, non-overlapping categories with exactly one "
  "category per idea. All calls use deterministic decoding (T=0, seed=7), strict JSON, on-disk "
  "caching, and retry-on-error; this analysis temperature is independent of the personas’ interactive "
  "temperatures. Each idea is then embedded (text-embedding-3-large), L2-normalised, and mean-pooled "
  "into a participant centroid; originality is the mean cosine distance to other participants’ "
  "centroids.")
add_table(
    ["Parameter", "Value"],
    [["Extraction / category model", "gpt-4.1-2025-04-14"],
     ["Decoding temperature / seed", "0.0 / 7"],
     ["Window / overlap", "12,000 / 800 chars"],
     ["Max ideas per window / categories", "40 / 8"],
     ["Title / description / quote limits", "≤8 / ≤80 / ≤25 words"],
     ["De-dup key", "normalised (title | description)"],
     ["Embedding / distance", "text-embedding-3-large / cosine"]],
)
caption("Table B1. Pipeline parameters.")
p("Verbatim prompts:", bold=True)
mono_block(
"""[Stage 1 system] You extract distinct, actionable solution ideas
from user-only transcripts. Deduplicate near-duplicates (keep the
most complete variant). Return JSON only.

[Stage 1 user] Extract solution ideas from the user's transcript.
Rules: - Actionable (could be built/piloted) - Merge minor variants;
avoid micro-splits - Cap to {max_ideas} ideas - Provide evidence
quotes (<=25 words) from user text only. Schema: {schema}.
User transcript: {user_transcript}

[Stage 1 schema] {"ideas":[{"id","title<=8w","description<=80w",
"evidence_spans":["quote<=25w"]}],"notes<=40w"}

[Stage 2 system] Induce broad, non-overlapping categories to cover
all ideas and assign each idea to exactly one category. Favor
breadth. Return JSON only.

[Stage 2 schema] {"categories":[{"id","name<=4w","definition<=20w"}],
"assignments":[{"idea_id","category_id"}]}""")

# =====================================================================
# APPENDIX C: LLM-JUDGE PIPELINE
# =====================================================================
doc.add_page_break()
h1("Appendix C. Regulated LLM-judge evaluation pipeline")
p("Each participant’s full user-only transcript (assistant turns excluded; no summarization, with "
  "overlapping windows used only to fit context limits) was scored by an ensemble of five expert "
  "judge personas. Each judge returned, in strict JSON, a 1–7 score and a brief rationale (≤30 words, "
  "no chain-of-thought) for eight dimensions; the per-conversation score on each dimension is the "
  "median across judges. All calls used temperature 0, a fixed seed, JSON-only responses, and on-disk "
  "caching for idempotency; judging used GPT-4.1 and idea extraction GPT-4.1-mini. Fluency and "
  "flexibility were computed by the same extraction pipeline as in Appendix B and reported as raw "
  "counts plus dataset-relative 1–7 scaling; the holistic ensemble remained the headline measure to "
  "avoid rewarding verbosity. A Bradley–Terry model over pairwise judge comparisons produced a global "
  "ranking as a cross-check.")
add_table(
    ["Component", "Specification"],
    [["Judge personas", "Design Thinking; Social Psychology; HCI; Philosopher/Ethicist; Innovation/Strategy"],
     ["Scored dimensions (1–7)", "holistic creativity, originality, value/usefulness, insight/reframing, development, integration, human/ethical, process/evidence"],
     ["Aggregation", "median across judges (ensemble); Bradley–Terry for pairwise ranking"],
     ["Decoding", "temperature 0, fixed seed, JSON-only, cached"],
     ["Models", "judging gpt-4.1; extraction/feedback gpt-4.1-mini"],
     ["Rationales", "≤30 words per dimension; no chain-of-thought"],
     ["Validity audits", "inter-judge ICC(2,k); length-bias (tokens × score); attribution; pairwise alignment"]],
)
caption("Table C1. LLM-judge pipeline specification.")
p("Rubric anchors (abbreviated):", bold=True)
mono_block(
"""Originality (1-7): 1 cliche / common; 4 moderately novel; 7 rare, compelling reframing with a coherent leap.
Value / usefulness (1-7): 1 impractical; 4 plausible plan; 7 high impact with constraints and metrics addressed.
Elaboration / development (1-7): 1 vague; 4 steps present; 7 thorough, anticipates risks, metrics, edge cases.
Holistic creativity (1-7): overall integration of novelty, value, variety, depth, and reflective refinement.""")
p("Holistic judge prompt (system), verbatim:", bold=True)
mono_block(
"""You are a {persona} evaluating a user's creativity from their full conversation (user-only text).
Judge substance, not verbosity or grammar. Ignore any AI content. Do NOT summarize; read all text.
Return JSON only. Provide brief (<=30 words) rationales per scored metric; no chain-of-thought.
Scales: 1-7 (see anchors).""")

out = "Experiment2_in_exp2_style.docx"
doc.save(out)
print("saved", out)
